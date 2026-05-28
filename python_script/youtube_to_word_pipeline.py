"""
YouTube Installation Video → Word Document Pipeline  (template-style Word layout)

Dark client template (logo top-left, coral title, QR/video panel, red rule) or light mode via DOC_DARK_THEME,
then overview/tools; each step shows instructions and screenshot side by side.

Pipeline = v13 (video pipeline, sharp screenshots, etc.) with this Word layout only.

Run:
  python youtube_to_word_pipeline.py "https://www.youtube.com/watch?v=VIDEO_ID" output --verbose
  python youtube_to_word_pipeline.py "https://..." output --doc-only

Env: LOGO_PATH, LOGO_WIDTH_INCHES, DELIVERABLE_OUTPUT_DIR, DOC_STEP_*_INCHES,
FRAME_EXTRACT_WIDTH, FINAL_SCREENSHOT_WIDTH, FRAME_FINAL_ACCURATE_SEEK, SCREENSHOT_BORDER_PX,
CONTACT_*, TEMPLATE_SUBTITLE, TEMPLATE_WARNING_TEXT, TEMPLATE_IMPORTANT_NOTE,
VISION_CONFIRM_MAX_DROP, VISION_CONFIRM_FLOOR, VISION_SHORTLIST_FRAMES,
MAX_TRANSCRIPT_WINDOW_SECONDS, VIDEO_OUTRO_EXCLUSION_SECONDS,
LOW_CONFIDENCE_REFINE_THRESHOLD, ENABLE_LOW_CONFIDENCE_REFINE,
YT_DLP_TITLE_TIMEOUT_SEC

Dependencies: anthropic openai python-docx pillow python-dotenv pydantic,
yt-dlp, ffmpeg; optional qrcode[pil], scipy. (google-genai is no longer
required for the default pipeline — OpenAI handles transcription + vision.)
"""

import argparse
import asyncio
import datetime
import json
import logging
import math
import os
import re
import shlex
import shutil
import subprocess
import sys
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Optional

from dotenv import load_dotenv
from pydantic import BaseModel, Field, ValidationError, field_validator

APP_ROOT = Path(__file__).resolve().parent


def _resolve_app_path(raw: str) -> str:
    """Resolve relative paths against this package dir (stable under pm2/systemd)."""
    if not raw:
        return raw
    p = Path(raw)
    return str(p.resolve()) if p.is_absolute() else str((APP_ROOT / p).resolve())


load_dotenv(APP_ROOT / ".env")

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────

CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-20250514")
CLAUDE_MAX_TOKENS = int(os.getenv("CLAUDE_MAX_TOKENS", "16000"))
# Deterministic step extraction: temperature 0 keeps Claude from collapsing 9
# atomic steps into 5 mega-steps on alternating runs of the same video.
CLAUDE_TEMPERATURE = float(os.getenv("CLAUDE_TEMPERATURE", "0"))

GEMINI_TRANSCRIBE_MODEL = os.getenv("GEMINI_TRANSCRIBE_MODEL", "gemini-2.5-flash")
GEMINI_VISION_MODEL = os.getenv("GEMINI_VISION_MODEL", "gemini-2.5-flash")
GEMINI_ENABLE_MODEL_FALLBACK = os.getenv("GEMINI_ENABLE_MODEL_FALLBACK", "true").lower() == "true"
# Transient failures on primary before trying the next model (last model uses max_retries).
GEMINI_RETRIES_BEFORE_FALLBACK = int(os.getenv("GEMINI_RETRIES_BEFORE_FALLBACK", "2"))
_DEFAULT_GEMINI_FALLBACK = "gemini-2.0-flash"


def _parse_gemini_model_chain(primary: str, env_key: str) -> list[str]:
    """Primary model first, then comma-separated fallbacks (deduped)."""
    raw = os.getenv(env_key, _DEFAULT_GEMINI_FALLBACK)
    chain = [primary.strip()]
    for part in raw.split(","):
        m = part.strip()
        if m and m not in chain:
            chain.append(m)
    return chain


GEMINI_TRANSCRIBE_MODELS = _parse_gemini_model_chain(
    GEMINI_TRANSCRIBE_MODEL, "GEMINI_TRANSCRIBE_FALLBACK_MODELS",
)
GEMINI_VISION_MODELS = _parse_gemini_model_chain(
    GEMINI_VISION_MODEL, "GEMINI_VISION_FALLBACK_MODELS",
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY", "")

# ── OpenAI (default provider for transcription + vision; replaces Gemini) ───
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
# Primary transcription model. whisper-1 returns segments with timestamps via
# verbose_json. The newer gpt-4o*-transcribe models are text-only (no segments),
# so we transparently fall back to synthetic segments when they are picked.
OPENAI_TRANSCRIBE_MODEL = os.getenv("OPENAI_TRANSCRIBE_MODEL", "whisper-1")
OPENAI_VISION_MODEL = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")
# Comma-separated fallback chain (tried in order on quota / availability errors)
OPENAI_TRANSCRIBE_FALLBACK_MODELS = os.getenv(
    "OPENAI_TRANSCRIBE_FALLBACK_MODELS", "gpt-4o-mini-transcribe",
)
OPENAI_VISION_FALLBACK_MODELS = os.getenv(
    "OPENAI_VISION_FALLBACK_MODELS", "gpt-4o-mini",
)
OPENAI_ENABLE_MODEL_FALLBACK = os.getenv("OPENAI_ENABLE_MODEL_FALLBACK", "true").lower() == "true"
OPENAI_RETRIES_BEFORE_FALLBACK = int(os.getenv("OPENAI_RETRIES_BEFORE_FALLBACK", "2"))


def _parse_openai_model_chain(primary: str, fallback_raw: str) -> list[str]:
    chain = [primary.strip()]
    for part in (fallback_raw or "").split(","):
        m = part.strip()
        if m and m not in chain:
            chain.append(m)
    return chain


OPENAI_TRANSCRIBE_MODELS = _parse_openai_model_chain(
    OPENAI_TRANSCRIBE_MODEL, OPENAI_TRANSCRIBE_FALLBACK_MODELS,
)
OPENAI_VISION_MODELS = _parse_openai_model_chain(
    OPENAI_VISION_MODEL, OPENAI_VISION_FALLBACK_MODELS,
)

OPENAI_VISION_TIMEOUT_SEC = float(os.getenv("OPENAI_VISION_TIMEOUT_SEC",
                                            os.getenv("GEMINI_VISION_TIMEOUT_SEC", "180")))
OPENAI_VISION_MAX_RETRIES = int(os.getenv("OPENAI_VISION_MAX_RETRIES",
                                          os.getenv("GEMINI_VISION_MAX_RETRIES", "5")))
OPENAI_VISION_BASE_DELAY = float(os.getenv("OPENAI_VISION_BASE_DELAY",
                                           os.getenv("GEMINI_VISION_BASE_DELAY", "6.0")))
# Image detail for vision: "low" (~85 tok, fast) | "high" (full-res, slow) |
# "auto" (OpenAI picks). "auto" is the safest default for installation frames.
OPENAI_VISION_IMAGE_DETAIL = os.getenv("OPENAI_VISION_IMAGE_DETAIL", "auto").lower()
if OPENAI_VISION_IMAGE_DETAIL not in {"low", "high", "auto"}:
    OPENAI_VISION_IMAGE_DETAIL = "auto"
# Max audio bytes per OpenAI Whisper request (real cap is 25 MB)
MAX_OPENAI_AUDIO_BYTES = int(os.getenv("MAX_OPENAI_AUDIO_BYTES", str(24 * 1024 * 1024)))

_OPENAI_CLIENT_SINGLETON: Any = None


def _openai_client() -> Any:
    """Lazy-init the OpenAI client so the module imports cleanly without the key."""
    global _OPENAI_CLIENT_SINGLETON
    if _OPENAI_CLIENT_SINGLETON is not None:
        return _OPENAI_CLIENT_SINGLETON
    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError(
            "openai package not installed. Run: pip install 'openai>=1.50'"
        ) from exc
    if not OPENAI_API_KEY:
        raise RuntimeError(
            "OPENAI_API_KEY is not set. Add it to .env to use OpenAI vision/transcription."
        )
    _OPENAI_CLIENT_SINGLETON = OpenAI(api_key=OPENAI_API_KEY)
    return _OPENAI_CLIENT_SINGLETON

# v12: Logo path. ONLY this env var is used — no directory scanning.
# Falls back to generated placeholder if unset or file fails to load.
LOGO_PATH = _resolve_app_path(os.getenv("LOGO_PATH", ""))

# v12: Single consistent width for both real logos and placeholders.
LOGO_WIDTH_INCHES = float(os.getenv("LOGO_WIDTH_INCHES", "2.5"))

# ── Frame sampling ──────────────────────────
VISUAL_PRE_ROLL = float(os.getenv("VISUAL_PRE_ROLL", "0.5"))
VISUAL_POST_ROLL = float(os.getenv("VISUAL_POST_ROLL", "5.0"))
VISUAL_SAMPLE_INTERVAL = float(os.getenv("VISUAL_SAMPLE_INTERVAL", "0.35"))

MAX_CANDIDATE_FRAMES = int(os.getenv("MAX_CANDIDATE_FRAMES", "20"))
UNIFORM_SAMPLE_RESERVED_SLOTS = int(os.getenv("UNIFORM_SAMPLE_RESERVED_SLOTS", "6"))

# ── Frame quality thresholds ──────────────────
SHARPNESS_MIN = float(os.getenv("SHARPNESS_MIN", "80.0"))
BRIGHTNESS_MIN = float(os.getenv("BRIGHTNESS_MIN", "30.0"))
BRIGHTNESS_MAX = float(os.getenv("BRIGHTNESS_MAX", "230.0"))
MIN_FRAMES_AFTER_QUALITY_FILTER = int(os.getenv("MIN_FRAMES_AFTER_QUALITY_FILTER", "6"))

# ── Color cast detection ─────────────────────
COLOR_CAST_MAX_RATIO = float(os.getenv("COLOR_CAST_MAX_RATIO", "1.45"))
COLOR_CAST_MIN_BRIGHTNESS = float(os.getenv("COLOR_CAST_MIN_BRIGHTNESS", "15.0"))

# Prototype banner detection
# v10 FIX: Tightened thresholds to avoid false positives on dark interior shots.
PROTOTYPE_BANNER_CHECK = os.getenv("PROTOTYPE_BANNER_CHECK", "true").lower() == "true"
PROTOTYPE_BANNER_TOP_FRACTION = float(os.getenv("PROTOTYPE_BANNER_TOP_FRACTION", "0.15"))
PROTOTYPE_BANNER_MEAN_MAX = float(os.getenv("PROTOTYPE_BANNER_MEAN_MAX", "40.0"))
PROTOTYPE_BANNER_CONTRAST_THRESHOLD = float(os.getenv("PROTOTYPE_BANNER_CONTRAST_THRESHOLD", "90.0"))

# ── Reinstall step search offset ─────────────
_REINSTALL_KW_ENV = os.getenv(
    "REINSTALL_KEYWORDS",
    "reinstall,replace,reattach,refit,remount,put back,bolt back,"
    "screw back,clip back,re-install,re-attach,reassemble,re-assemble,"
    "refasten,retighten,reseat",
)
REINSTALL_KEYWORDS: frozenset[str] = frozenset(
    k.strip().lower() for k in _REINSTALL_KW_ENV.split(",") if k.strip()
)
REINSTALL_EXTRA_FORWARD_SEARCH = float(os.getenv("REINSTALL_EXTRA_FORWARD_SEARCH", "18.0"))
REINSTALL_TRIGGER_GAP_SECONDS = float(os.getenv("REINSTALL_TRIGGER_GAP_SECONDS", "12.0"))

# Multi-person step keywords for Pass 2.5 rescue
_MULTI_PERSON_KW_ENV = os.getenv(
    "MULTI_PERSON_KEYWORDS",
    "friend,two people,both people,second person,helper,partner,together,with help",
)
MULTI_PERSON_KEYWORDS: frozenset[str] = frozenset(
    k.strip().lower() for k in _MULTI_PERSON_KW_ENV.split(",") if k.strip()
)
MULTI_PERSON_RESCUE_CONFIDENCE = float(os.getenv("MULTI_PERSON_RESCUE_CONFIDENCE", "0.80"))
MULTI_PERSON_RESCUE_FRAMES = int(os.getenv("MULTI_PERSON_RESCUE_FRAMES", "20"))

# ── Consecutive duplicate detection ──────────
CONSECUTIVE_FRAME_SIMILARITY_MAX = float(
    os.getenv("CONSECUTIVE_FRAME_SIMILARITY_MAX", "0.94")
)
HISTOGRAM_BINS = int(os.getenv("HISTOGRAM_BINS", "32"))

# ── Step extraction quality gate ─────────────
MIN_STEP_GAP_SECONDS = float(os.getenv("MIN_STEP_GAP_SECONDS", "12.0"))

# ── Vision confidence thresholds ─────────────
VISION_MIN_CONFIDENCE = float(os.getenv("VISION_MIN_CONFIDENCE", "0.70"))
VISION_CONFIRM_THRESHOLD = float(os.getenv("VISION_CONFIRM_THRESHOLD", "0.75"))
VISION_RETRY_MIN_CONFIDENCE = float(os.getenv("VISION_RETRY_MIN_CONFIDENCE", "0.50"))
# Max allowed drop in confidence after borderline confirmation when wrong_step is false.
VISION_CONFIRM_MAX_DROP = float(os.getenv("VISION_CONFIRM_MAX_DROP", "0.15"))
# Borderline frames in [VISION_CONFIRM_FLOOR, VISION_CONFIRM_THRESHOLD) get a second look.
VISION_CONFIRM_FLOOR = float(os.getenv("VISION_CONFIRM_FLOOR", "0.48"))
# Post-dedup pass: re-sample steps below this confidence with a denser candidate window.
LOW_CONFIDENCE_REFINE_THRESHOLD = float(os.getenv("LOW_CONFIDENCE_REFINE_THRESHOLD", "0.80"))
LOW_CONFIDENCE_REFINE_FRAMES = int(os.getenv("LOW_CONFIDENCE_REFINE_FRAMES", "28"))
ENABLE_LOW_CONFIDENCE_REFINE = os.getenv("ENABLE_LOW_CONFIDENCE_REFINE", "true").lower() == "true"
LOW_CONFIDENCE_REFINE_MIN_GAIN = float(os.getenv("LOW_CONFIDENCE_REFINE_MIN_GAIN", "0.05"))

# Cap oversized Claude segment citations so frame search stays on the spoken action.
MAX_TRANSCRIPT_WINDOW_SECONDS = float(os.getenv("MAX_TRANSCRIPT_WINDOW_SECONDS", "35.0"))
MAX_SOURCE_SEGMENT_SPAN_SECONDS = float(os.getenv("MAX_SOURCE_SEGMENT_SPAN_SECONDS", "45.0"))
# Exclude final N seconds from candidate pools (sponsor/outro slates).
VIDEO_OUTRO_EXCLUSION_SECONDS = float(os.getenv("VIDEO_OUTRO_EXCLUSION_SECONDS", "15.0"))
# After quality filter, send only the best-ranked frames to Gemini (clearer comparisons).
VISION_SHORTLIST_FRAMES = int(os.getenv("VISION_SHORTLIST_FRAMES", "16"))
VISION_RESCUE_SHORTLIST_FRAMES = int(os.getenv("VISION_RESCUE_SHORTLIST_FRAMES", "10"))
GEMINI_VISION_TIMEOUT_SEC = float(os.getenv("GEMINI_VISION_TIMEOUT_SEC", "120"))
GEMINI_VISION_MAX_RETRIES = int(os.getenv("GEMINI_VISION_MAX_RETRIES", "5"))
GEMINI_VISION_BASE_DELAY = float(os.getenv("GEMINI_VISION_BASE_DELAY", "6.0"))
# Accept a borderline normal pass instead of multi-person + wide rescue (saves API time).
VISION_SOFT_ACCEPT_CONFIDENCE = float(os.getenv("VISION_SOFT_ACCEPT_CONFIDENCE", "0.38"))
VISION_SKIP_EXTENDED_ON_API_ERROR = (
    os.getenv("VISION_SKIP_EXTENDED_ON_API_ERROR", "true").lower() == "true"
)
# Stricter vision + no "manual review" flags — for client-facing demo deliverables.
CLIENT_DEMO_MODE = os.getenv("CLIENT_DEMO_MODE", "false").lower() == "true"
ENABLE_VISION_CONFIRMATION = os.getenv("ENABLE_VISION_CONFIRMATION", "true").lower() == "true"
# How far back to search when a step timestamp lands in the outro / past content_end.
OUTRO_STEP_LOOKBACK_SECONDS = float(os.getenv("OUTRO_STEP_LOOKBACK_SECONDS", "90.0"))
# Clamp step timestamps that exceed content_end by more than this (seconds).
STEP_CLAMP_PAST_CONTENT_SECONDS = float(os.getenv("STEP_CLAMP_PAST_CONTENT_SECONDS", "2.0"))

_OUTRO_TITLE_KW_ENV = os.getenv(
    "OUTRO_STEP_TITLE_KEYWORDS",
    "warning,caution,disclaimer,legal,contact,subscribe,thank you,outro,review",
)
OUTRO_STEP_TITLE_KEYWORDS: frozenset[str] = frozenset(
    k.strip().lower() for k in _OUTRO_TITLE_KW_ENV.split(",") if k.strip()
)

# yt-dlp --get-title can exceed 30s on slow networks; override via env.
YT_DLP_TITLE_TIMEOUT_SEC = int(os.getenv("YT_DLP_TITLE_TIMEOUT_SEC", "120"))
YT_DLP_RETRIES = int(os.getenv("YT_DLP_RETRIES", "10"))
YT_DLP_FRAGMENT_RETRIES = int(os.getenv("YT_DLP_FRAGMENT_RETRIES", "10"))
YT_DLP_SOCKET_TIMEOUT = int(os.getenv("YT_DLP_SOCKET_TIMEOUT", "30"))
# Try smaller 720p downloads before 1080p (~half the bytes on many install videos).
YT_DLP_PREFER_720_FIRST = os.getenv("YT_DLP_PREFER_720_FIRST", "true").lower() == "true"
# YouTube auth on datacenter IPs — cookies file is the primary fix (see DEPLOY.md).
YT_DLP_BIN = os.getenv("YT_DLP_BIN", "yt-dlp").strip() or "yt-dlp"
YT_DLP_COOKIES_FILE = os.getenv("YT_DLP_COOKIES_FILE", "").strip()
YT_DLP_PROXY = os.getenv("YT_DLP_PROXY", "").strip()
YT_DLP_REMOTE_COMPONENTS = os.getenv("YT_DLP_REMOTE_COMPONENTS", "").strip()
YT_DLP_PLAYER_CLIENT = os.getenv("YT_DLP_PLAYER_CLIENT", "").strip()
YT_DLP_EXTRACTOR_ARGS = os.getenv("YT_DLP_EXTRACTOR_ARGS", "").strip()
YT_DLP_EXTRA_ARGS = os.getenv("YT_DLP_EXTRA_ARGS", "").strip()

# ── Rescue pass ──────────────────────────────
WIDE_RESCUE_PRE_ROLL = float(os.getenv("WIDE_RESCUE_PRE_ROLL", "20.0"))
WIDE_RESCUE_POST_ROLL = float(os.getenv("WIDE_RESCUE_POST_ROLL", "40.0"))
WIDE_RESCUE_FRAMES = int(os.getenv("WIDE_RESCUE_FRAMES", "16"))

# ── Extended fallback pass ───────────────────
EXTENDED_FALLBACK_PRE_ROLL = float(os.getenv("EXTENDED_FALLBACK_PRE_ROLL", "30.0"))
EXTENDED_FALLBACK_POST_ROLL = float(os.getenv("EXTENDED_FALLBACK_POST_ROLL", "30.0"))
EXTENDED_FALLBACK_FRAMES = int(os.getenv("EXTENDED_FALLBACK_FRAMES", "16"))

# ── Whole-video rescue pass (last resort before forced) ──────────
# When normal + wide-rescue + extended-fallback all fail, scan frames uniformly
# across the whole content range using the step's visual_query. This catches
# cases where Claude's transcript anchor lands on the wrong scene (e.g., outro
# voice-over describing an action that was shown earlier).
WHOLE_VIDEO_RESCUE_ENABLED = os.getenv("WHOLE_VIDEO_RESCUE_ENABLED", "true").lower() == "true"
WHOLE_VIDEO_RESCUE_FRAMES = int(os.getenv("WHOLE_VIDEO_RESCUE_FRAMES", "24"))
WHOLE_VIDEO_RESCUE_MIN_CONFIDENCE = float(os.getenv("WHOLE_VIDEO_RESCUE_MIN_CONFIDENCE", "0.65"))

# ── Scene detection ──────────────────────────
SCENE_THRESHOLD = float(os.getenv("SCENE_THRESHOLD", "0.35"))
SCENE_DETECTION_TIMEOUT_SECONDS = int(os.getenv("SCENE_DETECTION_TIMEOUT_SECONDS", "120"))

# ── General ──────────────────────────────────
ALWAYS_FILL_SCREENSHOTS = os.getenv("ALWAYS_FILL_SCREENSHOTS", "true").lower() == "true"
ALLOW_PARTIAL_VISUAL_MATCH = os.getenv("ALLOW_PARTIAL_VISUAL_MATCH", "true").lower() == "true"
PARTIAL_MATCH_MIN_CONFIDENCE = float(os.getenv("PARTIAL_MATCH_MIN_CONFIDENCE", "0.40"))
STRICT_SCREENSHOT_MODE = os.getenv("STRICT_SCREENSHOT_MODE", "false").lower() == "true"
DOC_INCLUDE_LOW_CONFIDENCE_SCREENSHOTS = (
    os.getenv("DOC_INCLUDE_LOW_CONFIDENCE_SCREENSHOTS", "true").lower() == "true"
)
FINAL_FALLBACK_STRATEGY = os.getenv("FINAL_FALLBACK_STRATEGY", "midpoint").lower()

SCREENSHOT_CONCURRENCY = int(os.getenv("SCREENSHOT_CONCURRENCY", "5"))
SAFE_END_PADDING_SECONDS = float(os.getenv("SAFE_END_PADDING_SECONDS", "2.0"))

MAX_GEMINI_AUDIO_BYTES = int(os.getenv("MAX_GEMINI_AUDIO_BYTES", str(20 * 1024 * 1024)))
AUDIO_CHUNK_SECONDS = int(os.getenv("AUDIO_CHUNK_SECONDS", "600"))

H264_PROCESSING_WIDTH = int(os.getenv("H264_PROCESSING_WIDTH", "1280"))

# v14: fast candidate extraction vs. high-res finals written to screenshots/
FRAME_EXTRACT_WIDTH = int(os.getenv("FRAME_EXTRACT_WIDTH", str(H264_PROCESSING_WIDTH)))
FINAL_SCREENSHOT_WIDTH = int(os.getenv("FINAL_SCREENSHOT_WIDTH", "1920"))
FRAME_FINAL_ACCURATE_SEEK = os.getenv("FRAME_FINAL_ACCURATE_SEEK", "true").lower() == "true"
# Hybrid seek: fast-seek to (ts - PRE_SECONDS), then accurate-seek the remainder.
# Keeps frame accuracy while avoiding O(video-length) decodes on long videos.
FRAME_FINAL_PRE_SEEK_SECONDS = float(os.getenv("FRAME_FINAL_PRE_SEEK_SECONDS", "2.0"))
FFMPEG_FINAL_SCREENSHOT_TIMEOUT_SEC = float(os.getenv("FFMPEG_FINAL_SCREENSHOT_TIMEOUT_SEC", "120"))
FFMPEG_CANDIDATE_FRAME_TIMEOUT_SEC = float(os.getenv("FFMPEG_CANDIDATE_FRAME_TIMEOUT_SEC", "75"))
SCREENSHOT_BORDER_PX = int(os.getenv("SCREENSHOT_BORDER_PX", "4"))

DOC_IMAGE_WIDTH_INCHES = float(os.getenv("DOC_IMAGE_WIDTH_INCHES", "6.2"))
DOC_STEP_IMAGE_WIDTH_INCHES = float(os.getenv("DOC_STEP_IMAGE_WIDTH_INCHES", "3.58"))
DOC_STEP_TEXT_WIDTH_INCHES = float(os.getenv("DOC_STEP_TEXT_WIDTH_INCHES", "2.92"))
DOC_LEFT_MARGIN_INCHES = float(os.getenv("DOC_LEFT_MARGIN_INCHES", "0.75"))
DOC_RIGHT_MARGIN_INCHES = float(os.getenv("DOC_RIGHT_MARGIN_INCHES", "0.75"))
DOC_DARK_THEME = os.getenv("DOC_DARK_THEME", "true").lower() == "true"
DOC_SHOW_STEP_TIMESTAMPS = os.getenv("DOC_SHOW_STEP_TIMESTAMPS", "true").lower() == "true"
DOC_ACCENT_COLOR = os.getenv("DOC_ACCENT_COLOR", "E8634B").strip().lstrip("#")
DOC_HEADER_LOGO_MAX_WIDTH_INCHES = float(os.getenv("DOC_HEADER_LOGO_MAX_WIDTH_INCHES", "2.75"))
DOC_HEADER_LOGO_MAX_HEIGHT_INCHES = float(os.getenv("DOC_HEADER_LOGO_MAX_HEIGHT_INCHES", "0.72"))
DOC_CONTENT_WIDTH_INCHES = float(os.getenv("DOC_CONTENT_WIDTH_INCHES", "6.5"))

# Server deliverable: copy final .docx here for the web app / email worker (optional).
DELIVERABLE_OUTPUT_DIR = _resolve_app_path(os.getenv("DELIVERABLE_OUTPUT_DIR", "").strip())
# slug = Title_slug_videoId.docx (default) | video_id = videoId.docx only (Sunil contract)
DELIVERABLE_FILENAME_STYLE = os.getenv("DELIVERABLE_FILENAME_STYLE", "slug").strip().lower()

DEBIAS_TIMESTAMPS = os.getenv("DEBIAS_TIMESTAMPS", "true").lower() == "true"

# ── Contact info ──────────────────────────────
CONTACT_INFO: dict = {
    "website": os.getenv("CONTACT_WEBSITE", "https://turnoffroad.com"),
    "email":   os.getenv("CONTACT_EMAIL",   "help@turnoffroad.com"),
    "phone":   os.getenv("CONTACT_PHONE",   "951-505-6437"),
    "brand":   os.getenv("CONTACT_BRAND",   "Turn Offroad"),
}

logger = logging.getLogger("yt_pipeline")

# ─────────────────────────────────────────────
# AI RETRY HELPER (covers OpenAI + Gemini error patterns)
# ─────────────────────────────────────────────

def _is_ai_transient_error(err_str: str) -> bool:
    err_lower = err_str.lower()
    return (
        "503" in err_str
        or "502" in err_str
        or "500" in err_str
        or "UNAVAILABLE" in err_str
        or "429" in err_str
        or "RESOURCE_EXHAUSTED" in err_str
        or "rate_limit" in err_lower
        or "rate limit" in err_lower
        or "ratelimiterror" in err_lower
        or "apiconnectionerror" in err_lower
        or "apitimeouterror" in err_lower
        or "internalservererror" in err_lower
        or "server_error" in err_lower
        or "Server disconnected" in err_str
        or "RemoteProtocolError" in err_str
        or "Connection reset" in err_str
        or "Connection aborted" in err_str
        or "timed out" in err_lower
        or "timeout" in err_lower
        or "getaddrinfo" in err_lower
        or "name or service not known" in err_lower
        or "temporary failure in name resolution" in err_lower
        or "errno 11001" in err_lower
        or "errno 11004" in err_lower
    )


def _is_ai_model_unavailable_error(err_str: str) -> bool:
    """404 / model-not-available / not-supported responses — skip to next fallback model."""
    err_lower = err_str.lower()
    return (
        "404" in err_str
        or "not_found" in err_lower
        or "is no longer available" in err_lower
        or "not available to new users" in err_lower
        or "model not found" in err_lower
        or "model_not_found" in err_lower
        or "does not exist" in err_lower
        or "is not supported" in err_lower
        or "permission_denied" in err_lower
        or "permissiondeniederror" in err_lower
        or "notfounderror" in err_lower
        or "unsupported model" in err_lower
    )


# Backward-compatible aliases (older code paths may import these names).
_is_gemini_transient_error = _is_ai_transient_error
_is_gemini_model_unavailable_error = _is_ai_model_unavailable_error


async def _ai_with_retry(
    fn,
    *args,
    max_retries: int = 3,
    base_delay: float = 4.0,
    step_label: str = "",
    timeout_sec: float | None = None,
    models: list[str] | None = None,
    model_param: str = "model",
    **kwargs,
):
    label = f"[{step_label}] " if step_label else ""
    last_exc: Exception = RuntimeError("No attempts made")
    model_chain = [m for m in (models or []) if m]

    async def _call_once(call_kwargs: dict[str, Any]) -> Any:
        coro = asyncio.to_thread(fn, *args, **call_kwargs)
        if timeout_sec and timeout_sec > 0:
            return await asyncio.wait_for(coro, timeout=timeout_sec)
        return await coro

    if not model_chain:
        for attempt in range(max_retries + 1):
            try:
                return await _call_once(kwargs)
            except asyncio.TimeoutError as exc:
                last_exc = exc
                err_str = "AI request timed out"
                is_transient = True
            except Exception as exc:
                last_exc = exc
                err_str = str(exc)
                is_transient = _is_ai_transient_error(err_str)
            if not is_transient or attempt >= max_retries:
                raise last_exc
            delay = base_delay * (2 ** attempt)
            logger.warning(
                "%sAI transient error (attempt %d/%d): %s — retrying in %.0fs",
                label, attempt + 1, max_retries, err_str[:120], delay,
            )
            await asyncio.sleep(delay)
        raise last_exc

    model_idx = 0
    while model_idx < len(model_chain):
        active_model = model_chain[model_idx]
        call_kwargs = {**kwargs, model_param: active_model}
        is_last_model = model_idx >= len(model_chain) - 1
        per_model_cap = (
            max_retries
            if is_last_model
            else min(max_retries, max(0, OPENAI_RETRIES_BEFORE_FALLBACK))
        )
        for attempt in range(per_model_cap + 1):
            try:
                result = await _call_once(call_kwargs)
                if model_idx > 0:
                    logger.info(
                        "%sAI succeeded with fallback model %s",
                        label, active_model,
                    )
                return result
            except asyncio.TimeoutError as exc:
                last_exc = exc
                err_str = "AI request timed out"
                is_transient = True
            except Exception as exc:
                last_exc = exc
                err_str = str(exc)
                is_transient = _is_ai_transient_error(err_str)
                if _is_ai_model_unavailable_error(err_str):
                    if model_idx + 1 < len(model_chain):
                        logger.warning(
                            "%sAI model %s unavailable on this account "
                            "(%s) — switching to next fallback.",
                            label, active_model, err_str[:120],
                        )
                        break
                    raise last_exc
            if not is_transient:
                raise last_exc
            if attempt >= per_model_cap:
                break
            delay = base_delay * (2 ** attempt)
            logger.warning(
                "%sAI transient error on %s (attempt %d/%d): %s — retrying in %.0fs",
                label, active_model, attempt + 1, per_model_cap, err_str[:120], delay,
            )
            await asyncio.sleep(delay)

        if model_idx + 1 < len(model_chain):
            next_model = model_chain[model_idx + 1]
            logger.warning(
                "%sSwitching AI model %s → %s after repeated demand/capacity errors",
                label, active_model, next_model,
            )
            model_idx += 1
        else:
            raise last_exc

    raise last_exc


# Backward-compat alias: older code still calls _gemini_with_retry.
_gemini_with_retry = _ai_with_retry


# ─────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────

def setup_logging(output_dir: Path, verbose: bool) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    log_path = output_dir / "pipeline.log"

    root = logging.getLogger()
    root.setLevel(logging.DEBUG)
    root.handlers.clear()

    fmt = "%(asctime)s [%(levelname)s] %(funcName)s: %(message)s"
    datefmt = "%H:%M:%S"

    fh = logging.FileHandler(log_path, mode="w", encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(fmt, datefmt=datefmt))
    root.addHandler(fh)

    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.DEBUG if verbose else logging.INFO)
    ch.setFormatter(logging.Formatter(fmt, datefmt=datefmt))
    root.addHandler(ch)

    for noisy in (
        "anthropic", "anthropic._base_client", "google", "google.genai",
        "google_genai", "httpx", "httpcore", "httpcore.connection",
        "httpcore.http11", "scenedetect", "PIL",
    ):
        logging.getLogger(noisy).setLevel(logging.WARNING)

    logger.info("Log file: %s", log_path)


# ─────────────────────────────────────────────
# DATA MODELS
# ─────────────────────────────────────────────

def parse_timestamp_value(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        text = (
            value.strip().lower()
            .replace("seconds", "").replace("second", "")
            .replace("secs", "").replace("sec", "").replace("s", "").strip()
        )
        if ":" in text:
            nums = [float(p) for p in text.split(":")]
            if len(nums) == 2:
                return nums[0] * 60 + nums[1]
            if len(nums) == 3:
                return nums[0] * 3600 + nums[1] * 60 + nums[2]
            raise ValueError(f"Invalid timestamp: {value}")
        return float(text)
    raise ValueError(f"Invalid timestamp value: {value}")


class TranscriptSegment(BaseModel):
    id: int = Field(ge=0)
    text: str
    start: float = Field(ge=0)
    end: float = Field(ge=0)

    @field_validator("text")
    @classmethod
    def clean_text(cls, value: str) -> str:
        value = value.strip()
        if not value:
            raise ValueError("Transcript text cannot be empty.")
        return value


class SceneCut(BaseModel):
    timestamp: float = Field(ge=0)
    frame_num: Optional[int] = None


class FrameCandidate(BaseModel):
    label: str
    timestamp: float
    image_path: str
    sharpness: float = 0.0
    brightness: float = 128.0
    color_cast_ratio: float = 1.0
    motion_score: float = 0.0
    quality_score: float = 0.0
    has_prototype_banner: bool = False


class PipelinePromptOverrides(BaseModel):
    """Optional API / admin text appended to Claude prompts (not a full replacement)."""
    step_prompt_extra: str = ""
    tools_prompt_extra: str = ""
    important_note: str = ""


class PipelineResult(BaseModel):
    output_dir: str
    deliverable_path: str
    video_title: str
    video_url: str
    step_count: int
    quality_report: dict[str, Any] = Field(default_factory=dict)


class InstallationStep(BaseModel):
    step_number: int = Field(ge=1)
    title: str
    description: str
    timestamp: float = Field(ge=0)
    source_segment_ids: list[int] = Field(default_factory=list)
    visual_query: str = ""

    transcript_start: Optional[float] = None
    transcript_end: Optional[float] = None

    selected_frame_label: Optional[str] = None
    selected_frame_reason: Optional[str] = None
    selected_frame_confidence: Optional[float] = None

    screenshot_timestamp: Optional[float] = None
    screenshot_path: Optional[str] = None

    @field_validator("title", "description")
    @classmethod
    def non_empty(cls, value: str) -> str:
        value = value.strip()
        if not value:
            raise ValueError("Field cannot be empty.")
        return value

    @field_validator(
        "timestamp", "transcript_start", "transcript_end", "screenshot_timestamp",
        mode="before",
    )
    @classmethod
    def parse_ts(cls, value: Any) -> Any:
        if value is None:
            return None
        return parse_timestamp_value(value)

    @field_validator("screenshot_path", mode="before")
    @classmethod
    def normalise_path(cls, value: Any) -> Any:
        if value is None:
            return None
        return str(value).replace("\\", "/")


# ─────────────────────────────────────────────
# JSON HELPERS
# ─────────────────────────────────────────────

def strip_fences(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```json"):
        raw = raw[7:].strip()
    elif raw.startswith("```"):
        raw = raw[3:].strip()
    if raw.endswith("```"):
        raw = raw[:-3].strip()
    return raw


def extract_first_json_block(text: str) -> str:
    text = text.strip()
    starts = [i for i, ch in enumerate(text) if ch in "{["]
    if not starts:
        raise json.JSONDecodeError("No JSON start found", text, 0)
    for start in starts:
        stack: list[str] = []
        in_str = False
        escape = False
        for i in range(start, len(text)):
            ch = text[i]
            if in_str:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_str = False
                continue
            if ch == '"':
                in_str = True
            elif ch in "{[":
                stack.append(ch)
            elif ch in "}]":
                if not stack:
                    break
                opening = stack.pop()
                if opening == "{" and ch != "}":
                    break
                if opening == "[" and ch != "]":
                    break
                if not stack:
                    return text[start:i + 1]
    raise json.JSONDecodeError("Could not locate complete JSON block", text, 0)


def parse_json_response(raw: str) -> Any:
    raw = strip_fences(raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return json.loads(extract_first_json_block(raw))


def unwrap_steps(data: Any) -> list[dict[str, Any]]:
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        for key in ("steps", "items", "installation_steps", "result"):
            val = data.get(key)
            if isinstance(val, list):
                return val
    raise ValueError(f"No steps list found in response type: {type(data)}")


def save_json(data: Any, path: Path, label: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    logger.info("Saved %s → %s", label, path)


# ─────────────────────────────────────────────
# SUBPROCESS
# ─────────────────────────────────────────────

async def run_cmd(
    cmd: list[str],
    step_label: str = "",
    retries: int = 0,
    retry_delay: float = 1.5,
    timeout: Optional[float] = None,
) -> subprocess.CompletedProcess:
    label = f"[{step_label}] " if step_label else ""
    last_stdout = ""
    last_stderr = ""
    last_returncode = 1

    for attempt in range(retries + 1):
        if attempt > 0:
            delay = retry_delay * attempt
            logger.warning("%sRetrying in %.1fs...", label, delay)
            await asyncio.sleep(delay)

        logger.debug("%sRunning: %s", label, " ".join(str(x) for x in cmd))

        try:
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            try:
                stdout_b, stderr_b = await asyncio.wait_for(proc.communicate(), timeout=timeout)
            except asyncio.TimeoutError:
                proc.kill()
                await proc.wait()
                raise TimeoutError(f"Command timed out after {timeout}s")

            stdout = stdout_b.decode(errors="replace")
            stderr = stderr_b.decode(errors="replace")
            last_stdout = stdout
            last_stderr = stderr
            last_returncode = proc.returncode if proc.returncode is not None else 1

            if stdout:
                logger.debug("%sstdout: %s", label, stdout.strip()[:500])
            if stderr:
                logger.debug("%sstderr: %s", label, stderr.strip()[:500])

            if proc.returncode == 0:
                return subprocess.CompletedProcess(cmd, 0, stdout, stderr)

            logger.warning("%sExit %s", label, proc.returncode)
            if stderr.strip():
                logger.warning("%sstderr: %s", label, stderr.strip()[:800])

        except TimeoutError:
            raise
        except Exception as exc:
            last_stderr = str(exc)
            logger.warning("%sException: %s", label, exc)

    raise subprocess.CalledProcessError(
        last_returncode, cmd, output=last_stdout, stderr=last_stderr,
    )


def normalize_youtube_url(url: str) -> str:
    """
    Strip playlist/query junk so yt-dlp downloads one video, not a whole playlist.
    https://www.youtube.com/watch?v=ID&list=...&index=7 → watch?v=ID
    """
    raw = (url or "").strip()
    m = re.search(r"(?:[?&]v=|youtu\.be/)([A-Za-z0-9_-]{6,})", raw)
    if m:
        clean = f"https://www.youtube.com/watch?v={m.group(1)}"
        if clean != raw.split("#")[0].rstrip("/"):
            logger.info("Normalized YouTube URL (dropped playlist params): %s", clean)
        return clean
    return raw


def format_video_timestamp(seconds: float) -> str:
    """Human-readable timestamp for Word docs (M:SS or H:MM:SS)."""
    total = max(0, int(round(float(seconds))))
    hours, rem = divmod(total, 3600)
    minutes, secs = divmod(rem, 60)
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def step_video_timestamp_seconds(step: InstallationStep) -> float:
    """Seconds into the source video for this step's screenshot."""
    if step.screenshot_timestamp is not None:
        return float(step.screenshot_timestamp)
    return float(step.timestamp)


def youtube_watch_url_at(video_url: str, seconds: float) -> str:
    """YouTube watch URL that opens at the given second (for review / re-screenshot)."""
    base = normalize_youtube_url((video_url or "").strip())
    if not base:
        return ""
    t = max(0, int(round(float(seconds))))
    joiner = "&" if "?" in base else "?"
    return f"{base}{joiner}t={t}s"


def add_docx_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    color_hex: str = "0563C1",
    font_size_pt: int = 9,
) -> None:
    """Add a clickable external hyperlink run to a python-docx paragraph."""
    if not text or not url:
        return
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size_pt * 2)))
    r_pr.append(sz)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex.lstrip("#"))
    r_pr.append(color)
    new_run.append(r_pr)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    new_run.append(t_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def require_tool(name: str) -> None:
    if shutil.which(name) is None:
        logger.error("Required tool not found on PATH: %s", name)
        if name in {"ffmpeg", "ffprobe"}:
            logger.error("Windows: winget install Gyan.FFmpeg")
            logger.error("Mac:     brew install ffmpeg")
            logger.error("Linux:   sudo apt install ffmpeg")
        if name == "yt-dlp":
            logger.error("Install: pip install yt-dlp")
        sys.exit(1)


def validate_config() -> None:
    if not GEMINI_API_KEY:
        logger.error("GEMINI_API_KEY is required.")
        sys.exit(1)
    if not CLAUDE_API_KEY:
        logger.error("CLAUDE_API_KEY is required.")
        sys.exit(1)

    model = CLAUDE_MODEL.strip()
    if not model or "claude" not in model.lower():
        logger.error(
            "CLAUDE_MODEL=%r does not look like a valid Anthropic model string. "
            "Examples: claude-sonnet-4-20250514  claude-opus-4-20250514  claude-haiku-4-5-20251001",
            CLAUDE_MODEL,
        )
        sys.exit(1)

    if not (4 <= MAX_CANDIDATE_FRAMES <= 26):
        logger.error("MAX_CANDIDATE_FRAMES must be 4–26.")
        sys.exit(1)
    if not (4 <= WIDE_RESCUE_FRAMES <= 26):
        logger.error("WIDE_RESCUE_FRAMES must be 4–26.")
        sys.exit(1)
    if UNIFORM_SAMPLE_RESERVED_SLOTS < 1:
        logger.error("UNIFORM_SAMPLE_RESERVED_SLOTS must be >= 1.")
        sys.exit(1)
    if UNIFORM_SAMPLE_RESERVED_SLOTS >= MAX_CANDIDATE_FRAMES:
        logger.error("UNIFORM_SAMPLE_RESERVED_SLOTS must be < MAX_CANDIDATE_FRAMES.")
        sys.exit(1)
    if VISUAL_SAMPLE_INTERVAL <= 0:
        logger.error("VISUAL_SAMPLE_INTERVAL must be > 0.")
        sys.exit(1)

    require_tool("yt-dlp")
    require_tool("ffmpeg")
    require_tool("ffprobe")

    if not OPENAI_API_KEY:
        logger.error(
            "OPENAI_API_KEY is not set. Add it to .env — required for transcription and vision."
        )
        sys.exit(1)

    if OPENAI_ENABLE_MODEL_FALLBACK:
        logger.debug(
            "Config OK. Claude=%s | OpenAI-transcribe=%s (fallbacks: %s) | "
            "OpenAI-vision=%s (fallbacks: %s)",
            CLAUDE_MODEL,
            OPENAI_TRANSCRIBE_MODEL,
            " → ".join(OPENAI_TRANSCRIBE_MODELS[1:]) or "none",
            OPENAI_VISION_MODEL,
            " → ".join(OPENAI_VISION_MODELS[1:]) or "none",
        )
    else:
        logger.debug(
            "Config OK. Claude=%s | OpenAI-transcribe=%s | OpenAI-vision=%s",
            CLAUDE_MODEL, OPENAI_TRANSCRIBE_MODEL, OPENAI_VISION_MODEL,
        )

    yt_status = yt_dlp_config_status()
    if yt_status["cookies_file"]:
        if yt_status["cookies_loaded"]:
            logger.info("yt-dlp cookies: %s", yt_status["cookies_file"])
        else:
            logger.warning(
                "YT_DLP_COOKIES_FILE is set but missing/empty: %s",
                yt_status["cookies_file"],
            )
    elif yt_status["proxy_configured"]:
        logger.info("yt-dlp proxy configured (no cookies file).")
    else:
        logger.warning(
            "No YT_DLP_COOKIES_FILE or YT_DLP_PROXY set — YouTube may block datacenter IPs."
        )


# ─────────────────────────────────────────────
# VIDEO DOWNLOAD / NORMALIZE
# ─────────────────────────────────────────────

def _yt_dlp_cookies_path() -> Optional[Path]:
    if not YT_DLP_COOKIES_FILE:
        return None
    p = Path(YT_DLP_COOKIES_FILE)
    if not p.is_absolute():
        p = Path(_resolve_app_path(YT_DLP_COOKIES_FILE))
    return p


def _is_youtube_bot_error(exc_or_text: str | Exception) -> bool:
    text = str(exc_or_text).lower()
    return any(
        marker in text
        for marker in (
            "sign in to confirm",
            "confirm you're not a bot",
            "not a bot",
            "cookies-from-browser",
            "http error 403",
            "unable to extract uploader id",
            "this content isn't available",
            "bot",
        )
    )


def yt_dlp_config_status() -> dict[str, Any]:
    """Summarize yt-dlp auth settings for /health (no secrets)."""
    cookie_path = _yt_dlp_cookies_path()
    cookie_ok = bool(cookie_path and cookie_path.is_file() and cookie_path.stat().st_size > 0)
    return {
        "bin": YT_DLP_BIN,
        "cookies_file": str(cookie_path) if cookie_path else None,
        "cookies_loaded": cookie_ok,
        "proxy_configured": bool(YT_DLP_PROXY),
        "remote_components": YT_DLP_REMOTE_COMPONENTS or None,
        "player_client": YT_DLP_PLAYER_CLIENT or None,
    }


def _yt_dlp_auth_args(*, player_client: Optional[str] = None) -> list[str]:
    """Cookies, proxy, EJS solver, and player-client overrides for YouTube bot checks."""
    args: list[str] = []

    cookie_path = _yt_dlp_cookies_path()
    if cookie_path:
        if cookie_path.is_file():
            args.extend(["--cookies", str(cookie_path)])
        else:
            logger.warning("YT_DLP_COOKIES_FILE not found: %s", cookie_path)

    if YT_DLP_PROXY:
        args.extend(["--proxy", YT_DLP_PROXY])

    if YT_DLP_REMOTE_COMPONENTS:
        args.extend(["--remote-components", YT_DLP_REMOTE_COMPONENTS])

    clients = (player_client or YT_DLP_PLAYER_CLIENT).strip()
    if clients:
        args.extend(["--extractor-args", f"youtube:player_client={clients}"])

    if YT_DLP_EXTRACTOR_ARGS:
        args.extend(["--extractor-args", YT_DLP_EXTRACTOR_ARGS])

    if YT_DLP_EXTRA_ARGS:
        args.extend(shlex.split(YT_DLP_EXTRA_ARGS, posix=os.name != "nt"))

    return args


def _yt_dlp_base_cmd(*, player_client: Optional[str] = None) -> list[str]:
    """Shared yt-dlp flags for downloads (resume-friendly, more retries on flaky links)."""
    return [
        YT_DLP_BIN,
        "--no-check-certificates",
        "--no-playlist",
        "--retries", str(YT_DLP_RETRIES),
        "--fragment-retries", str(YT_DLP_FRAGMENT_RETRIES),
        "--socket-timeout", str(YT_DLP_SOCKET_TIMEOUT),
        "--continue",
        "--write-info-json",
        "--merge-output-format", "mp4",
        *_yt_dlp_auth_args(player_client=player_client),
    ]


def _yt_dlp_format_chain() -> list[str]:
    """720p-first when YT_DLP_PREFER_720_FIRST — smaller files survive slow networks better."""
    fmt_720 = (
        "bestvideo[height<=720][vcodec^=avc1][ext=mp4]+bestaudio[ext=m4a]"
        "/bestvideo[height<=720][ext=mp4]+bestaudio/best[height<=720][ext=mp4]/best"
    )
    fmt_1080 = (
        "bestvideo[height<=1080][vcodec^=avc1][ext=mp4]+bestaudio[ext=m4a]"
        "/bestvideo[height<=720][vcodec^=avc1][ext=mp4]+bestaudio[ext=m4a]"
        "/best[height<=1080][ext=mp4]/best"
    )
    fmt_small_merge = "best[height<=720][ext=mp4]/b[height<=720]/best"
    if YT_DLP_PREFER_720_FIRST:
        return [fmt_720, fmt_small_merge, fmt_1080, "best"]
    return [fmt_1080, fmt_720, fmt_small_merge, "best"]


async def ingest_local_video(source: Path, output_dir: Path) -> Path:
    """Copy a local file to output_dir/video.mp4 and skip yt-dlp."""
    source = source.expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(f"--video-file not found: {source}")
    if source.stat().st_size == 0:
        raise ValueError(f"--video-file is empty: {source}")

    output_dir.mkdir(parents=True, exist_ok=True)
    final_path = output_dir / "video.mp4"
    logger.info(
        "[1/7] Using local video (skip download): %s (%.1f MB)",
        source.name, source.stat().st_size / 1_048_576,
    )
    if final_path.resolve() != source.resolve():
        shutil.copy2(source, final_path)
    return final_path


async def download_video(url: str, output_dir: Path) -> Path:
    url = normalize_youtube_url(url)
    logger.info("[1/7] Downloading video...")
    output_dir.mkdir(parents=True, exist_ok=True)

    final_path = output_dir / "video.mp4"
    output_template = output_dir / "video.%(ext)s"

    if final_path.exists() and final_path.stat().st_size > 0:
        logger.info(
            "Already downloaded: %s (%.1f MB)",
            final_path, final_path.stat().st_size / 1_048_576,
        )
        return final_path

    format_chain = _yt_dlp_format_chain()
    last_error: Optional[Exception] = None

    player_client_attempts: list[Optional[str]] = [None]
    seen_clients: set[str] = set()
    if YT_DLP_PLAYER_CLIENT:
        for client in [c.strip() for c in YT_DLP_PLAYER_CLIENT.split(",") if c.strip()]:
            if client not in seen_clients:
                seen_clients.add(client)
                player_client_attempts.append(client)
    for client in ("mweb", "android", "web"):
        if client not in seen_clients:
            player_client_attempts.append(client)

    for player_client in player_client_attempts:
        if player_client:
            logger.info("yt-dlp player client attempt: %s", player_client)

        for fmt in format_chain:
            try:
                logger.info("Trying format: %s", fmt[:80])
                await run_cmd(
                    _yt_dlp_base_cmd(player_client=player_client)
                    + ["-f", fmt, "-o", str(output_template), url],
                    step_label="yt-dlp", retries=1,
                )

                candidates = [
                    p for p in output_dir.glob("video.*")
                    if p.suffix.lower() in {".mp4", ".mkv", ".webm", ".mov"}
                    and not p.name.endswith(".part")
                    and p.stat().st_size > 0
                ]

                if not candidates:
                    raise FileNotFoundError("yt-dlp completed but no video file was found.")

                downloaded = max(candidates, key=lambda p: p.stat().st_size)

                if downloaded.resolve() == final_path.resolve():
                    logger.info(
                        "Downloaded → %s (%.1f MB)",
                        final_path.name, final_path.stat().st_size / 1_048_576,
                    )
                    return final_path

                logger.info("Remuxing %s → %s", downloaded.name, final_path.name)
                try:
                    await run_cmd(
                        [
                            "ffmpeg", "-hide_banner", "-loglevel", "error",
                            "-y", "-i", str(downloaded), "-c", "copy", str(final_path),
                        ],
                        step_label="ffmpeg-remux",
                    )
                except subprocess.CalledProcessError:
                    logger.warning("Remux failed. Re-encoding to H.264.")
                    await run_cmd(
                        [
                            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
                            "-i", str(downloaded),
                            "-vf", f"scale={H264_PROCESSING_WIDTH}:-2",
                            "-c:v", "libx264", "-preset", "veryfast", "-crf", "22",
                            "-g", "15", "-keyint_min", "15", "-sc_threshold", "0",
                            "-c:a", "aac", "-b:a", "128k", str(final_path),
                        ],
                        step_label="ffmpeg-transcode", timeout=900,
                    )

                if not final_path.exists() or final_path.stat().st_size == 0:
                    raise FileNotFoundError(f"Normalized file missing: {final_path}")

                logger.info(
                    "Downloaded → %s (%.1f MB)",
                    final_path.name, final_path.stat().st_size / 1_048_576,
                )
                return final_path

            except Exception as exc:
                last_error = exc
                logger.warning("Format failed: %s", exc)
                if not _is_youtube_bot_error(exc):
                    break

    bot_hint = ""
    if last_error and _is_youtube_bot_error(last_error):
        bot_hint = (
            "\nYouTube blocked this server IP (bot check). Fix:\n"
            "  1) Export fresh browser cookies to YT_DLP_COOKIES_FILE (Netscape format).\n"
            "  2) Set YT_DLP_REMOTE_COMPONENTS=ejs:github and ensure deno is installed.\n"
            "  3) Optional fallback: YT_DLP_PROXY with a residential proxy URL.\n"
            "See python_script/DEPLOY.md → YouTube auth on the droplet.\n"
        )

    raise RuntimeError(
        f"All yt-dlp attempts failed. Last error: {last_error}{bot_hint}\n"
        "Other tips:\n"
        "  1) Download outside the pipeline, then rerun with --video-file.\n"
        "  2) Do NOT use --fresh if a partial download exists (yt-dlp can resume).\n"
        "  3) pip install -U yt-dlp; set YT_DLP_PREFER_720_FIRST=true (default)."
    )


async def get_video_codec(video_path: Path) -> str:
    result = await run_cmd(
        [
            "ffprobe", "-v", "quiet", "-print_format", "json",
            "-select_streams", "v:0", "-show_streams", str(video_path),
        ],
        step_label="ffprobe-codec",
    )
    data = json.loads(result.stdout)
    streams = data.get("streams", [])
    if not streams:
        return ""
    return str(streams[0].get("codec_name", "")).lower()


async def normalize_video_for_processing(video_path: Path) -> Path:
    codec = await get_video_codec(video_path)
    logger.info("Video codec: %s", codec or "unknown")

    if codec == "h264":
        logger.info("Already H.264. No transcode needed.")
        return video_path

    logger.warning(
        "Codec is %s. Transcoding to H.264 for reliable frame extraction.",
        codec or "unknown",
    )

    output_dir = video_path.parent
    backup_path = output_dir / "video_original_backup.mp4"
    h264_path = output_dir / "video_h264.mp4"

    if h264_path.exists():
        h264_path.unlink()

    await run_cmd(
        [
            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
            "-i", str(video_path),
            "-vf", f"scale={H264_PROCESSING_WIDTH}:-2",
            "-c:v", "libx264", "-preset", "veryfast", "-crf", "22",
            "-g", "15", "-keyint_min", "15", "-sc_threshold", "0",
            "-c:a", "copy", str(h264_path),
        ],
        step_label="h264-normalize", timeout=900,
    )

    if not h264_path.exists() or h264_path.stat().st_size == 0:
        raise RuntimeError(f"H.264 transcode failed: {h264_path}")

    if backup_path.exists():
        backup_path.unlink()
    video_path.rename(backup_path)
    h264_path.rename(video_path)

    logger.info("H.264 transcode complete. Backup: %s", backup_path.name)
    return video_path


def read_video_title_from_info_json(output_dir: Path) -> Optional[str]:
    """Read title from yt-dlp --write-info-json sidecar (video.info.json)."""
    info_path = output_dir / "video.info.json"
    if not info_path.exists():
        return None
    try:
        with open(info_path, encoding="utf-8") as f:
            data = json.load(f)
        title = (data.get("title") or "").strip()
        return title or None
    except Exception as exc:
        logger.debug("Could not read title from %s: %s", info_path, exc)
        return None


async def get_video_title(url: str, output_dir: Optional[Path] = None) -> str:
    url = normalize_youtube_url(url)
    if output_dir is not None:
        t = read_video_title_from_info_json(output_dir)
        if t:
            return t
    try:
        result = await run_cmd(
            _yt_dlp_base_cmd() + ["--get-title", url],
            step_label="yt-title", retries=1, timeout=YT_DLP_TITLE_TIMEOUT_SEC,
        )
        title = result.stdout.strip()
        if title:
            return title
    except Exception as exc:
        logger.debug("Could not fetch title: %s", exc)
    return "Installation Guide"


async def get_duration(video_path: Path) -> float:
    result = await run_cmd(
        [
            "ffprobe", "-v", "quiet", "-print_format", "json",
            "-show_format", str(video_path),
        ],
        step_label="ffprobe-duration",
    )
    return float(json.loads(result.stdout)["format"]["duration"])


def ffprobe_duration_sync(path: Path) -> float:
    result = subprocess.run(
        [
            "ffprobe", "-v", "quiet", "-print_format", "json",
            "-show_format", str(path),
        ],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr)
    return float(json.loads(result.stdout)["format"]["duration"])


# ─────────────────────────────────────────────
# SCENE DETECTION
# ─────────────────────────────────────────────

async def detect_scenes(video_path: Path) -> list[SceneCut]:
    logger.info("[2/7] Detecting scene cuts with ffmpeg scene filter...")

    try:
        result = await asyncio.wait_for(
            run_cmd(
                [
                    "ffmpeg", "-hide_banner",
                    "-i", str(video_path),
                    "-vf", f"select='gt(scene,{SCENE_THRESHOLD})',showinfo",
                    "-vsync", "vfr",
                    "-f", "null", "-",
                ],
                step_label="scene-detect",
            ),
            timeout=SCENE_DETECTION_TIMEOUT_SECONDS,
        )

        output = result.stderr if result.stderr else result.stdout

        cuts: list[SceneCut] = []
        for m in re.finditer(r"pts_time:([\d.]+)", output):
            ts = round(float(m.group(1)), 3)
            if ts > 0.5:
                cuts.append(SceneCut(timestamp=ts))

        merged: list[SceneCut] = []
        for c in sorted(cuts, key=lambda x: x.timestamp):
            if not merged or c.timestamp - merged[-1].timestamp > 0.5:
                merged.append(c)

        logger.info("Detected %d scene cuts.", len(merged))
        return merged

    except asyncio.TimeoutError:
        logger.warning("Scene detection timed out. Continuing without scene cuts.")
        return []
    except Exception as exc:
        logger.warning("Scene detection failed: %s. Continuing without.", exc)
        return []


# ─────────────────────────────────────────────
# TRANSCRIPT TIMESTAMP NORMALISATION
# ─────────────────────────────────────────────

def _convert_m_ss_to_seconds(value: float) -> float:
    minutes = int(value)
    frac = value - minutes
    seconds = round(frac * 100, 3)
    return round(minutes * 60 + seconds, 3)


def _looks_like_m_ss(segments: list[TranscriptSegment], video_duration: float) -> bool:
    if len(segments) < 5 or video_duration <= 0:
        return False
    starts = [s.start for s in segments]
    max_start = max(starts)
    if max_start <= 0:
        return False
    if max_start >= video_duration * 0.25:
        return False
    valid = sum(
        1 for v in starts
        if 0.0 <= round((v - int(v)) * 100, 1) <= 59.99
    )
    if valid / len(starts) < 0.70:
        return False
    converted_max = _convert_m_ss_to_seconds(max_start)
    if converted_max > video_duration + 30:
        return False
    return True


def normalise_transcript_timestamps(
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[TranscriptSegment]:
    if not DEBIAS_TIMESTAMPS or len(segments) < 5:
        return segments

    from collections import Counter

    if _looks_like_m_ss(segments, video_duration):
        converted = [
            s.model_copy(update={
                "start": max(0.0, _convert_m_ss_to_seconds(s.start)),
                "end":   max(0.0, _convert_m_ss_to_seconds(s.end)),
            })
            for s in segments
        ]
        starts_c = [s.start for s in converted]
        if starts_c[-1] <= video_duration + 30 and starts_c == sorted(starts_c):
            logger.warning(
                "Transcript timestamps were in M.SS format (max raw=%.3f). "
                "Converted to seconds (max converted=%.1fs). Video duration=%.1fs.",
                max(s.start for s in segments), starts_c[-1], video_duration,
            )
            return [s.model_copy(update={"id": i}) for i, s in enumerate(converted)]
        else:
            logger.warning("M.SS conversion produced implausible timestamps — skipping.")

    fracs = [round(s.start % 1, 2) for s in segments]
    counts = Counter(fracs)
    most_common_frac, frac_count = counts.most_common(1)[0]
    ratio = frac_count / len(fracs)

    if ratio >= 0.60 and most_common_frac > 0.01:
        logger.info(
            "Timestamp bias detected: fractional part %.2f in %.0f%% of segments. "
            "Subtracting %.2fs from all timestamps.",
            most_common_frac, ratio * 100, most_common_frac,
        )
        debiased = [
            s.model_copy(update={
                "start": max(0.0, round(s.start - most_common_frac, 3)),
                "end":   max(0.0, round(s.end   - most_common_frac, 3)),
            })
            for s in segments
        ]
        return debiased

    logger.debug("Transcript timestamps look clean. No normalisation needed.")
    return segments


_MARKETING_OUTRO_PHRASES: tuple[str, ...] = (
    "congratulations",
    "transform your",
    "until your next turn",
    "don't hesitate to contact",
    "be safe out there",
    "more than an accessory",
    "iconic heritage",
    "perfect fusion of classic design",
)


def _is_marketing_outro_text(text: str) -> bool:
    t = text.lower()
    return any(p in t for p in _MARKETING_OUTRO_PHRASES)


# Minimum length of a run of segments with identical start times to flag transcript as broken.
COLLAPSED_TRANSCRIPT_RUN_THRESHOLD = int(
    os.getenv("COLLAPSED_TRANSCRIPT_RUN_THRESHOLD", "4")
)


def _transcript_appears_collapsed(segments: list[TranscriptSegment]) -> tuple[bool, int]:
    """Return (is_collapsed, longest_run). Gemini occasionally returns many late segments
    all stamped with the same start time — those timestamps are unreliable."""
    if len(segments) < COLLAPSED_TRANSCRIPT_RUN_THRESHOLD:
        return False, 0
    longest = 1
    run = 1
    for i in range(1, len(segments)):
        if abs(segments[i].start - segments[i - 1].start) < 0.6:
            run += 1
            longest = max(longest, run)
        else:
            run = 1
    return longest >= COLLAPSED_TRANSCRIPT_RUN_THRESHOLD, longest


def redistribute_collapsed_segments(
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[TranscriptSegment]:
    """Spread runs of >=3 segments sharing the same start time across the
    available timeline so Claude gets plausible (if approximate) timestamps."""
    if len(segments) < 3 or video_duration <= 0:
        return segments

    out = list(segments)
    n = len(out)
    safe_end = max(2.0, video_duration - SAFE_END_PADDING_SECONDS)

    i = 0
    while i < n:
        start_i = out[i].start
        j = i + 1
        while j < n and abs(out[j].start - start_i) < 0.6:
            j += 1
        run_len = j - i
        if run_len >= 3:
            prev_end = out[i - 1].end if i > 0 else 0.0
            lower = max(prev_end + 0.1, start_i - 0.5)
            if j < n:
                upper = max(lower + 1.0, out[j].start - 0.1)
            else:
                upper = safe_end
            if upper - lower < 1.0 and i > 0:
                lookback = max(0.0, prev_end - run_len * 3.0)
                lower = max(lookback, 0.0)
                upper = max(lower + 1.0, prev_end + 1.0, safe_end)
            if upper > lower + 0.5:
                step = (upper - lower) / run_len
                logger.warning(
                    "Redistributing %d collapsed transcript segments (ids %d-%d) "
                    "from %.1fs across [%.1f, %.1f].",
                    run_len, out[i].id, out[j - 1].id, start_i, lower, upper,
                )
                for k in range(run_len):
                    new_start = lower + k * step
                    new_end = lower + (k + 1) * step - 0.05
                    out[i + k] = out[i + k].model_copy(update={
                        "start": round(max(0.0, new_start), 3),
                        "end": round(max(new_start + 0.3, new_end), 3),
                    })
        i = j

    return out


def fit_transcript_to_video_duration(
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[TranscriptSegment]:
    """
    Gemini transcription often drifts past real video length on the closing narration.
    Clamp and chain-fix segments so Claude step timestamps stay inside the MP4.
    """
    if not segments or video_duration <= 0:
        return segments

    hard_cap = max(1.0, video_duration - SAFE_END_PADDING_SECONDS)

    trimmed: list[TranscriptSegment] = []
    for s in segments:
        if _is_marketing_outro_text(s.text):
            continue
        trimmed.append(s)

    fixed: list[TranscriptSegment] = []
    last_end = 0.0
    drift_fixes = 0
    for s in trimmed:
        start, end = float(s.start), float(s.end)
        dur = max(0.5, end - start)
        if start > hard_cap:
            start = min(last_end + 0.35, max(0.0, hard_cap - 2.0))
            end = min(start + dur * 0.45, video_duration)
            drift_fixes += 1
        start = max(0.0, min(start, hard_cap))
        end = max(start + 0.35, min(end, video_duration))
        last_end = end
        fixed.append(
            s.model_copy(update={"start": round(start, 3), "end": round(end, 3)}),
        )

    max_end_after_fix = max((s.end for s in fixed), default=0.0)
    if max_end_after_fix > hard_cap + 2.0:
        scale = hard_cap / max_end_after_fix
        fixed = [
            s.model_copy(update={
                "start": round(s.start * scale, 3),
                "end": round(min(s.end * scale, hard_cap), 3),
            })
            for s in fixed
        ]
        logger.warning(
            "Post-drift-fix rescale by %.3f (max_end=%.1fs → %.1fs).",
            scale, max_end_after_fix, hard_cap,
        )

    max_end = max((s.end for s in fixed), default=0.0)
    if max_end > video_duration + 5.0:
        scale = hard_cap / max_end
        fixed = [
            s.model_copy(update={
                "start": round(s.start * scale, 3),
                "end": round(min(s.end * scale, hard_cap), 3),
            })
            for s in fixed
        ]
        logger.warning(
            "Scaled transcript timestamps by %.3f to fit %.1fs video.",
            scale, video_duration,
        )

    if drift_fixes or len(fixed) != len(segments):
        logger.warning(
            "fit_transcript_to_video_duration: %d drift fixes, %d → %d segments "
            "(video %.1fs, was max %.1fs).",
            drift_fixes,
            len(segments),
            len(fixed),
            video_duration,
            max((s.end for s in segments), default=0.0),
        )

    return [s.model_copy(update={"id": i}) for i, s in enumerate(fixed)]


def _apply_transcript_duration_fixes(
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[TranscriptSegment]:
    if video_duration <= 0:
        return segments
    segments = normalise_transcript_timestamps(segments, video_duration)
    pre_collapsed, pre_run = _transcript_appears_collapsed(segments)
    if pre_collapsed:
        logger.warning(
            "Transcript collapsed before drift-fix (run=%d). Redistributing timestamps.",
            pre_run,
        )
        segments = redistribute_collapsed_segments(segments, video_duration)
    fitted = fit_transcript_to_video_duration(segments, video_duration)
    post_collapsed, post_run = _transcript_appears_collapsed(fitted)
    if post_collapsed:
        logger.warning(
            "Transcript collapsed after drift-fix (run=%d) — drift loop ceilinged "
            "multiple late segments at hard_cap. Redistributing timestamps.",
            post_run,
        )
        fitted = redistribute_collapsed_segments(fitted, video_duration)
    return fitted


# ─────────────────────────────────────────────
# TRANSCRIPT TEXT CLEANING
# ─────────────────────────────────────────────

_TRANSCRIPT_FIXES: list[tuple[str, str]] = [
    ("hard wire up top",    "hardware up top"),
    ("hard wire",           "hardware"),
    ("hardwire up top",     "hardware up top"),
    ("hardwire",            "hardware"),
    ("hard ware",           "hardware"),
    ("target top",          "tailgate top"),
    ("the target",          "the tailgate"),
    ("clamp down points",   "clamp-down points"),
    ("clamp down",          "clamp-down"),
    ("t 30",                "T30"),
    ("t30",                 "T30"),
    ("t 40",                "T40"),
    ("t40",                 "T40"),
    ("t 47",                "T47"),
    ("t47",                 "T47"),
    ("t 50",                "T50"),
    ("t50",                 "T50"),
    ("  ",                  " "),
]


def clean_transcript_text(text: str) -> str:
    for wrong, right in _TRANSCRIPT_FIXES:
        text = re.sub(re.escape(wrong), right, text, flags=re.IGNORECASE)
    return text.strip()


def clean_segments_for_prompt(segments: list[TranscriptSegment]) -> list[TranscriptSegment]:
    return [
        s.model_copy(update={"text": clean_transcript_text(s.text)})
        for s in segments
    ]


# ─────────────────────────────────────────────
# GEMINI TRANSCRIPTION
# ─────────────────────────────────────────────

def load_cached_transcript(path: Path) -> Optional[list[TranscriptSegment]]:
    if not path.exists():
        return None
    try:
        with open(path, encoding="utf-8") as f:
            raw = json.load(f)
        segs = []
        for i, item in enumerate(raw):
            if "id" not in item:
                item["id"] = i
            segs.append(TranscriptSegment.model_validate(item))
        logger.info("Transcript cache hit: %d segments.", len(segs))
        return segs
    except Exception as exc:
        logger.warning("Transcript cache invalid: %s. Retranscribing.", exc)
        return None


def extract_audio_mp3_sync(video_path: Path, audio_path: Path) -> None:
    result = subprocess.run(
        [
            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
            "-i", str(video_path),
            "-vn", "-ac", "1", "-ar", "16000", "-b:a", "64k",
            str(audio_path),
        ],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Audio extraction failed:\n{result.stderr}")


def split_audio_sync(audio_path: Path, chunks_dir: Path) -> list[Path]:
    chunks_dir.mkdir(parents=True, exist_ok=True)
    for old in chunks_dir.glob("chunk_*.mp3"):
        old.unlink()

    chunk_tmpl = chunks_dir / "chunk_%03d.mp3"
    result = subprocess.run(
        [
            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
            "-i", str(audio_path), "-f", "segment",
            "-segment_time", str(AUDIO_CHUNK_SECONDS), "-c", "copy",
            str(chunk_tmpl),
        ],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"Audio splitting failed:\n{result.stderr}")

    chunks = sorted(chunks_dir.glob("chunk_*.mp3"))
    if not chunks:
        raise RuntimeError("No audio chunks created.")
    return chunks


def text_to_synthetic_segments(
    text: str,
    offset: float,
    start_id: int,
    window: float = 10.0,
) -> list[TranscriptSegment]:
    words = text.split()
    if not words:
        return []
    w_per_win = max(1, int(window * 2.5))
    segs: list[TranscriptSegment] = []
    for i, chunk_start in enumerate(range(0, len(words), w_per_win)):
        chunk = words[chunk_start:chunk_start + w_per_win]
        start = offset + i * window
        segs.append(
            TranscriptSegment(
                id=start_id + len(segs), text=" ".join(chunk),
                start=round(start, 3), end=round(start + window, 3),
            )
        )
    logger.warning("Synthetic transcript segments created: %d.", len(segs))
    return segs


def transcribe_one_openai_sync(
    audio_path: Path,
    offset_seconds: float,
    start_id: int,
    model: str | None = None,
) -> list[TranscriptSegment]:
    """
    Transcribe a single audio file with OpenAI.

    Primary path: ``whisper-1`` with ``response_format="verbose_json"`` to get
    real per-segment timestamps. Fallback path (``gpt-4o-transcribe`` /
    ``gpt-4o-mini-transcribe``) returns text only — we synthesize timestamps
    so downstream step extraction still has window anchors to work with.
    """
    client = _openai_client()
    active_model = (model or OPENAI_TRANSCRIBE_MODEL).strip()
    is_whisper = "whisper" in active_model.lower()

    prompt = (
        "Preserve tool names, part names, sizes, and installation action words exactly. "
        "Do not summarize. 'tailgate top' is a vehicle part — transcribe it exactly, "
        "not as 'target top'. 'hardware' refers to physical fasteners/clips — never "
        "'hard wire' or 'hardwire'."
    )

    with open(audio_path, "rb") as audio_file:
        if is_whisper:
            response = client.audio.transcriptions.create(
                model=active_model,
                file=audio_file,
                response_format="verbose_json",
                timestamp_granularities=["segment"],
                temperature=0,
                prompt=prompt,
            )
        else:
            # gpt-4o*-transcribe: text only, no segments
            response = client.audio.transcriptions.create(
                model=active_model,
                file=audio_file,
                response_format="text",
                temperature=0,
                prompt=prompt,
            )

    if is_whisper:
        raw_segments = getattr(response, "segments", None) or []
        if not raw_segments:
            text_blob = (getattr(response, "text", "") or "").strip()
            logger.warning(
                "OpenAI %s returned no segments — using synthetic timestamps.",
                active_model,
            )
            return text_to_synthetic_segments(text_blob, offset_seconds, start_id)
        segs: list[TranscriptSegment] = []
        for item in raw_segments:
            if hasattr(item, "model_dump"):
                obj = item.model_dump()
            elif isinstance(item, dict):
                obj = item
            else:
                obj = {
                    "text": getattr(item, "text", ""),
                    "start": getattr(item, "start", 0),
                    "end":   getattr(item, "end", 0),
                }
            text = str(obj.get("text", "")).strip()
            if not text:
                continue
            try:
                start = parse_timestamp_value(obj.get("start", 0)) + offset_seconds
                end = parse_timestamp_value(obj.get("end", start + 5)) + offset_seconds
            except Exception:
                continue
            if end <= start:
                end = start + 5.0
            segs.append(
                TranscriptSegment(
                    id=start_id + len(segs), text=text,
                    start=round(start, 3), end=round(end, 3),
                )
            )
        if segs:
            return segs
        text_blob = (getattr(response, "text", "") or "").strip()
        logger.warning(
            "OpenAI %s segments produced no usable rows — using synthetic timestamps.",
            active_model,
        )
        return text_to_synthetic_segments(text_blob, offset_seconds, start_id)

    text_blob = response if isinstance(response, str) else str(getattr(response, "text", response))
    text_blob = (text_blob or "").strip()
    logger.info(
        "OpenAI %s (text-only) — synthesizing segment timestamps from %d chars of text.",
        active_model, len(text_blob),
    )
    return text_to_synthetic_segments(text_blob, offset_seconds, start_id)


def transcribe_one_gemini_sync(
    audio_path: Path,
    offset_seconds: float,
    start_id: int,
    model: str | None = None,
) -> list[TranscriptSegment]:
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        raise RuntimeError("google-genai not installed. Run: pip install google-genai")

    client = genai.Client(api_key=GEMINI_API_KEY)

    prompt = (
        "Transcribe this audio into timestamped segments.\n\n"
        "Return ONLY valid JSON. No markdown. No explanation.\n\n"
        "Use this exact shape:\n"
        '{"segments":[{"id":0,"start":0.0,"end":4.2,"text":"spoken text"}]}\n\n'
        "Rules:\n"
        "- start and end are seconds from the beginning of THIS audio file.\n"
        "- Each segment should be 3 to 10 seconds long.\n"
        "- Preserve tool names, part names, sizes, and installation action words exactly.\n"
        "- Do not summarize. Transcribe the spoken words verbatim.\n"
        "- Use clean decimal timestamps, not fractions with recurring digits.\n"
        "- 'tailgate top' is a vehicle part — transcribe it exactly, not as 'target top'.\n"
        "- 'hardware' refers to physical fasteners/clips — not 'hard wire' or 'hardwire'.\n"
    )

    audio_bytes = audio_path.read_bytes()

    active_model = model or GEMINI_TRANSCRIBE_MODEL
    response = client.models.generate_content(
        model=active_model,
        contents=[
            prompt,
            types.Part.from_bytes(data=audio_bytes, mime_type="audio/mpeg"),
        ],
        config=types.GenerateContentConfig(
            temperature=0,
            response_mime_type="application/json",
        ),
    )

    raw = (response.text or "").strip()
    logger.debug("Gemini transcript raw %d chars: %s...", len(raw), raw[:200])

    try:
        data = parse_json_response(raw)
    except Exception as exc:
        logger.warning("Gemini transcript JSON parse failed: %s. Using synthetic.", exc)
        return text_to_synthetic_segments(raw, offset_seconds, start_id)

    if isinstance(data, dict) and isinstance(data.get("segments"), list):
        raw_segs = data["segments"]
    elif isinstance(data, list):
        raw_segs = data
    else:
        logger.warning("Gemini transcript unexpected shape. Using synthetic.")
        return text_to_synthetic_segments(raw, offset_seconds, start_id)

    segs: list[TranscriptSegment] = []
    for item in raw_segs:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        try:
            start = parse_timestamp_value(item.get("start", 0)) + offset_seconds
            end = parse_timestamp_value(item.get("end", start + 5)) + offset_seconds
        except Exception:
            continue
        if end <= start:
            end = start + 5.0
        segs.append(
            TranscriptSegment(
                id=start_id + len(segs), text=text,
                start=round(start, 3), end=round(end, 3),
            )
        )

    if segs:
        return segs

    logger.warning("Gemini returned no usable transcript segments. Using synthetic.")
    return text_to_synthetic_segments(raw, offset_seconds, start_id)


async def transcribe_audio(
    video_path: Path,
    output_dir: Path,
    video_duration: float = 0.0,
) -> list[TranscriptSegment]:
    logger.info("[3/7] Transcribing audio with OpenAI %s...", OPENAI_TRANSCRIBE_MODEL)

    transcript_path = output_dir / "transcript.json"
    cached = load_cached_transcript(transcript_path)
    if cached:
        if video_duration > 0:
            fixed = _apply_transcript_duration_fixes(cached, video_duration)
            if fixed is not cached:
                logger.info("Re-fitted cached transcript to video duration.")
                with open(transcript_path, "w", encoding="utf-8") as f:
                    json.dump([s.model_dump() for s in fixed], f, indent=2, ensure_ascii=False)
                return fixed
        return cached

    audio_path = output_dir / "audio.mp3"
    logger.info("Extracting audio → %s", audio_path)
    await asyncio.to_thread(extract_audio_mp3_sync, video_path, audio_path)

    if audio_path.stat().st_size <= MAX_OPENAI_AUDIO_BYTES:
        logger.info("Sending full audio to OpenAI %s.", OPENAI_TRANSCRIBE_MODEL)
        segments = await _ai_with_retry(
            transcribe_one_openai_sync,
            audio_path, 0.0, 0,
            max_retries=4,
            base_delay=8.0,
            step_label="openai-transcribe",
            models=OPENAI_TRANSCRIBE_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
        )
        collapsed, run_len = _transcript_appears_collapsed(segments)
        if collapsed and OPENAI_ENABLE_MODEL_FALLBACK and len(OPENAI_TRANSCRIBE_MODELS) > 1:
            logger.warning(
                "Primary OpenAI transcript has %d consecutive segments with identical start time "
                "— retrying with fallback transcription models.",
                run_len,
            )
            for fallback in OPENAI_TRANSCRIBE_MODELS[1:]:
                try:
                    candidate = await asyncio.to_thread(
                        transcribe_one_openai_sync, audio_path, 0.0, 0, fallback,
                    )
                except Exception as exc:
                    logger.warning(
                        "Fallback transcription model %s failed: %s", fallback, exc,
                    )
                    continue
                cand_collapsed, cand_run = _transcript_appears_collapsed(candidate)
                if not cand_collapsed:
                    logger.info(
                        "Fallback transcription model %s produced clean timestamps "
                        "(%d segments).",
                        fallback, len(candidate),
                    )
                    segments = candidate
                    break
                logger.warning(
                    "Fallback model %s also returned collapsed transcript "
                    "(longest run %d).",
                    fallback, cand_run,
                )
    else:
        logger.info(
            "Audio too large (%.1f MB). Splitting into chunks.",
            audio_path.stat().st_size / 1_048_576,
        )
        chunks_dir = output_dir / "audio_chunks"
        chunks = await asyncio.to_thread(split_audio_sync, audio_path, chunks_dir)
        segments = []
        offset = 0.0
        for chunk in chunks:
            logger.info("Chunk %s at offset %.1fs", chunk.name, offset)
            chunk_segs = await _ai_with_retry(
                transcribe_one_openai_sync,
                chunk, offset, len(segments),
                max_retries=4,
                base_delay=8.0,
                step_label=f"openai-transcribe-{chunk.name}",
                models=OPENAI_TRANSCRIBE_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
            )
            segments.extend(chunk_segs)
            offset += await asyncio.to_thread(ffprobe_duration_sync, chunk)

    if not segments:
        raise RuntimeError("OpenAI transcription returned no segments.")

    segments = [s.model_copy(update={"id": i}) for i, s in enumerate(segments)]

    if video_duration > 0:
        segments = _apply_transcript_duration_fixes(segments, video_duration)
        segments = [s.model_copy(update={"id": i}) for i, s in enumerate(segments)]

    with open(transcript_path, "w", encoding="utf-8") as f:
        json.dump([s.model_dump() for s in segments], f, indent=2, ensure_ascii=False)

    logger.info("%d transcript segments saved → %s", len(segments), transcript_path)
    return segments


# ─────────────────────────────────────────────
# CLAUDE STEP EXTRACTION
# ─────────────────────────────────────────────

def build_step_prompt(
    segments: list[TranscriptSegment],
    extra_instructions: str = "",
    video_duration: float = 0.0,
) -> str:
    clean_segs = clean_segments_for_prompt(segments)
    transcript_text = "\n".join(
        f"[segment_id={s.id}] [{s.start:.2f}s - {s.end:.2f}s] {s.text}"
        for s in clean_segs
    )
    extra_block = ""
    if extra_instructions and extra_instructions.strip():
        extra_block = (
            "\n━━━ CLIENT / ADMIN EXTRA INSTRUCTIONS ━━━\n"
            f"{extra_instructions.strip()}\n"
            "Apply only when consistent with the rules above.\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        )
    duration_block = ""
    if video_duration > 0:
        content_end = content_max_timestamp(video_duration, segments)
        duration_block = (
            f"\nVIDEO FACTS:\n"
            f"- Total video length: {video_duration:.1f} seconds.\n"
            f"- Latest valid step timestamp: {content_end:.1f} seconds.\n"
            "- Every step timestamp MUST fall inside the video and inside the cited segment times.\n"
            "- Do NOT create steps from marketing, congratulations, or contact/closing narration.\n"
            "- If the narrator describes an action that is not shown on camera before the video ends, "
            "omit that step.\n\n"
        )
    return (
        "You are extracting procedural installation steps from a video transcript.\n\n"
        f"{duration_block}"
        "Think internally about whether each sentence is a real installation action. "
        "Do not output reasoning.\n\n"
        "INCLUDE concrete physical actions:\n"
        "- removing, attaching, installing, positioning, aligning parts\n"
        "- tightening bolts, screws, nuts, clamps\n"
        "- connecting wires, cables, hoses, plugs\n"
        "- applying adhesive, sealant, tape, lubricant\n"
        "- drilling, cutting, measuring, marking\n"
        "- checking or verifying installed result\n\n"
        "EXCLUDE:\n"
        "- introductions, greetings, marketing, sponsorships\n"
        "- tool lists without an active installation action\n"
        "- warnings with no associated action\n"
        "- repeated narration of the same exact action\n"
        "- outro, contact, subscribe, help text, legal warnings, or caution disclaimers\n"
        "  (do NOT emit steps for 'warnings & cautions' title cards — those are not install steps)\n\n"
        "━━━ CRITICAL STEP GRANULARITY RULES ━━━\n"
        "ONE STEP = ONE CAMERA SHOT OF ONE ATOMIC ACTION. If a frame from the\n"
        "step's timestamp can show the entire action being performed, it's atomic.\n"
        "If you'd need multiple frames at different timestamps to illustrate the\n"
        "step, you have bundled actions — SPLIT them into separate steps.\n\n"
        "ANTI-BUNDLING RULES (these are violations — never produce them):\n"
        "- 'Finger-tighten X AND then torque-tighten X with a 19mm socket' → 2 steps\n"
        "  (one frame can't show both finger-tight AND wrench-tight at the same time).\n"
        "- 'Insert bolts AND adjust the bumper AND tighten the bolts' → 3 steps.\n"
        "- 'Tighten the hitch bolts AND tighten the frame mount bolts' → 2 steps\n"
        "  (different physical components on the vehicle, different camera angle).\n"
        "- 'Thread backing nuts AND tighten hitch bolts AND tighten frame mounts' → 3 steps.\n"
        "- 'Position the bumper AND thread the bolts' → 2 steps (lift vs. thread).\n"
        "- ANY description containing 'then' followed by a new physical action → SPLIT.\n"
        "- ANY description containing 'Adjust… until… then tighten…' → SPLIT.\n\n"
        "WHEN TO MERGE (only these cases):\n"
        f"- Two timestamps closer than {MIN_STEP_GAP_SECONDS:.0f}s AND describing the\n"
        "  SAME atomic action on the SAME component (e.g. 'first bed rail bolt'\n"
        f"  followed by 'remaining bed rail bolts' within {MIN_STEP_GAP_SECONDS:.0f}s).\n"
        f"- A 'reinstall' / 'put back' action within {MIN_STEP_GAP_SECONDS:.0f}s of its\n"
        "  matching 'remove' step (append as 'Reverse to reinstall.' sentence).\n"
        "- Otherwise: do NOT merge. Emit separate steps even if they're consecutive\n"
        "  in the narration, as long as the camera shows them as distinct shots.\n\n"
        "COUNT EXPECTATION: A 5-7 minute install video typically has 7-12 distinct\n"
        "atomic steps. If you produce fewer than 6, you have almost certainly\n"
        "bundled multiple actions. Re-read each step's description: if it contains\n"
        "the word 'then', 'adjust… until', or describes work on more than one\n"
        "physical component, SPLIT it before emitting.\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
        "For each step:\n"
        "  step_number        : integer starting at 1\n"
        "  title              : short imperative phrase (max 7 words, NO slash-separated alternatives)\n"
        "  description        : 1 to 3 clear sentences. Include specific part names, "
        "tool sizes, and directional cues. Use compound words correctly: 'hardware' "
        "not 'hard ware', 'hardtop' or 'hard top' as appropriate. Never write "
        "'hardwire' when you mean 'hardware'. Use 'tailgate top' not 'target top'.\n"
        "  timestamp          : float seconds — must be the START of the action\n"
        "  source_segment_ids : list of transcript segment IDs used\n"
        "  visual_query       : one sentence describing what the PERFECT screenshot "
        "should show — be specific about the physical object, human hand(s), and "
        "action motion visible in frame. If the step requires two people, say so.\n\n"
        "Rules:\n"
        "- Never invent timestamps.\n"
        "- source_segment_ids MUST directly contain the spoken instruction for that step.\n"
        "- TIMESTAMP ACCURACY IS CRITICAL: the timestamp must be the exact second the\n"
        "  physical action begins on camera. If the narrator describes an action while\n"
        "  showing something unrelated, use the transcript segment where the action is\n"
        "  BEING PERFORMED, not where it is introduced or explained.\n"
        f"- If two consecutive steps describe actions in the same physical area of the\n"
        f"  vehicle, double-check their timestamps are at least {MIN_STEP_GAP_SECONDS:.0f}s\n"
        "  apart and point to visually distinct moments.\n"
        "- When a narrator says 'now we will...' or 'next we...' before doing the action,\n"
        "  the timestamp should be AFTER that intro phrase, at the segment where doing\n"
        "  begins.\n"
        "- source_segment_ids MUST NOT cite distant unrelated segments — only IDs whose\n"
        "  spoken text describes THIS step (never bundle an entire video section).\n"
        "- timestamp MUST be between the earliest start and latest end of source_segment_ids.\n"
        "- Do NOT merge unrelated actions that occur far apart in time.\n"
        "- Timestamp should be at the first visual moment of the action, not mid-way.\n"
        "- Emit one step per atomic physical action — see granularity rules above.\n"
        "  Err on the side of MORE granular steps when in doubt; under-segmenting\n"
        "  produces low-confidence screenshots downstream.\n"
        "- Preserve ALL part names, sizes, torx bit sizes, and orientation details.\n\n"
        f"{extra_block}"
        f"TRANSCRIPT:\n{transcript_text}"
    )


def normalize_step(raw: dict[str, Any], idx: int) -> dict[str, Any]:
    return {
        "step_number": raw.get("step_number", raw.get("index", raw.get("step", idx))),
        "title": raw.get("title", raw.get("name", raw.get("heading", f"Step {idx}"))),
        "description": raw.get("description", raw.get("details", raw.get("instruction", ""))),
        "timestamp": raw.get(
            "timestamp", raw.get("time", raw.get("start", raw.get("start_time", 0))),
        ),
        "source_segment_ids": raw.get(
            "source_segment_ids", raw.get("segment_ids", raw.get("segments", [])),
        ),
        "visual_query": raw.get("visual_query", raw.get("visual_description", "")),
    }


def validate_steps(raw_steps: list[dict[str, Any]]) -> list[InstallationStep]:
    steps: list[InstallationStep] = []
    for i, raw in enumerate(raw_steps, start=1):
        normalized = normalize_step(raw, i)
        try:
            steps.append(InstallationStep.model_validate(normalized))
        except ValidationError as exc:
            logger.error("Step %d validation failed.\nRaw: %s\nError: %s", i, raw, exc)
            raise
    return [s.model_copy(update={"step_number": i}) for i, s in enumerate(steps, start=1)]


def call_claude_steps_sync(prompt: str) -> list[InstallationStep]:
    try:
        import anthropic
    except ImportError:
        raise RuntimeError("anthropic not installed. Run: pip install anthropic")

    logger.info("Calling Claude with schema tool: %s", CLAUDE_MODEL)
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    tool_schema = {
        "name": "emit_installation_steps",
        "description": "Emit extracted installation steps as structured JSON.",
        "input_schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "steps": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "step_number": {"type": "integer", "minimum": 1},
                            "title": {"type": "string", "minLength": 1},
                            "description": {"type": "string", "minLength": 1},
                            "timestamp": {"type": "number", "minimum": 0},
                            "source_segment_ids": {
                                "type": "array",
                                "items": {"type": "integer", "minimum": 0},
                            },
                            "visual_query": {"type": "string"},
                        },
                        "required": [
                            "step_number", "title", "description",
                            "timestamp", "source_segment_ids", "visual_query",
                        ],
                    },
                }
            },
            "required": ["steps"],
        },
    }

    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=CLAUDE_MAX_TOKENS,
        temperature=CLAUDE_TEMPERATURE,
        system=(
            "You extract installation procedure steps. "
            "Use the provided tool exactly once. "
            "Do not output free-form text."
        ),
        tools=[tool_schema],
        tool_choice={"type": "tool", "name": "emit_installation_steps"},
        messages=[{"role": "user", "content": prompt}],
    )

    for block in message.content:
        if getattr(block, "type", None) == "tool_use":
            tool_input = getattr(block, "input", None)
            if isinstance(tool_input, dict):
                raw_steps = unwrap_steps(tool_input)
                return validate_steps(raw_steps)

    parts = [getattr(b, "text", "") for b in message.content if getattr(b, "text", "")]
    raw = "\n".join(parts).strip()
    if raw:
        logger.warning("Claude returned text instead of tool_use. Parsing fallback JSON.")
        return validate_steps(unwrap_steps(parse_json_response(raw)))

    raise RuntimeError("Claude returned no tool_use content and no text fallback.")


def normalized_text_for_similarity(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def text_similarity(a: str, b: str) -> float:
    a = normalized_text_for_similarity(a)
    b = normalized_text_for_similarity(b)
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def steps_are_probable_duplicates(a: InstallationStep, b: InstallationStep) -> bool:
    title_sim = text_similarity(a.title, b.title)
    desc_sim = text_similarity(a.description, b.description)
    time_close = abs(a.timestamp - b.timestamp) <= 8.0
    ids_a = set(a.source_segment_ids or [])
    ids_b = set(b.source_segment_ids or [])
    segment_overlap = bool(ids_a and ids_b and ids_a.intersection(ids_b))
    highly_similar_text = title_sim >= 0.94 and desc_sim >= 0.84
    same_time_area = time_close or segment_overlap
    return highly_similar_text and same_time_area


def deduplicate_steps(steps: list[InstallationStep]) -> list[InstallationStep]:
    kept: list[InstallationStep] = []
    for step in steps:
        duplicate_of: Optional[InstallationStep] = None
        for previous in kept:
            if steps_are_probable_duplicates(step, previous):
                duplicate_of = previous
                break
        if duplicate_of:
            logger.info(
                "Removed probable duplicate: '%s' near %.1fs. Duplicate of '%s' near %.1fs.",
                step.title, step.timestamp, duplicate_of.title, duplicate_of.timestamp,
            )
            continue
        kept.append(step)
    return [s.model_copy(update={"step_number": i}) for i, s in enumerate(kept, start=1)]


def _clean_slash_title(title: str) -> str:
    if " / " in title:
        return title.split(" / ")[0].strip()
    return title


_MERGE_ACTION_VERBS = frozenset({
    "remove", "install", "slide", "pull", "push", "tighten", "torque",
    "mount", "attach", "unbolt", "loosen", "position", "align", "drill",
    "cut", "clip", "unclip", "lower", "lift", "reinstall",
})


def _step_action_verbs(title: str) -> set[str]:
    t = title.lower()
    return {v for v in _MERGE_ACTION_VERBS if v in t}


def _should_merge_close_steps(prev: InstallationStep, step: InstallationStep, gap: float) -> bool:
    """Merge only when steps are truly the same beat — not two different actions 9s apart."""
    if gap >= MIN_STEP_GAP_SECONDS:
        return False
    prev_verbs = _step_action_verbs(prev.title)
    step_verbs = _step_action_verbs(step.title)
    if prev_verbs and step_verbs and prev_verbs != step_verbs:
        logger.info(
            "Keeping separate steps '%s' + '%s' (%.1fs gap, distinct actions).",
            prev.title, step.title, gap,
        )
        return False
    return True


def enforce_min_step_gap(steps: list[InstallationStep]) -> list[InstallationStep]:
    if len(steps) <= 1:
        return steps

    merged: list[InstallationStep] = [steps[0]]
    for step in steps[1:]:
        prev = merged[-1]
        gap = step.timestamp - prev.timestamp
        if _should_merge_close_steps(prev, step, gap):
            logger.info(
                "Merging step '%s' (%.1fs) into '%s' (%.1fs) — gap %.1fs < %.0fs.",
                step.title, step.timestamp, prev.title, prev.timestamp, gap, MIN_STEP_GAP_SECONDS,
            )
            combined_title = _clean_slash_title(prev.title)
            combined_desc = prev.description.rstrip(". ") + ". " + step.description
            combined_ids = sorted(set(prev.source_segment_ids or []) | set(step.source_segment_ids or []))
            combined_vq = prev.visual_query or step.visual_query
            merged_timestamp = round((prev.timestamp + step.timestamp) / 2.0, 3)
            prev_start = prev.transcript_start if prev.transcript_start is not None else prev.timestamp
            step_start = step.transcript_start if step.transcript_start is not None else step.timestamp
            prev_end = prev.transcript_end if prev.transcript_end is not None else prev.timestamp
            step_end = step.transcript_end if step.transcript_end is not None else step.timestamp
            merged[-1] = prev.model_copy(update={
                "title": combined_title,
                "description": combined_desc,
                "source_segment_ids": combined_ids,
                "visual_query": combined_vq,
                "timestamp": merged_timestamp,
                "transcript_start": round(min(prev_start, step_start), 3),
                "transcript_end": round(max(prev_end, step_end), 3),
            })
        else:
            merged.append(step)

    if len(merged) < len(steps):
        logger.info(
            "Step gap enforcement: reduced %d → %d steps.", len(steps), len(merged)
        )

    cleaned = [
        s.model_copy(update={"title": _clean_slash_title(s.title)})
        for s in merged
    ]

    return [s.model_copy(update={"step_number": i}) for i, s in enumerate(cleaned, start=1)]


def segment_window_for_ids(
    source_ids: list[int],
    by_id: dict[int, TranscriptSegment],
) -> Optional[tuple[float, float]]:
    chosen = [by_id[sid] for sid in source_ids if sid in by_id]
    if not chosen:
        return None
    return min(s.start for s in chosen), max(s.end for s in chosen)


def filter_source_ids_near_timestamp(
    source_ids: list[int],
    by_id: dict[int, TranscriptSegment],
    anchor_ts: float,
    max_span: float = MAX_SOURCE_SEGMENT_SPAN_SECONDS,
) -> list[int]:
    """Drop transcript segment IDs far from the step anchor (Claude often over-cites)."""
    chosen = [by_id[sid] for sid in source_ids if sid in by_id]
    if not chosen:
        return source_ids

    nearby = [
        s for s in chosen
        if s.end >= anchor_ts - 8.0 and s.start <= anchor_ts + max_span
    ]
    if nearby:
        return sorted({s.id for s in nearby})

    chosen.sort(
        key=lambda s: min(abs(s.start - anchor_ts), abs(s.end - anchor_ts)),
    )
    return [chosen[0].id]


def cap_transcript_window(
    start: float,
    end: float,
    anchor_ts: float,
    max_span: float = MAX_TRANSCRIPT_WINDOW_SECONDS,
) -> tuple[float, float]:
    """Shrink an oversized transcript window to stay centered on the action."""
    if end - start <= max_span:
        return start, end

    half = max_span / 2.0
    capped_start = max(0.0, anchor_ts - half)
    capped_end = capped_start + max_span
    logger.debug(
        "Capping transcript window %.1f-%.1fs → %.1f-%.1fs (anchor %.1fs, max %.0fs).",
        start, end, capped_start, capped_end, anchor_ts, max_span,
    )
    return capped_start, capped_end


def nearest_segment_to_timestamp(
    timestamp: float,
    segments: list[TranscriptSegment],
) -> TranscriptSegment:
    return min(
        segments,
        key=lambda s: min(
            abs(s.start - timestamp),
            abs(s.end - timestamp),
            abs((s.start + s.end) / 2 - timestamp),
        ),
    )


def timestamp_inside_window(ts: float, start: float, end: float, tolerance: float = 2.0) -> bool:
    return start - tolerance <= ts <= end + tolerance


_ANTICIPATORY_PHRASES = (
    "now we", "next we", "we're going to", "we will", "going to",
    "you're going to", "you will", "what we want to", "let's go ahead",
    "first thing", "make sure you", "the next step",
)


def _segment_is_anticipatory(text: str) -> bool:
    t = text.lower().strip()
    return any(t.startswith(p) for p in _ANTICIPATORY_PHRASES)


def _advance_timestamp_past_intro(
    timestamp: float,
    source_ids: list[int],
    by_id: dict[int, TranscriptSegment],
) -> float:
    """
    If the first cited segment is anticipatory narration, nudge the timestamp
    forward to the next segment so vision search lands on the action itself.
    """
    if not source_ids:
        return timestamp
    first = by_id.get(source_ids[0])
    if first is None or not _segment_is_anticipatory(first.text):
        return timestamp
    next_seg = by_id.get(source_ids[1]) if len(source_ids) > 1 else by_id.get(source_ids[0] + 1)
    if next_seg and next_seg.start > timestamp:
        nudged = round(next_seg.start + 0.5, 3)
        logger.debug(
            "Anticipatory segment detected ('%s...') — nudging timestamp %.2fs → %.2fs.",
            first.text[:40], timestamp, nudged,
        )
        return nudged
    return timestamp


def enrich_steps_with_transcript_windows(
    steps: list[InstallationStep],
    segments: list[TranscriptSegment],
) -> list[InstallationStep]:
    if not segments:
        raise RuntimeError("Cannot enrich steps: transcript segments are empty.")

    by_id = {s.id: s for s in segments}
    enriched: list[InstallationStep] = []

    for step in steps:
        filtered_ids = filter_source_ids_near_timestamp(
            step.source_segment_ids, by_id, step.timestamp,
        )
        if filtered_ids != step.source_segment_ids:
            logger.info(
                "Step %02d: trimmed source_segment_ids %s → %s (anchor %.1fs).",
                step.step_number, step.source_segment_ids, filtered_ids, step.timestamp,
            )

        window = segment_window_for_ids(filtered_ids, by_id)

        if window:
            start, end = window
            if not timestamp_inside_window(step.timestamp, start, end, tolerance=3.0):
                logger.warning(
                    "Step %02d timestamp %.2fs is outside cited window %.2f-%.2fs. "
                    "Replacing timestamp with %.2fs.",
                    step.step_number, step.timestamp, start, end, start,
                )
                fixed_timestamp = start
            else:
                fixed_timestamp = step.timestamp
        else:
            nearest = nearest_segment_to_timestamp(step.timestamp, segments)
            start, end = nearest.start, nearest.end
            fixed_timestamp = nearest.start
            logger.warning(
                "Step %02d has no valid source_segment_ids. Using nearest segment %d at %.2f-%.2fs.",
                step.step_number, nearest.id, start, end,
            )

        fixed_timestamp = _advance_timestamp_past_intro(
            fixed_timestamp, filtered_ids, by_id,
        )

        start, end = cap_transcript_window(start, end, fixed_timestamp)

        enriched.append(
            step.model_copy(update={
                "timestamp": round(fixed_timestamp, 3),
                "source_segment_ids": filtered_ids,
                "transcript_start": round(start, 3),
                "transcript_end": round(end, 3),
                "visual_query": step.visual_query or f"{step.title}. {step.description}",
            })
        )

    capped: list[InstallationStep] = []
    for idx, step in enumerate(enriched):
        end = step.transcript_end or step.timestamp + 8.0
        if idx + 1 < len(enriched):
            next_start = enriched[idx + 1].timestamp
            end = min(end, max(step.timestamp + 1.0, next_start - 0.25))
        if step.transcript_start is not None and end <= step.transcript_start:
            end = step.transcript_start + 1.0
        capped.append(step.model_copy(update={"transcript_end": round(end, 3)}))

    return capped


_STEP_TITLE_STOPWORDS = frozenset({
    "step", "with", "from", "into", "your", "this", "that", "the", "and",
    "for", "using", "use", "place", "onto", "next", "side", "panel",
    "before", "after", "down", "over", "back", "front", "rear", "left",
    "right", "both", "their", "them", "have", "make", "sure", "then",
})

_STEP_ACTION_VERBS = frozenset({
    "install", "installing", "installed",
    "mount", "mounting", "mounted",
    "attach", "attaching", "attached",
    "remove", "removing", "removed",
    "pull", "pulling", "pulled",
    "push", "pushing", "pushed",
    "slide", "sliding", "slid",
    "tighten", "tightening", "tightened",
    "loosen", "loosening", "loosened",
    "position", "positioning", "positioned",
    "align", "aligning", "aligned",
    "screw", "screwing", "screwed", "unscrew",
    "fasten", "fastening", "fastened",
    "thread", "threading", "threaded",
    "lift", "lifting", "lifted",
    "snug", "torque", "torqued",
    "insert", "inserting", "inserted",
    "secure", "securing", "secured",
    "connect", "connecting", "connected",
    "disconnect", "disconnecting",
    "place", "placing", "placed",
    "drop", "dropping", "dropped",
    "grab", "grabbing", "grabbed",
    "flip", "flipping", "flipped",
})

# Earliest-mention relocation is destructive when the transcript timestamps
# are wrong; gate it behind an env flag and add strong correctness checks.
RELOCATE_LATE_STEPS_ENABLED = (
    os.getenv("RELOCATE_LATE_STEPS_ENABLED", "true").lower() == "true"
)
# Steps cannot be relocated earlier than this fraction of the video duration
# (intro/parts list / tool overview lives in the first ~25%).
RELOCATE_EARLIEST_FRACTION = float(os.getenv("RELOCATE_EARLIEST_FRACTION", "0.30"))


def _step_title_keywords(title: str) -> set[str]:
    return {
        w.lower().strip(".,!?:;'\"")
        for w in re.findall(r"\b[a-zA-Z][a-zA-Z0-9]{3,}\b", title)
        if w.lower() not in _STEP_TITLE_STOPWORDS
        and w.lower() not in _STEP_ACTION_VERBS
    }


def _step_action_keywords(title: str, description: str = "") -> set[str]:
    text = f"{title} {description}".lower()
    return {w for w in re.findall(r"\b[a-zA-Z]+\b", text) if w in _STEP_ACTION_VERBS}


def relocate_late_steps_to_earliest_mention(
    steps: list[InstallationStep],
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[InstallationStep]:
    """For each step whose timestamp lands near the outro, find a transcript segment
    that contains BOTH a step-title noun AND a step-title action verb, and pull the
    step's timestamp back to that mention. Requires monotonic ordering and excludes
    intro/parts-list segments. Critical when Gemini's transcript bunches late segments."""
    if not RELOCATE_LATE_STEPS_ENABLED:
        return steps
    if not steps or not segments or video_duration <= 0:
        return steps

    content_end = content_max_timestamp(video_duration, segments)
    late_threshold = max(0.0, content_end - 25.0)
    earliest_allowed = max(20.0, video_duration * RELOCATE_EARLIEST_FRACTION)

    relocated: list[InstallationStep] = []
    prev_anchor = 0.0
    min_gap = max(4.0, MIN_STEP_GAP_SECONDS * 0.6)

    for step in steps:
        if step.timestamp < late_threshold:
            prev_anchor = max(prev_anchor, step.timestamp)
            relocated.append(step)
            continue

        nouns = _step_title_keywords(step.title)
        actions = _step_action_keywords(step.title, step.description)
        if len(nouns) < 1 or not actions:
            relocated.append(step)
            continue

        floor = max(earliest_allowed, prev_anchor + min_gap)
        ceiling = step.timestamp - 1.0
        if floor >= ceiling:
            relocated.append(step)
            continue

        best: TranscriptSegment | None = None
        best_score = 0
        for seg in segments:
            if seg.start < floor:
                continue
            if seg.start >= ceiling:
                break
            words_in_seg = {
                w.lower() for w in re.findall(r"\b[a-zA-Z]+\b", seg.text)
            }
            noun_hits = len(nouns & words_in_seg)
            action_hits = len(actions & words_in_seg)
            if noun_hits >= 1 and action_hits >= 1:
                score = noun_hits * 2 + action_hits
                if score > best_score:
                    best = seg
                    best_score = score

        if best is None:
            relocated.append(step)
            continue

        new_ts = max(prev_anchor + min_gap, round(best.start + 0.4, 3))
        new_ts = min(new_ts, ceiling - 0.5)
        new_start = round(max(0.0, new_ts - 4.0), 3)
        new_end = round(min(content_end, max(new_ts + 12.0, best.end + 2.0)), 3)
        logger.warning(
            "Step %02d '%s': relocated %.1fs → %.1fs via noun+action match "
            "(seg %d: '%s...').",
            step.step_number, step.title, step.timestamp, new_ts,
            best.id, best.text[:60],
        )
        prev_anchor = new_ts
        relocated.append(step.model_copy(update={
            "timestamp": new_ts,
            "transcript_start": new_start,
            "transcript_end": new_end,
            "source_segment_ids": [best.id],
        }))

    return relocated


def clamp_steps_to_video_content(
    steps: list[InstallationStep],
    segments: list[TranscriptSegment],
    video_duration: float,
) -> list[InstallationStep]:
    """Drop or clamp steps whose timestamps cannot exist in the MP4."""
    if video_duration <= 0:
        return steps

    content_end = content_max_timestamp(video_duration, segments)
    by_id = {s.id: s for s in segments}
    kept: list[InstallationStep] = []

    for step in steps:
        cited = [by_id[sid] for sid in step.source_segment_ids if sid in by_id]
        if cited and all(_is_marketing_outro_text(s.text) for s in cited):
            logger.info(
                "Dropping step %02d '%s' — cites only outro/marketing transcript.",
                step.step_number, step.title,
            )
            continue

        ts = step.timestamp
        if ts > content_end + STEP_CLAMP_PAST_CONTENT_SECONDS:
            if cited:
                ts = min(cited[0].start, content_end - 2.0)
            else:
                ts = content_end - 5.0
            logger.warning(
                "Step %02d '%s': clamped timestamp %.1fs → %.1fs (content end %.1fs).",
                step.step_number, step.title, step.timestamp, ts, content_end,
            )
            step = step.model_copy(update={
                "timestamp": round(max(0.0, ts), 3),
                "transcript_start": round(max(0.0, ts - 8.0), 3),
                "transcript_end": round(min(content_end, ts + 12.0), 3),
            })

        kept.append(step)

    if len(kept) < len(steps):
        kept = [s.model_copy(update={"step_number": i}) for i, s in enumerate(kept, start=1)]
    return kept


async def extract_steps_with_ai(
    segments: list[TranscriptSegment],
    step_prompt_extra: str = "",
    video_duration: float = 0.0,
) -> list[InstallationStep]:
    logger.info("[4/7] Extracting steps with Claude...")
    prompt = build_step_prompt(
        segments,
        extra_instructions=step_prompt_extra,
        video_duration=video_duration,
    )
    steps = await asyncio.to_thread(call_claude_steps_sync, prompt)
    before = len(steps)
    steps = deduplicate_steps(steps)
    if len(steps) < before:
        logger.info("Removed %d duplicate steps.", before - len(steps))

    steps = enrich_steps_with_transcript_windows(steps, segments)
    steps = relocate_late_steps_to_earliest_mention(steps, segments, video_duration)
    steps = clamp_steps_to_video_content(steps, segments, video_duration)
    steps = enforce_min_step_gap(steps)
    steps = filter_non_install_steps(steps)
    steps.sort(key=lambda s: s.timestamp)

    logger.info("Extracted %d steps.", len(steps))
    return steps


# ─────────────────────────────────────────────
# TOOLS / TIME / DIFFICULTY EXTRACTION
# ─────────────────────────────────────────────

# v12: Keywords used to classify items as hardware when auto-splitting a flat list.
_HARDWARE_KEYWORDS = frozenset({
    "washer", "bolt", "bolts", "nut", "nuts", "clip", "clips", "screw", "screws",
    "rivet", "pin", "pins", "ring", "o-ring", "hex", "fastener",
    "seal", "gasket", "adhesive", "tape", "spring", "bracket", "stud", "studs",
    "spacer", "bushing", "grommet", "retainer", "cap", "plug", "insert",
    "consumable", "cable tie", "zip tie",
})


def _auto_split_tools(flat_tools: list[str]) -> tuple[list[str], list[str]]:
    """Split a flat tools list into (tools, hardware) based on keyword matching."""
    tools_out, hardware_out = [], []
    for item in flat_tools:
        lower = item.lower()
        if any(kw in lower for kw in _HARDWARE_KEYWORDS):
            hardware_out.append(item)
        else:
            tools_out.append(item)
    return tools_out, hardware_out


def _call_claude_tools_info_sync(
    transcript_text: str,
    extra_instructions: str = "",
) -> dict:
    try:
        import anthropic
    except ImportError:
        return {}

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    tool_schema = {
        "name": "emit_install_metadata",
        "description": "Emit install summary metadata for the cover page.",
        "input_schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "tools": {
                    "type": "array",
                    "description": (
                        "Hand tools and power tools only — things you hold and operate. "
                        "Examples: 'T30 Torx bit', 'T40 Torx bit', 'T50 Torx bit', "
                        "'13 mm socket', '10 mm wrench', 'power drill'. "
                        "Do NOT include fasteners, clips, or consumables here."
                    ),
                    "items": {"type": "string"},
                },
                "hardware": {
                    "type": "array",
                    "description": (
                        "Physical fasteners, clips, seals, and consumables — things you install "
                        "or attach to the vehicle. Examples: 'spring washers', 'body clips', "
                        "'bed rail bolts', 'washers', 'jam nuts', 'T47 bolts (spares)'. "
                        "Do NOT include hand tools or power tools here."
                    ),
                    "items": {"type": "string"},
                },
                "estimated_time_minutes": {
                    "type": "integer",
                    "description": "Rough total install time in minutes.",
                },
                "difficulty": {
                    "type": "string",
                    "enum": ["Easy", "Moderate", "Advanced"],
                    "description": "Overall difficulty level.",
                },
                "people_required": {
                    "type": "integer",
                    "description": "Minimum number of people needed (1 or 2).",
                },
            },
            "required": ["tools", "hardware", "estimated_time_minutes", "difficulty", "people_required"],
        },
    }

    prompt = (
        "Read this installation video transcript and extract:\n"
        "1. Tools — hand tools / power tools / bit sizes only (e.g. T30 Torx bit, 13 mm socket).\n"
        "2. Hardware — fasteners, clips, washers, consumables installed on the vehicle.\n"
        "3. Estimated total installation time in minutes.\n"
        "4. Overall difficulty (Easy / Moderate / Advanced).\n"
        "5. Minimum number of people required (1 or 2).\n\n"
        "CRITICAL: Keep tools and hardware in SEPARATE lists. "
        "A T47 bolt is hardware (it's installed on the vehicle). "
        "A T47 Torx bit is a tool (you hold it in your hand). "
        "Spring washers and body clips are hardware. A 10 mm wrench is a tool.\n"
        "If a video mentions no fasteners at all, hardware may be an empty list.\n\n"
        "Return only the tool call, no free text.\n\n"
    )
    if extra_instructions and extra_instructions.strip():
        prompt += (
            "Additional guidance from the client:\n"
            f"{extra_instructions.strip()}\n\n"
        )
    prompt += f"TRANSCRIPT (first 8000 chars):\n{transcript_text[:8000]}"

    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=1000,
            system=(
                "Extract install metadata with tools and hardware in separate lists. "
                "Use the tool exactly once."
            ),
            tools=[tool_schema],
            tool_choice={"type": "tool", "name": "emit_install_metadata"},
            messages=[{"role": "user", "content": prompt}],
        )
        for block in msg.content:
            if getattr(block, "type", None) == "tool_use":
                inp = getattr(block, "input", {})
                if isinstance(inp, dict):
                    return inp
    except Exception as exc:
        logger.warning("tools-info extraction failed: %s", exc)
    return {}


async def extract_tools_info_with_ai(
    segments: list[TranscriptSegment],
    tools_prompt_extra: str = "",
) -> dict:
    logger.info("Extracting tools / time / difficulty with Claude...")
    transcript_text = "\n".join(f"[{s.start:.1f}s] {s.text}" for s in segments)
    result = await asyncio.to_thread(
        _call_claude_tools_info_sync, transcript_text, tools_prompt_extra,
    )
    if result:
        logger.info(
            "Tools: %d items | Hardware: %d items | Time: %s min | Difficulty: %s | People: %s",
            len(result.get("tools", [])),
            len(result.get("hardware", [])),
            result.get("estimated_time_minutes", "?"),
            result.get("difficulty", "?"),
            result.get("people_required", "?"),
        )
    return result


# ─────────────────────────────────────────────
# STEP HELPERS
# ─────────────────────────────────────────────

def is_reinstall_step(step: InstallationStep) -> bool:
    combined = (step.title + " " + step.description).lower()
    return any(kw in combined for kw in REINSTALL_KEYWORDS)


def is_multi_person_step(step: InstallationStep) -> bool:
    combined = (step.title + " " + step.description + " " + step.visual_query).lower()
    return any(kw in combined for kw in MULTI_PERSON_KEYWORDS)


def is_non_install_step(step: InstallationStep) -> bool:
    """Title-card / legal / outro steps that are not real install actions."""
    title = step.title.lower()
    return any(kw in title for kw in OUTRO_STEP_TITLE_KEYWORDS)


def filter_non_install_steps(steps: list[InstallationStep]) -> list[InstallationStep]:
    kept: list[InstallationStep] = []
    for step in steps:
        if is_non_install_step(step):
            logger.info(
                "Dropping non-install step %02d '%s' (outro/warnings title).",
                step.step_number, step.title,
            )
            continue
        kept.append(step)
    if len(kept) < len(steps):
        return [s.model_copy(update={"step_number": i}) for i, s in enumerate(kept, start=1)]
    return steps


def prepare_steps_for_vision(
    steps: list[InstallationStep],
    video_duration: float,
) -> list[InstallationStep]:
    """
    Clamp timestamps to searchable video content and pull outro-timestamped steps
    back into the last stretch of install footage.
    """
    content_end = content_max_timestamp(video_duration)
    ordered = sorted(steps, key=lambda s: s.timestamp)
    prepared: list[InstallationStep] = []

    past_count = sum(
        1 for s in ordered
        if s.timestamp > content_end
        or (s.transcript_start is not None and s.transcript_start > content_end)
    )
    past_index = 0
    min_gap = max(4.0, MIN_STEP_GAP_SECONDS * 0.6)
    spread_span = max(15.0, min_gap * max(1, past_count - 1) + 6.0)
    spread_start = max(0.0, content_end - spread_span)

    for step in ordered:
        updates: dict[str, Any] = {}
        ts = step.timestamp
        t_start = step.transcript_start if step.transcript_start is not None else ts
        t_end = step.transcript_end if step.transcript_end is not None else ts + 8.0

        if ts > content_end or t_start > content_end:
            if past_count > 1:
                target = spread_start + past_index * (
                    (content_end - 2.0 - spread_start) / max(1, past_count - 1)
                )
                anchor = max(0.0, min(target, content_end - 2.0))
            else:
                eff_start = max(0.0, min(t_start, content_end - 0.5))
                eff_end = max(eff_start + 0.5, min(t_end, content_end))
                if eff_end - eff_start >= 2.0:
                    anchor = max(eff_start + 0.5, min(ts, eff_end - 1.0))
                else:
                    anchor = max(0.0, min(ts, content_end - 3.0))
            if prepared:
                anchor = max(anchor, prepared[-1].timestamp + min_gap)
            anchor = max(0.0, min(anchor, content_end - 2.0))
            past_index += 1
            logger.warning(
                "Step %02d '%s': timestamp %.1fs past content end %.1fs — "
                "relocating vision search to ~%.1fs.",
                step.step_number, step.title, ts, content_end, anchor,
            )
            updates["timestamp"] = round(anchor, 3)
            updates["transcript_start"] = round(max(0.0, anchor - 12.0), 3)
            updates["transcript_end"] = round(min(content_end, anchor + 14.0), 3)

        if updates:
            step = step.model_copy(update=updates)
            ts = step.timestamp
            t_start = step.transcript_start or ts
            t_end = step.transcript_end or ts + 8.0

        if t_start > content_end:
            updates2 = {
                "transcript_start": round(max(0.0, content_end - 20.0), 3),
                "transcript_end": round(content_end, 3),
            }
            step = step.model_copy(update=updates2)

        if t_end > content_end:
            step = step.model_copy(update={"transcript_end": round(content_end, 3)})

        prepared.append(step)

    return prepared


def merge_confirmation_result(
    confidence: float,
    reason: str,
    is_wrong_step: bool,
    is_action_visible: bool,
    quality_flags: list[str],
    new_conf: float,
    new_reason: str,
    new_wrong: bool,
    new_action: bool,
    new_flags: list[str],
    step_number: int,
) -> tuple[float, str, bool, bool, list[str]]:
    """
  Apply a single-frame confirmation pass without trusting harmful swings.
  Rejects large drops and false confidence boosts when the frame is the wrong step.
    """
    if abs(new_conf - confidence) <= 0.05:
        return confidence, reason, is_wrong_step, is_action_visible, quality_flags

    mismatch_phrases = (
        "does not match", "does not depict", "does not show", "not show",
        "wrong step", "irrelevant", "fails to", "not visible", "not the described",
    )
    reason_lower = new_reason.lower()

    # False boost: wrong step, no action, or confirmation text contradicts a higher score.
    if new_conf > confidence and (
        new_wrong
        or not new_action
        or any(p in reason_lower for p in mismatch_phrases)
    ):
        logger.info(
            "Step %02d: confirmation boost rejected (%.2f -> %.2f, wrong=%s action=%s).",
            step_number, confidence, new_conf, new_wrong, new_action,
        )
        return confidence, reason, is_wrong_step, is_action_visible, quality_flags

    # Harmful drop: keep the original borderline pick instead of 0.05–0.15.
    if new_conf < confidence and (new_wrong or new_conf < VISION_CONFIRM_FLOOR):
        logger.info(
            "Step %02d: confirmation drop rejected (%.2f -> %.2f, wrong=%s), keeping %.2f.",
            step_number, confidence, new_conf, new_wrong, confidence,
        )
        return confidence, reason, is_wrong_step, is_action_visible, quality_flags

    applied_conf = new_conf
    if not new_wrong and new_conf < confidence and (confidence - new_conf) > VISION_CONFIRM_MAX_DROP:
        applied_conf = max(new_conf, confidence - VISION_CONFIRM_MAX_DROP)
        logger.info(
            "Step %02d: confirmation conf %.2f -> %.2f (clamped drop to %.2f).",
            step_number, confidence, applied_conf, confidence - applied_conf,
        )
    else:
        logger.info(
            "Step %02d: confirmation updated conf %.2f -> %.2f",
            step_number, confidence, applied_conf,
        )

    merged_reason = f"{reason} [confirmed: {new_reason}]"
    merged_flags = list(set(quality_flags + new_flags)) if new_flags else quality_flags
    return applied_conf, merged_reason, new_wrong, new_action, merged_flags


def normalize_visual_search_window(
    w_start: float,
    w_end: float,
    step: InstallationStep,
    video_duration: float,
) -> tuple[float, float]:
    """Guarantee w_start < w_end and both lie within searchable content."""
    safe_end = content_max_timestamp(video_duration)
    start = step.transcript_start if step.transcript_start is not None else step.timestamp
    end = step.transcript_end if step.transcript_end is not None else step.timestamp + 8.0
    anchor = min(step.timestamp, safe_end)

    w_start = min(w_start, safe_end)
    w_end = min(w_end, safe_end)

    if w_end <= w_start:
        half = min(14.0, MAX_TRANSCRIPT_WINDOW_SECONDS / 2.0)
        w_start = max(0.0, anchor - half)
        w_end = min(safe_end, anchor + half)
        if w_end <= w_start:
            w_end = min(safe_end, w_start + max(4.0, min(8.0, end - start)))
        logger.warning(
            "Step %02d: inverted/empty search window — using [%.1f, %.1f] around anchor %.1fs.",
            step.step_number, w_start, w_end, anchor,
        )

    return w_start, w_end


# ─────────────────────────────────────────────
# FRAME QUALITY SCORING
# ─────────────────────────────────────────────

def _check_prototype_banner_sync(image_path: Path) -> bool:
    if not PROTOTYPE_BANNER_CHECK:
        return False
    try:
        from PIL import Image
        import numpy as np

        with Image.open(image_path) as img:
            gray = np.array(img.convert("L").resize((320, 180), Image.BILINEAR), dtype=np.float32)

        top_rows = int(gray.shape[0] * PROTOTYPE_BANNER_TOP_FRACTION)
        region = gray[:top_rows, :]
        mean_bright = float(np.mean(region))
        std_bright = float(np.std(region))

        if mean_bright < PROTOTYPE_BANNER_MEAN_MAX and std_bright > PROTOTYPE_BANNER_CONTRAST_THRESHOLD:
            logger.debug(
                "Banner detected: mean=%.1f < %.0f and std=%.1f > %.0f in %s",
                mean_bright, PROTOTYPE_BANNER_MEAN_MAX, std_bright,
                PROTOTYPE_BANNER_CONTRAST_THRESHOLD, image_path.name,
            )
            return True
        return False
    except Exception:
        return False


def score_frame_sync(image_path: Path) -> tuple[float, float, float, bool]:
    try:
        from PIL import Image
        import numpy as np

        with Image.open(image_path) as img:
            rgb_img = img.convert("RGB")
            gray_img = rgb_img.convert("L")

            arr_gray = np.array(gray_img, dtype=np.float32)
            arr_rgb  = np.array(rgb_img,  dtype=np.float32)

            mean_brightness = float(np.mean(arr_gray))
            if mean_brightness < 60.0 and mean_brightness > 5.0:
                scale = min(2.5, 100.0 / mean_brightness)
                arr_gray = np.clip(arr_gray * scale, 0, 255)

            try:
                from scipy.ndimage import convolve as sci_convolve
                kernel = np.array([[0, 1, 0], [1, -4, 1], [0, 1, 0]], dtype=np.float32)
                lap = sci_convolve(arr_gray, kernel)
                sharpness = float(np.var(lap))
            except ImportError:
                dy = np.diff(arr_gray, axis=0)
                dx = np.diff(arr_gray, axis=1)
                dy_pad = np.pad(dy, ((0, 1), (0, 0)), mode="edge")
                dx_pad = np.pad(dx, ((0, 0), (0, 1)), mode="edge")
                grad_mag = np.sqrt(dy_pad ** 2 + dx_pad ** 2)
                sharpness = float(np.var(grad_mag)) * 8.0

            brightness = float(np.mean(np.array(gray_img, dtype=np.float32)))

            r_mean = float(np.mean(arr_rgb[:, :, 0]))
            g_mean = float(np.mean(arr_rgb[:, :, 1]))
            b_mean = float(np.mean(arr_rgb[:, :, 2]))

            channel_means = sorted([r_mean, g_mean, b_mean])
            dominant   = channel_means[2]
            others_avg = (channel_means[0] + channel_means[1]) / 2.0

            if others_avg >= COLOR_CAST_MIN_BRIGHTNESS:
                color_cast_ratio = dominant / others_avg
            else:
                color_cast_ratio = 1.0

            has_banner = _check_prototype_banner_sync(image_path)

            return sharpness, brightness, color_cast_ratio, has_banner

    except ImportError:
        return 100.0, 128.0, 1.0, False


def compute_quality_score(
    sharpness: float,
    brightness: float,
    color_cast_ratio: float = 1.0,
) -> float:
    sharpness_score   = min(sharpness / 500.0, 1.0)
    brightness_score  = max(0.0, 1.0 - abs(brightness - 128.0) / 128.0)
    cast_penalty_range = max(0.1, COLOR_CAST_MAX_RATIO - 1.0 + 0.5)
    cast_excess  = max(0.0, color_cast_ratio - 1.0)
    cast_penalty = min(1.0, cast_excess / cast_penalty_range)
    color_score  = 1.0 - cast_penalty
    return 0.55 * sharpness_score + 0.35 * brightness_score + 0.10 * color_score


async def score_frame(image_path: Path) -> tuple[float, float, float, bool]:
    return await asyncio.to_thread(score_frame_sync, image_path)


def filter_by_quality(
    candidates: list[FrameCandidate],
    reject_banners: bool = True,
) -> list[FrameCandidate]:
    if reject_banners:
        no_banner = [c for c in candidates if not c.has_prototype_banner]
        banner_count = len(candidates) - len(no_banner)
        if banner_count > 0:
            logger.debug("Prototype banner filter removed %d frames.", banner_count)
        if len(no_banner) >= MIN_FRAMES_AFTER_QUALITY_FILTER:
            candidates = no_banner

    good = [
        c for c in candidates
        if c.sharpness >= SHARPNESS_MIN
        and BRIGHTNESS_MIN <= c.brightness <= BRIGHTNESS_MAX
        and c.color_cast_ratio <= COLOR_CAST_MAX_RATIO
    ]

    if len(good) >= MIN_FRAMES_AFTER_QUALITY_FILTER:
        removed = len(candidates) - len(good)
        if removed:
            logger.debug(
                "Quality filter: kept %d / %d (removed %d; color-cast thresh=%.2f).",
                len(good), len(candidates), removed, COLOR_CAST_MAX_RATIO,
            )
        return good

    sorted_by_q = sorted(candidates, key=lambda c: c.quality_score, reverse=True)
    kept = sorted_by_q[:max(MIN_FRAMES_AFTER_QUALITY_FILTER, len(candidates) // 2)]
    logger.debug("Quality filter fallback: keeping top-%d by score.", len(kept))
    return kept


# ─────────────────────────────────────────────
# MOTION SCORING
# ─────────────────────────────────────────────

def compute_motion_scores_sync(candidates: list[FrameCandidate]) -> list[float]:
    if len(candidates) < 2:
        return [0.5] * len(candidates)

    try:
        from PIL import Image
        import numpy as np

        def load_gray(path: str) -> "np.ndarray":
            with Image.open(path) as img:
                small = img.convert("L").resize((160, 90), Image.BILINEAR)
                return np.array(small, dtype=np.float32)

        arrays = [load_gray(c.image_path) for c in candidates]
        diffs: list[float] = []

        for i, arr in enumerate(arrays):
            left  = arrays[i - 1] if i > 0 else arr
            right = arrays[i + 1] if i + 1 < len(arrays) else arr
            d_left  = float(np.mean(np.abs(arr - left)))
            d_right = float(np.mean(np.abs(arr - right)))
            diffs.append((d_left + d_right) / 2.0)

        max_diff = max(diffs) or 1.0
        return [d / max_diff for d in diffs]

    except ImportError:
        return [0.5] * len(candidates)


# ─────────────────────────────────────────────
# HISTOGRAM SIMILARITY
# ─────────────────────────────────────────────

def image_histogram_similarity_sync(path_a: str, path_b: str) -> float:
    try:
        from PIL import Image
        import numpy as np

        def normalised_hist(path: str) -> "np.ndarray":
            with Image.open(path) as img:
                arr = np.array(img.convert("RGB"), dtype=np.float32)
            hists = []
            for ch in range(3):
                h, _ = np.histogram(arr[:, :, ch], bins=HISTOGRAM_BINS, range=(0, 256))
                total = h.sum() + 1e-8
                hists.append(h / total)
            return np.concatenate(hists)

        ha = normalised_hist(path_a)
        hb = normalised_hist(path_b)
        similarity = float(np.sum(np.sqrt(ha * hb + 1e-10)))
        return min(1.0, similarity / 3.0)

    except Exception as exc:
        logger.debug("Histogram similarity error: %s", exc)
        return 0.0


# ─────────────────────────────────────────────
# FRAME CANDIDATE BUILDING
# ─────────────────────────────────────────────

def safe_max_timestamp(duration: float) -> float:
    return max(0.0, duration - SAFE_END_PADDING_SECONDS)


def content_max_timestamp(
    duration: float,
    segments: list[TranscriptSegment] | None = None,
) -> float:
    """Latest timestamp eligible for installation-action screenshots (skips outro)."""
    hard_cap = max(0.0, duration - SAFE_END_PADDING_SECONDS)
    exclusion = min(VIDEO_OUTRO_EXCLUSION_SECONDS, max(8.0, duration * 0.03))
    by_exclusion = max(0.0, hard_cap - exclusion)

    if not segments:
        return by_exclusion

    last_install_end = 0.0
    for seg in segments:
        if _is_marketing_outro_text(seg.text):
            break
        last_install_end = max(last_install_end, min(seg.end, hard_cap))

    if last_install_end > 0:
        return min(hard_cap, max(by_exclusion, last_install_end - 1.0))
    return by_exclusion


def unique_sorted_timestamps(values: list[float], duration: float) -> list[float]:
    cleaned = []
    max_ts = content_max_timestamp(duration)
    for v in values:
        if math.isnan(v) or math.isinf(v):
            continue
        cleaned.append(round(max(0.0, min(max_ts, float(v))), 3))
    return sorted(set(cleaned))


def shortlist_candidates_for_vision(
    candidates: list[FrameCandidate],
    step: InstallationStep,
    video_duration: float,
    max_frames: int = VISION_SHORTLIST_FRAMES,
) -> list[FrameCandidate]:
    """Keep the sharpest, most on-action frames so Gemini compares fewer, better options."""
    if len(candidates) <= max_frames:
        return candidates

    anchor = step.timestamp
    t_start = step.transcript_start if step.transcript_start is not None else anchor
    t_end = step.transcript_end if step.transcript_end is not None else anchor + 8.0
    mid = (t_start + t_end) / 2.0
    span_half = max(4.0, (t_end - t_start) / 2.0 + 2.0)
    content_end = content_max_timestamp(video_duration)

    def rank_score(c: FrameCandidate) -> float:
        dist = min(abs(c.timestamp - anchor), abs(c.timestamp - mid))
        proximity = max(0.0, 1.0 - dist / span_half)
        outro_penalty = 0.4 if c.timestamp > content_end - 4.0 else 0.0
        banner_penalty = 0.55 if c.has_prototype_banner else 0.0
        motion = c.motion_score if c.motion_score > 0 else 0.5
        return (
            0.42 * c.quality_score
            + 0.33 * proximity
            + 0.15 * motion
            + 0.10 * (1.0 if t_start <= c.timestamp <= t_end else 0.35)
        ) - outro_penalty - banner_penalty

    ranked = sorted(candidates, key=rank_score, reverse=True)
    shortlist = ranked[:max_frames]
    logger.debug(
        "Step %02d vision shortlist: %d → %d frames.",
        step.step_number, len(candidates), len(shortlist),
    )
    return shortlist


def pick_local_fallback_candidate(
    candidates: list[FrameCandidate],
    step: InstallationStep,
    video_duration: float,
) -> Optional[FrameCandidate]:
    """Best sharp frame near the step anchor when Gemini vision is unavailable."""
    if not candidates:
        return None
    content_end = content_max_timestamp(video_duration)
    in_content = [c for c in candidates if c.timestamp <= content_end]
    pool_base = in_content if in_content else candidates
    non_banner = [c for c in pool_base if not c.has_prototype_banner]
    pool = non_banner if non_banner else pool_base
    anchor = step.timestamp
    return max(
        pool,
        key=lambda c: (
            c.quality_score
            - 0.25 * abs(c.timestamp - anchor)
            - (0.5 if c.timestamp > content_end - 3.0 else 0.0)
        ),
    )


def should_soft_accept_borderline(
    selected: Optional[FrameCandidate],
    confidence: float,
    is_wrong_step: bool,
    is_action_visible: bool,
) -> bool:
    """Use a usable normal-pass frame instead of expensive rescue API calls."""
    if selected is None or not is_action_visible:
        return False
    if is_wrong_step:
        return False
    if CLIENT_DEMO_MODE:
        return False
    if confidence < VISION_SOFT_ACCEPT_CONFIDENCE:
        return False
    return True


def _vision_skip_extended_on_api_error() -> bool:
    if CLIENT_DEMO_MODE:
        return False
    return VISION_SKIP_EXTENDED_ON_API_ERROR


def _extended_fallback_min_confidence() -> float:
    return VISION_MIN_CONFIDENCE if CLIENT_DEMO_MODE else 0.35


def build_refine_timestamps(
    step: InstallationStep,
    video_duration: float,
    n_frames: int = LOW_CONFIDENCE_REFINE_FRAMES,
) -> list[float]:
    """Dense, action-centered sampling for the low-confidence refine pass."""
    start = step.transcript_start if step.transcript_start is not None else step.timestamp
    end = step.transcript_end if step.transcript_end is not None else step.timestamp + 8.0
    anchor = step.timestamp
    mid = (start + end) / 2.0
    duration = max(1.0, end - start)

    priority: list[float] = [
        anchor,
        mid,
        start,
        start + duration * 0.15,
        start + duration * 0.35,
        start + duration * 0.50,
        start + duration * 0.65,
        start + duration * 0.85,
        end,
        anchor + 0.4,
        anchor + 0.8,
        anchor + 1.2,
        max(start, anchor - 0.6),
        max(start, anchor - 1.2),
    ]
    priority = unique_sorted_timestamps(priority, video_duration)

    w_start = max(0.0, start - VISUAL_PRE_ROLL)
    w_end = min(content_max_timestamp(video_duration), end + VISUAL_POST_ROLL)
    if w_end <= w_start:
        w_end = min(content_max_timestamp(video_duration), w_start + 6.0)

    uniform: list[float] = []
    t = w_start
    interval = max(0.25, (w_end - w_start) / max(1, n_frames - 1))
    while t <= w_end:
        uniform.append(t)
        t += interval

    seen = set(priority)
    merged = list(priority)
    for ts in unique_sorted_timestamps(uniform, video_duration):
        if ts not in seen:
            merged.append(ts)
            seen.add(ts)

    return merged[:n_frames]


def reduce_to_n(values: list[float], n: int) -> list[float]:
    values = sorted(set(values))
    if n <= 0:
        return []
    if len(values) <= n:
        return values
    if n == 1:
        return [values[len(values) // 2]]
    return sorted(set(
        values[round(i * (len(values) - 1) / (n - 1))]
        for i in range(n)
    ))


def build_candidate_timestamps(
    step: InstallationStep,
    scene_cuts: list[SceneCut],
    video_duration: float,
    previous_step: Optional[InstallationStep] = None,
    next_step: Optional[InstallationStep] = None,
) -> list[float]:
    start = step.transcript_start if step.transcript_start is not None else step.timestamp
    end   = step.transcript_end   if step.transcript_end   is not None else step.timestamp + 8.0
    safe_end = content_max_timestamp(video_duration)
    mid = (start + end) / 2.0

    transcript_span = end - start
    if transcript_span < 12.0:
        anchor = step.timestamp
        half = 6.0
        start = max(0.0, anchor - half)
        end = min(safe_end, anchor + half)
        logger.debug(
            "Step %02d: tight transcript span %.1fs — re-centered on anchor %.2fs → [%.1f, %.1f].",
            step.step_number, transcript_span, anchor, start, end,
        )

    w_start = max(0.0,     start - VISUAL_PRE_ROLL)
    w_end   = min(safe_end, end  + VISUAL_POST_ROLL)

    if next_step is not None:
        next_step_start = next_step.transcript_start if next_step.transcript_start is not None else next_step.timestamp
        if next_step_start > step.timestamp + 0.5:
            next_step_cap = min(safe_end, next_step_start + 1.0)
            if w_end > next_step_cap:
                logger.debug(
                    "Step %02d: capping w_end from %.2fs → %.2fs (next step starts at %.2fs).",
                    step.step_number, w_end, next_step_cap, next_step_start,
                )
                w_end = next_step_cap

    reinstall_shifted = False
    if is_reinstall_step(step) and previous_step is not None:
        gap = step.timestamp - previous_step.timestamp
        if gap < REINSTALL_TRIGGER_GAP_SECONDS:
            next_start = next_step.transcript_start if next_step else safe_end
            max_shift  = max(0.0, next_start - w_end - 2.0) if next_step else REINSTALL_EXTRA_FORWARD_SEARCH
            shift = min(REINSTALL_EXTRA_FORWARD_SEARCH, gap * 1.5, max_shift)
            if shift > 1.0:
                logger.info(
                    "Step %02d (%s): reinstall type, only %.1fs after prev — "
                    "shifting search window +%.1fs.",
                    step.step_number, step.title, gap, shift,
                )
                w_start = min(safe_end, w_start + shift)
                w_end   = min(safe_end, w_end   + shift)
                reinstall_shifted = True

    w_start, w_end = normalize_visual_search_window(w_start, w_end, step, video_duration)

    window_span = w_end - w_start
    if window_span < 8.0:
        expansion = (8.0 - window_span) / 2.0
        w_start = max(0.0, w_start - expansion)
        w_end = min(safe_end, w_end + expansion)
        logger.debug(
            "Step %02d: narrow window %.1fs expanded to [%.1f, %.1f].",
            step.step_number, window_span, w_start, w_end,
        )

    step_duration = max(0.5, end - start)
    priority: list[float] = [
        step.timestamp,
        mid,
        start + step_duration * 0.40,
        start,
        start + 0.3,
        start + 0.6,
        start + 1.0,
        start + 1.5,
        start + step_duration * 0.25,
        start + step_duration * 0.50,
        start + step_duration * 0.75,
        max(start, end - 1.0),
        end,
        min(w_end, end + 1.0),
        min(w_end, end + 2.0),
    ]

    if reinstall_shifted:
        shift_center = (w_start + w_end) / 2.0
        priority.extend([
            w_start,
            w_start + (w_end - w_start) * 0.25,
            shift_center,
            w_start + (w_end - w_start) * 0.75,
            w_end,
        ])

    SCENE_CUT_FRAME_OFFSETS = [0.15, 0.35, 0.65, 1.0, 1.8]
    for cut in scene_cuts:
        if w_start - 0.5 <= cut.timestamp <= w_end + 0.5:
            for offset in SCENE_CUT_FRAME_OFFSETS:
                priority.append(cut.timestamp + offset)

    priority = unique_sorted_timestamps(priority, video_duration)

    max_priority_slots = MAX_CANDIDATE_FRAMES - UNIFORM_SAMPLE_RESERVED_SLOTS
    if len(priority) > max_priority_slots:
        priority = reduce_to_n(priority, max_priority_slots)

    sampled: list[float] = []
    t = w_start
    while t <= w_end:
        sampled.append(t)
        t += VISUAL_SAMPLE_INTERVAL
    sampled = unique_sorted_timestamps(sampled, video_duration)

    remaining = MAX_CANDIDATE_FRAMES - len(priority)
    if remaining > 0:
        sampled = reduce_to_n(sampled, remaining)
    else:
        sampled = []

    seen: set[float] = set(priority)
    final: list[float] = list(priority)
    for ts in sampled:
        if ts not in seen:
            final.append(ts)
            seen.add(ts)

    final = sorted(final[:MAX_CANDIDATE_FRAMES])
    logger.debug(
        "Step %02d candidates: %d window=[%.1f, %.1f] reinstall_shift=%s",
        step.step_number, len(final), w_start, w_end, reinstall_shifted,
    )
    return final


def build_multi_person_rescue_timestamps(
    step: InstallationStep,
    video_duration: float,
    previous_step: Optional[InstallationStep] = None,
    next_step: Optional[InstallationStep] = None,
) -> list[float]:
    t_start = step.transcript_start or step.timestamp
    t_end   = step.transcript_end   or (step.timestamp + 8.0)
    midpoint = (t_start + t_end) / 2.0
    safe_end = content_max_timestamp(video_duration)

    half = max(4.0, (t_end - t_start) / 2.0 + 3.0)
    w_start = max(0.0,     midpoint - half)
    w_end   = min(safe_end, midpoint + half)

    if next_step is not None:
        next_step_start = next_step.transcript_start if next_step.transcript_start is not None else next_step.timestamp
        w_end = min(w_end, next_step_start + 1.0)

    n = MULTI_PERSON_RESCUE_FRAMES
    if n <= 1:
        return [round(midpoint, 3)]

    values = [
        w_start + i * ((w_end - w_start) / (n - 1))
        for i in range(n)
    ]
    final = unique_sorted_timestamps(values, video_duration)
    logger.debug(
        "Step %02d multi-person rescue: %d timestamps window=[%.1f, %.1f]",
        step.step_number, len(final), w_start, w_end,
    )
    return final


def build_wide_rescue_timestamps(
    step: InstallationStep,
    video_duration: float,
    previous_step: Optional[InstallationStep] = None,
    next_step: Optional[InstallationStep] = None,
) -> list[float]:
    center   = step.timestamp
    safe_end = content_max_timestamp(video_duration)

    start = max(0.0,     center - WIDE_RESCUE_PRE_ROLL)
    end   = min(safe_end, center + WIDE_RESCUE_POST_ROLL)

    if previous_step is not None:
        prev_end = previous_step.transcript_end or previous_step.timestamp + 5.0
        start = max(start, max(0.0, prev_end - 2.0))

    if next_step is not None:
        next_start = next_step.transcript_start or next_step.timestamp
        end = min(end, next_start + 5.0, safe_end)

    if end <= start:
        start = max(0.0,     center - 15.0)
        end   = min(safe_end, center + 25.0)

    if WIDE_RESCUE_FRAMES <= 1:
        return [round((start + end) / 2.0, 3)]

    values = [
        start + i * ((end - start) / (WIDE_RESCUE_FRAMES - 1))
        for i in range(WIDE_RESCUE_FRAMES)
    ]
    final = unique_sorted_timestamps(values, video_duration)
    logger.debug("Step %02d rescue candidates: %d", step.step_number, len(final))
    return final


def build_extended_fallback_timestamps(
    step: InstallationStep,
    video_duration: float,
    previous_step: Optional[InstallationStep] = None,
    next_step: Optional[InstallationStep] = None,
) -> list[float]:
    start_ts = step.transcript_start or step.timestamp
    end_ts   = step.transcript_end   or step.timestamp + 8.0
    midpoint = (start_ts + end_ts) / 2.0
    safe_end = content_max_timestamp(video_duration)

    fb_start = max(0.0,     midpoint - EXTENDED_FALLBACK_PRE_ROLL)
    fb_end   = min(safe_end, midpoint + EXTENDED_FALLBACK_POST_ROLL)

    if previous_step is not None:
        prev_end = previous_step.transcript_end or previous_step.timestamp + 5.0
        fb_start = max(fb_start, max(0.0, prev_end - 2.0))

    if next_step is not None:
        next_start = next_step.transcript_start or next_step.timestamp
        fb_end = min(fb_end, next_start + 2.0)

    if fb_end - fb_start < 10.0:
        center = min((start_ts + end_ts) / 2.0, safe_end)
        fb_start = max(0.0, center - 8.0)
        fb_end = min(safe_end, center + 8.0)
        logger.debug(
            "Step %02d extended-fallback: boundary-capped window too small, "
            "using centered 16s window [%.1f, %.1f]",
            step.step_number, fb_start, fb_end,
        )

    fb_start, fb_end = normalize_visual_search_window(fb_start, fb_end, step, video_duration)

    n = max(4, EXTENDED_FALLBACK_FRAMES)
    if n <= 1:
        return [round(midpoint, 3)]

    values = [
        fb_start + i * ((fb_end - fb_start) / (n - 1))
        for i in range(n)
    ]
    final = unique_sorted_timestamps(values, video_duration)
    logger.debug(
        "Step %02d extended-fallback candidates: %d window=[%.1f, %.1f]",
        step.step_number, len(final), fb_start, fb_end,
    )
    return final


def build_global_rescue_timestamps(
    step: InstallationStep,
    video_duration: float,
    scene_cuts: list[SceneCut] | None = None,
    previous_screenshot_timestamp: Optional[float] = None,
) -> list[float]:
    """
    Whole-video rescue: sample frames uniformly across the entire content range,
    biased toward scene-cut boundaries. Used as a last resort when the transcript
    anchor turns out to be wrong (e.g., outro voice-over referencing an earlier
    action). The vision model uses the step's visual_query to find the right scene
    independently of Claude's anchor.
    """
    safe_end = content_max_timestamp(video_duration)
    intro_skip = 5.0  # skip the very-beginning host intro / banner
    start = min(intro_skip, max(0.0, safe_end - 10.0))
    end   = max(start + 10.0, safe_end)

    n = max(8, WHOLE_VIDEO_RESCUE_FRAMES)
    uniform = [
        start + i * ((end - start) / (n - 1))
        for i in range(n)
    ]

    extras: list[float] = []
    if scene_cuts:
        for cut in scene_cuts:
            ts = float(getattr(cut, "timestamp", 0.0) or 0.0)
            if start <= ts <= end:
                extras.append(ts + 0.5)

    skip_window: tuple[float, float] | None = None
    if previous_screenshot_timestamp is not None:
        skip_window = (
            previous_screenshot_timestamp - 1.0,
            previous_screenshot_timestamp + 1.0,
        )

    combined = uniform + extras
    if skip_window:
        combined = [t for t in combined if not (skip_window[0] <= t <= skip_window[1])]

    final = unique_sorted_timestamps(combined, video_duration)
    logger.debug(
        "Step %02d global-rescue candidates: %d window=[%.1f, %.1f] (scene_extras=%d)",
        step.step_number, len(final), start, end, len(extras),
    )
    return final


async def extract_frame_at_timestamp(
    video_path: Path,
    timestamp: float,
    output_path: Path,
    *,
    for_final_screenshot: bool = False,
) -> None:
    """
    Extract a single frame. Candidates use input-seek (fast) at FRAME_EXTRACT_WIDTH.
    Final step images use FINAL_SCREENSHOT_WIDTH and a *hybrid* seek strategy:
      - fast-seek to (timestamp - PRE_SEEK_SECONDS)  → keyframe boundary, near-instant
      - then output-seek the remaining PRE_SEEK_SECONDS for frame accuracy
    This keeps accuracy while avoiding O(video-length) full decodes on long videos
    (the previous pure output-seek could exceed 180s on 12+ minute videos).
    If the accurate attempt fails or times out, falls back to a pure fast input-seek
    so a single bad screenshot never kills the pipeline.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    scale_w = FINAL_SCREENSHOT_WIDTH if for_final_screenshot else FRAME_EXTRACT_WIDTH
    vf = f"scale={scale_w}:-2,format=yuvj420p"
    jpeg_q = "2" if for_final_screenshot else "3"

    want_accurate = bool(for_final_screenshot and FRAME_FINAL_ACCURATE_SEEK)

    def _fast_cmd() -> list[str]:
        return [
            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
            "-ss", f"{timestamp:.3f}",
            "-i", str(video_path),
            "-frames:v", "1", "-vf", vf,
            "-update", "1", "-q:v", jpeg_q,
            str(output_path),
        ]

    def _hybrid_cmd() -> list[str]:
        pre = max(0.0, FRAME_FINAL_PRE_SEEK_SECONDS)
        fast_to = max(0.0, timestamp - pre)
        out_seek = timestamp - fast_to
        return [
            "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
            "-ss", f"{fast_to:.3f}",
            "-i", str(video_path),
            "-ss", f"{out_seek:.3f}",
            "-frames:v", "1", "-vf", vf,
            "-update", "1", "-q:v", jpeg_q,
            str(output_path),
        ]

    if want_accurate:
        try:
            await run_cmd(
                _hybrid_cmd(),
                step_label=f"frame_{timestamp:.3f}",
                timeout=FFMPEG_FINAL_SCREENSHOT_TIMEOUT_SEC,
            )
            if output_path.exists() and output_path.stat().st_size > 0:
                return
            logger.warning(
                "Hybrid-seek extraction produced empty file at %.3fs — falling back to fast seek.",
                timestamp,
            )
        except (TimeoutError, subprocess.CalledProcessError) as exc:
            logger.warning(
                "Hybrid-seek ffmpeg failed at %.3fs (%s) — falling back to fast seek.",
                timestamp, exc,
            )

    await run_cmd(
        _fast_cmd(),
        step_label=f"frame_{timestamp:.3f}",
        timeout=FFMPEG_CANDIDATE_FRAME_TIMEOUT_SEC,
    )
    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError(f"Empty frame: {output_path}")


async def extract_candidate_frames(
    video_path: Path,
    step: InstallationStep,
    timestamps: list[float],
    candidates_dir: Path,
) -> list[FrameCandidate]:
    labels   = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    step_dir = candidates_dir / f"step_{step.step_number:02d}"
    step_dir.mkdir(parents=True, exist_ok=True)

    semaphore = asyncio.Semaphore(SCREENSHOT_CONCURRENCY)

    async def one(label: str, ts: float) -> Optional[FrameCandidate]:
        async with semaphore:
            path = step_dir / f"{label}_{ts:.3f}.jpg"
            try:
                if not path.exists() or path.stat().st_size == 0:
                    await extract_frame_at_timestamp(video_path, ts, path)

                sharpness, brightness, color_cast_ratio, has_banner = await score_frame(path)
                quality_score = compute_quality_score(sharpness, brightness, color_cast_ratio)

                return FrameCandidate(
                    label=label, timestamp=ts, image_path=str(path),
                    sharpness=sharpness, brightness=brightness,
                    color_cast_ratio=color_cast_ratio,
                    quality_score=quality_score,
                    has_prototype_banner=has_banner,
                )
            except Exception as exc:
                logger.warning(
                    "Step %02d frame %s@%.2fs failed: %s",
                    step.step_number, label, ts, exc,
                )
                return None

    raw_results = await asyncio.gather(
        *[one(labels[i], ts) for i, ts in enumerate(timestamps[:len(labels)])]
    )
    candidates = [r for r in raw_results if r is not None]

    if len(candidates) >= 2:
        motion_scores = await asyncio.to_thread(compute_motion_scores_sync, candidates)
        updated: list[FrameCandidate] = []
        for cand, ms in zip(candidates, motion_scores):
            blended_q = cand.quality_score * 0.85 + ms * 0.15
            updated.append(cand.model_copy(update={"motion_score": ms, "quality_score": blended_q}))
        candidates = updated

    return candidates


# ─────────────────────────────────────────────
# OPENAI VISION
# ─────────────────────────────────────────────

def _build_vision_prompt(
    step: "InstallationStep",
    candidates: list["FrameCandidate"],
    previous_step_title: str = "",
    next_step_title: str = "",
    confirmation_mode: bool = False,
    confirmation_label: Optional[str] = None,
    previous_screenshot_timestamp: Optional[float] = None,
    requires_two_people: bool = False,
) -> str:
    """Shared prompt builder for OpenAI / Gemini vision selectors."""
    candidate_text = "\n".join(f"{c.label}: {c.timestamp:.2f}s" for c in candidates)

    neighbor_context = ""
    if previous_step_title:
        neighbor_context += f"PREVIOUS STEP (do NOT pick frames for this): {previous_step_title}\n"
    if next_step_title:
        neighbor_context += f"NEXT STEP (do NOT pick frames for this): {next_step_title}\n"

    prev_ss_warning = ""
    if previous_screenshot_timestamp is not None:
        prev_ss_warning = (
            f"\n\u26a0 IMPORTANT \u2014 The previous step's screenshot was taken at "
            f"{previous_screenshot_timestamp:.2f}s.  Do NOT select a frame from "
            f"within \u00b12 seconds of that timestamp unless it clearly shows a "
            f"distinctly different action or camera angle.\n"
        )

    two_person_note = ""
    if requires_two_people:
        two_person_note = (
            "\nTWO-PERSON STEP \u2014 This step explicitly requires two people. "
            "Strongly prefer a frame where BOTH installers are visible and actively "
            "working together. If no such frame exists in the candidates, choose the "
            "frame that best shows the action even with only one person, and set "
            "confidence to 0.65-0.75 to flag for manual review.\n"
        )

    if confirmation_mode and confirmation_label:
        return (
            "You previously selected this frame as the best match for a step in an "
            "installation manual. Look carefully at this single image and re-evaluate.\n\n"
            f"STEP NUMBER: {step.step_number}\n"
            f"STEP TITLE: {step.title}\n"
            f"STEP DESCRIPTION: {step.description}\n"
            f"VISUAL TARGET: {step.visual_query}\n\n"
            f"Frame label: {confirmation_label}\n\n"
            "Is this frame actually showing the described installation action or the "
            "correct physical component? Give an honest re-assessment.\n\n"
            "Return ONLY valid JSON:\n"
            "{"
            f'"best_label":"{confirmation_label}",'
            '"confidence":0.85,'
            '"reason":"honest re-assessment",'
            '"is_action_visible":true,'
            '"is_wrong_step":false,'
            '"quality_flags":["clear"]'
            "}\n"
        )

    return (
        "You are selecting the best screenshot frame for a PRINTED INSTALLATION MANUAL.\n\n"
        "Each frame is preceded by its Label and Timestamp. "
        "Evaluate each decoded candidate JPEG carefully.\n\n"
        "You MUST choose exactly ONE frame from the provided labels. "
        "Return NONE only if every single frame is completely black or blank.\n\n"
        f"{neighbor_context}"
        f"{prev_ss_warning}"
        f"{two_person_note}\n"
        "=== MANDATORY REJECTION ===\n"
        "Reject any frame containing a 'prototype only', 'working prototype',\n"
        "'pre-production', 'does not reflect final product', or 'prototype'\n"
        "text overlay -- even if partially cut off or semi-transparent.\n"
        "Also reject: completely black/white frames, video transitions,\n"
        "solid color-tint frames, talking-head-only frames, title cards,\n"
        "sponsor slates, subscribe/contact overlays, social end screens, and\n"
        "branded outros — even if they are the only non-black frames.\n\n"
        "=== SELECTION PRIORITY (highest = best) ===\n"
        "1. The described action is ACTIVELY HAPPENING -- hand in motion, tool\n"
        "   engaging fastener, part being positioned, component being lifted.\n"
        "   Hands and hardware must BOTH be clearly visible.\n"
        "   Prefer a frame whose timestamp falls INSIDE the TRANSCRIPT WINDOW;\n"
        "   if unsure, choose the frame closest to the middle of that window.\n"
        "2. A tight close-up of the correct part at the moment of the action.\n"
        "3. The immediately completed state -- part just installed/tightened.\n"
        "4. A clear wide shot showing the relevant vehicle area with the part.\n\n"
        "=== PARTIAL MATCH ===\n"
        "If the exact action is not visible, choose the CLOSEST useful frame:\n"
        "- Relevant part close-up without active motion -> confidence 0.60-0.72\n"
        "- Completed installed state -> confidence 0.68-0.80\n"
        "- Wide context shot of correct area -> confidence 0.45-0.58\n"
        "- Only one person visible when two are needed -> confidence 0.60-0.72\n"
        "DO NOT return NONE just because the action is not perfectly framed.\n"
        "If every frame looks like an outro/sponsor/title card, still pick the frame\n"
        "whose timestamp is closest to the TRANSCRIPT WINDOW and shows any relevant\n"
        "vehicle hardware or install context — never a branded end slate.\n\n"
        "=== STEP DETAILS ===\n"
        f"STEP NUMBER: {step.step_number}\n"
        f"STEP TITLE: {step.title}\n"
        f"STEP DESCRIPTION: {step.description}\n"
        f"VISUAL TARGET: {step.visual_query}\n"
        f"STEP TIMESTAMP: {step.timestamp:.2f}s\n"
        f"TRANSCRIPT WINDOW: {step.transcript_start}s - {step.transcript_end}s\n\n"
        f"CANDIDATES:\n{candidate_text}\n\n"
        "Return ONLY valid JSON with this exact shape (no markdown, no text before/after):\n"
        '{"best_label":"A","confidence":0.85,"reason":"brief visual reason",'
        '"is_action_visible":true,"is_wrong_step":false,"quality_flags":["clear","exact_action"]}\n\n'
        "Confidence guide:\n"
        "0.90-1.00 = exact action visible, both hands and hardware clearly shown\n"
        "0.75-0.89 = clearly useful close-up or just-completed state\n"
        "0.55-0.74 = partial/context match, still informative\n"
        "0.40-0.54 = weak but usable -- better than no screenshot\n"
        "0.25-0.39 = very weak fallback -- flag for manual replacement\n"
    )


def select_frame_with_openai_vision_sync(
    step: "InstallationStep",
    candidates: list["FrameCandidate"],
    previous_step_title: str = "",
    next_step_title: str = "",
    confirmation_mode: bool = False,
    confirmation_label: Optional[str] = None,
    previous_screenshot_timestamp: Optional[float] = None,
    requires_two_people: bool = False,
    model: str | None = None,
) -> dict[str, Any]:
    """Pick the best candidate frame for ``step`` using OpenAI vision (gpt-4o-class)."""
    import base64

    client = _openai_client()
    active_model = (model or OPENAI_VISION_MODEL).strip()

    prompt = _build_vision_prompt(
        step, candidates,
        previous_step_title=previous_step_title,
        next_step_title=next_step_title,
        confirmation_mode=confirmation_mode,
        confirmation_label=confirmation_label,
        previous_screenshot_timestamp=previous_screenshot_timestamp,
        requires_two_people=requires_two_people,
    )

    content_parts: list[dict[str, Any]] = [{"type": "text", "text": prompt}]
    for c in candidates:
        try:
            image_bytes = Path(c.image_path).read_bytes()
        except Exception as exc:
            logger.warning(
                "OpenAI vision: failed to read candidate %s (%s) — skipping.",
                c.image_path, exc,
            )
            continue
        b64 = base64.b64encode(image_bytes).decode("ascii")
        content_parts.append({
            "type": "text",
            "text": f"Label: {c.label} | Timestamp: {c.timestamp:.2f}s",
        })
        content_parts.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{b64}",
                "detail": OPENAI_VISION_IMAGE_DETAIL,
            },
        })

    response = client.chat.completions.create(
        model=active_model,
        messages=[{"role": "user", "content": content_parts}],
        temperature=0,
        response_format={"type": "json_object"},
        max_tokens=400,
    )

    try:
        raw = (response.choices[0].message.content or "{}").strip()
    except Exception as exc:
        raise RuntimeError(f"OpenAI vision returned empty response: {exc}") from exc

    data = parse_json_response(raw)
    logger.debug(
        "OpenAI vision step %02d (%s): %s",
        step.step_number, active_model, raw[:300],
    )

    if not isinstance(data, dict):
        raise ValueError(f"OpenAI vision returned non-object: {data}")

    return data


# ─────────────────────────────────────────────
# GEMINI VISION (legacy — kept for optional fallback)
# ─────────────────────────────────────────────

def select_frame_with_gemini_vision_sync(
    step: InstallationStep,
    candidates: list[FrameCandidate],
    previous_step_title: str = "",
    next_step_title: str = "",
    confirmation_mode: bool = False,
    confirmation_label: Optional[str] = None,
    previous_screenshot_timestamp: Optional[float] = None,
    requires_two_people: bool = False,
    model: str | None = None,
) -> dict[str, Any]:
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        raise RuntimeError("google-genai not installed. Run: pip install google-genai")

    client = genai.Client(api_key=GEMINI_API_KEY)

    candidate_text = "\n".join(f"{c.label}: {c.timestamp:.2f}s" for c in candidates)

    neighbor_context = ""
    if previous_step_title:
        neighbor_context += f"PREVIOUS STEP (do NOT pick frames for this): {previous_step_title}\n"
    if next_step_title:
        neighbor_context += f"NEXT STEP (do NOT pick frames for this): {next_step_title}\n"

    prev_ss_warning = ""
    if previous_screenshot_timestamp is not None:
        prev_ss_warning = (
            f"\n\u26a0 IMPORTANT \u2014 The previous step's screenshot was taken at "
            f"{previous_screenshot_timestamp:.2f}s.  Do NOT select a frame from "
            f"within \u00b12 seconds of that timestamp unless it clearly shows a "
            f"distinctly different action or camera angle.\n"
        )

    two_person_note = ""
    if requires_two_people:
        two_person_note = (
            "\nTWO-PERSON STEP \u2014 This step explicitly requires two people. "
            "Strongly prefer a frame where BOTH installers are visible and actively "
            "working together. If no such frame exists in the candidates, choose the "
            "frame that best shows the action even with only one person, and set "
            "confidence to 0.65-0.75 to flag for manual review.\n"
        )

    if confirmation_mode and confirmation_label:
        prompt = (
            "You previously selected this frame as the best match for a step in an "
            "installation manual. Look carefully at this single image and re-evaluate.\n\n"
            f"STEP NUMBER: {step.step_number}\n"
            f"STEP TITLE: {step.title}\n"
            f"STEP DESCRIPTION: {step.description}\n"
            f"VISUAL TARGET: {step.visual_query}\n\n"
            f"Frame label: {confirmation_label}\n\n"
            "Is this frame actually showing the described installation action or the "
            "correct physical component? Give an honest re-assessment.\n\n"
            "Return ONLY valid JSON:\n"
            "{"
            f'"best_label":"{confirmation_label}",'
            '"confidence":0.85,'
            '"reason":"honest re-assessment",'
            '"is_action_visible":true,'
            '"is_wrong_step":false,'
            '"quality_flags":["clear"]'
            "}\n"
        )
    else:
        prompt = (
            "You are selecting the best screenshot frame for a PRINTED INSTALLATION MANUAL.\n\n"
            "Each frame is preceded by its Label and Timestamp. "
            "Evaluate each decoded candidate JPEG carefully.\n\n"
            "You MUST choose exactly ONE frame from the provided labels. "
            "Return NONE only if every single frame is completely black or blank.\n\n"
            f"{neighbor_context}"
            f"{prev_ss_warning}"
            f"{two_person_note}\n"
            "=== MANDATORY REJECTION ===\n"
            "Reject any frame containing a 'prototype only', 'working prototype',\n"
            "'pre-production', 'does not reflect final product', or 'prototype'\n"
            "text overlay -- even if partially cut off or semi-transparent.\n"
            "Also reject: completely black/white frames, video transitions,\n"
            "solid color-tint frames, talking-head-only frames, title cards,\n"
            "sponsor slates, subscribe/contact overlays, social end screens, and\n"
            "branded outros — even if they are the only non-black frames.\n\n"
            "=== SELECTION PRIORITY (highest = best) ===\n"
            "1. The described action is ACTIVELY HAPPENING -- hand in motion, tool\n"
            "   engaging fastener, part being positioned, component being lifted.\n"
            "   Hands and hardware must BOTH be clearly visible.\n"
            "   Prefer a frame whose timestamp falls INSIDE the TRANSCRIPT WINDOW;\n"
            "   if unsure, choose the frame closest to the middle of that window.\n"
            "2. A tight close-up of the correct part at the moment of the action.\n"
            "3. The immediately completed state -- part just installed/tightened.\n"
            "4. A clear wide shot showing the relevant vehicle area with the part.\n\n"
            "=== PARTIAL MATCH ===\n"
            "If the exact action is not visible, choose the CLOSEST useful frame:\n"
            "- Relevant part close-up without active motion -> confidence 0.60-0.72\n"
            "- Completed installed state -> confidence 0.68-0.80\n"
            "- Wide context shot of correct area -> confidence 0.45-0.58\n"
            "- Only one person visible when two are needed -> confidence 0.60-0.72\n"
            "DO NOT return NONE just because the action is not perfectly framed.\n"
            "If every frame looks like an outro/sponsor/title card, still pick the frame\n"
            "whose timestamp is closest to the TRANSCRIPT WINDOW and shows any relevant\n"
            "vehicle hardware or install context — never a branded end slate.\n\n"
            "=== STEP DETAILS ===\n"
            f"STEP NUMBER: {step.step_number}\n"
            f"STEP TITLE: {step.title}\n"
            f"STEP DESCRIPTION: {step.description}\n"
            f"VISUAL TARGET: {step.visual_query}\n"
            f"STEP TIMESTAMP: {step.timestamp:.2f}s\n"
            f"TRANSCRIPT WINDOW: {step.transcript_start}s - {step.transcript_end}s\n\n"
            f"CANDIDATES:\n{candidate_text}\n\n"
            "Return ONLY valid JSON with this exact shape (no markdown, no text before/after):\n"
            '{"best_label":"A","confidence":0.85,"reason":"brief visual reason",'
            '"is_action_visible":true,"is_wrong_step":false,"quality_flags":["clear","exact_action"]}\n\n'
            "Confidence guide:\n"
            "0.90-1.00 = exact action visible, both hands and hardware clearly shown\n"
            "0.75-0.89 = clearly useful close-up or just-completed state\n"
            "0.55-0.74 = partial/context match, still informative\n"
            "0.40-0.54 = weak but usable -- better than no screenshot\n"
            "0.25-0.39 = very weak fallback -- flag for manual replacement\n"
        )

    contents: list[Any] = [prompt]
    for c in candidates:
        contents.append(f"Label: {c.label} | Timestamp: {c.timestamp:.2f}s")
        contents.append(
            types.Part.from_bytes(
                data=Path(c.image_path).read_bytes(),
                mime_type="image/jpeg",
            )
        )

    active_model = model or GEMINI_VISION_MODEL
    response = client.models.generate_content(
        model=active_model,
        contents=contents,
        config=types.GenerateContentConfig(
            temperature=0,
            response_mime_type="application/json",
        ),
    )

    raw  = (response.text or "{}").strip()
    data = parse_json_response(raw)
    logger.debug(
        "Gemini vision step %02d (%s): %s",
        step.step_number, active_model, raw[:300],
    )

    if not isinstance(data, dict):
        raise ValueError(f"Gemini vision returned non-object: {data}")

    return data


def interpret_vision_decision(
    decision: dict[str, Any],
    candidates: list[FrameCandidate],
) -> tuple[Optional[FrameCandidate], float, str, bool, bool, list[str]]:
    best_label       = str(decision.get("best_label", "")).strip().upper()
    confidence       = float(decision.get("confidence", 0.0) or 0.0)
    reason           = str(decision.get("reason", "")).strip()
    is_wrong_step    = bool(decision.get("is_wrong_step", False))
    is_action_visible = bool(decision.get("is_action_visible", False))
    raw_flags        = decision.get("quality_flags", [])
    quality_flags    = [str(x) for x in raw_flags] if isinstance(raw_flags, list) else []
    by_label         = {c.label.upper(): c for c in candidates}

    if best_label in {"", "NONE", "NULL", "N/A"}:
        return None, max(confidence, 0.01), reason or "Vision did not choose a label.", \
               is_wrong_step, is_action_visible, quality_flags

    selected = by_label.get(best_label)
    if selected is None:
        return None, 0.01, f"Invalid label {best_label!r}.", True, False, quality_flags

    return selected, confidence, reason, is_wrong_step, is_action_visible, quality_flags


# ─────────────────────────────────────────────
# SCREENSHOT SELECTION ORCHESTRATION
# ─────────────────────────────────────────────

async def select_best_frame_for_step(
    video_path: Path,
    step: InstallationStep,
    scene_cuts: list[SceneCut],
    video_duration: float,
    candidates_root: Path,
    screenshots_dir: Path,
    previous_step: Optional[InstallationStep] = None,
    next_step: Optional[InstallationStep] = None,
    previous_screenshot_timestamp: Optional[float] = None,
) -> InstallationStep:

    prev_title = previous_step.title if previous_step else ""
    next_title = next_step.title     if next_step     else ""
    needs_two = is_multi_person_step(step)

    async def commit_frame(
        selected: FrameCandidate,
        confidence: float,
        reason: str,
        log_tag: str,
    ) -> InstallationStep:
        screenshots_dir.mkdir(parents=True, exist_ok=True)
        final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
        try:
            await extract_frame_at_timestamp(
                video_path, selected.timestamp, final_path, for_final_screenshot=True,
            )
        except Exception as exc:
            # Last-ditch: reuse the already-extracted candidate jpg so the pipeline
            # keeps running even if the high-res re-extraction fails entirely.
            logger.warning(
                "Step %02d final screenshot extraction failed (%s); reusing candidate %s.",
                step.step_number, exc, selected.image_path,
            )
            try:
                src = Path(selected.image_path)
                if src.exists() and src.stat().st_size > 0:
                    shutil.copyfile(src, final_path)
                else:
                    raise RuntimeError(f"Candidate frame missing: {src}")
            except Exception as copy_exc:
                logger.error(
                    "Step %02d could not even reuse candidate frame (%s).",
                    step.step_number, copy_exc,
                )
                return step.model_copy(update={
                    "selected_frame_label":      selected.label,
                    "selected_frame_reason":     f"{reason} [extract failed: {exc}]",
                    "selected_frame_confidence": 0.0,
                    "screenshot_timestamp":      selected.timestamp,
                    "screenshot_path":           "",
                })
        logger.info(
            "Step %02d %s frame %s @%.2fs conf=%.2f",
            step.step_number, log_tag, selected.label,
            selected.timestamp, confidence,
        )
        return step.model_copy(update={
            "selected_frame_label":      selected.label,
            "selected_frame_reason":     reason,
            "selected_frame_confidence": confidence,
            "screenshot_timestamp":      selected.timestamp,
            "screenshot_path":           str(final_path).replace("\\", "/"),
        })

    async def run_vision_pass(
        timestamps: list[float],
        pass_name: str,
        apply_quality_filter: bool = True,
        two_person: bool = False,
        max_vision_frames: int | None = None,
    ) -> tuple[
        Optional[FrameCandidate], float, str, bool, bool, list[str], list[FrameCandidate],
    ]:
        pass_root      = candidates_root / pass_name
        all_candidates = await extract_candidate_frames(video_path, step, timestamps, pass_root)

        if not all_candidates:
            return None, 0.0, f"{pass_name}: no candidates extracted.", True, False, [], []

        candidates = filter_by_quality(all_candidates) if apply_quality_filter else all_candidates
        if not candidates:
            candidates = all_candidates
        shortlist_max = max_vision_frames or (
            VISION_SHORTLIST_FRAMES
            if pass_name == "normal"
            else VISION_RESCUE_SHORTLIST_FRAMES
        )
        candidates = shortlist_candidates_for_vision(
            candidates, step, video_duration, max_frames=shortlist_max,
        )

        def local_fallback_result(reason_suffix: str) -> tuple[
            Optional[FrameCandidate], float, str, bool, bool, list[str], list[FrameCandidate],
        ]:
            local = pick_local_fallback_candidate(all_candidates, step, video_duration)
            if local is None:
                return (
                    None, 0.0, f"{pass_name}: {reason_suffix}",
                    True, False, [], all_candidates,
                )
            conf = max(PARTIAL_MATCH_MIN_CONFIDENCE, VISION_SOFT_ACCEPT_CONFIDENCE)
            return (
                local, conf,
                f"{pass_name}: local fallback ({reason_suffix})",
                False, True, ["local_fallback"], all_candidates,
            )

        try:
            decision = await _ai_with_retry(
                select_frame_with_openai_vision_sync,
                step, candidates,
                prev_title, next_title,
                False, None,
                previous_screenshot_timestamp,
                two_person,
                step_label=f"vision-step{step.step_number:02d}-{pass_name}",
                max_retries=OPENAI_VISION_MAX_RETRIES,
                base_delay=OPENAI_VISION_BASE_DELAY,
                timeout_sec=OPENAI_VISION_TIMEOUT_SEC,
                models=OPENAI_VISION_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
            )

            selected, confidence, reason, is_wrong_step, is_action_visible, quality_flags = \
                interpret_vision_decision(decision, candidates)

            if (
                pass_name == "normal"
                and ENABLE_VISION_CONFIRMATION
                and selected is not None
                and VISION_CONFIRM_THRESHOLD > confidence >= VISION_CONFIRM_FLOOR
                and not is_wrong_step
            ):
                logger.debug(
                    "Step %02d: confidence %.2f borderline — running confirmation pass.",
                    step.step_number, confidence,
                )
                try:
                    confirm_decision = await _ai_with_retry(
                        select_frame_with_openai_vision_sync,
                        step, [selected],
                        prev_title, next_title,
                        True, selected.label,
                        previous_screenshot_timestamp,
                        two_person,
                        step_label=f"confirm-step{step.step_number:02d}",
                        max_retries=OPENAI_VISION_MAX_RETRIES,
                        base_delay=OPENAI_VISION_BASE_DELAY,
                        timeout_sec=min(OPENAI_VISION_TIMEOUT_SEC, 90.0),
                        models=OPENAI_VISION_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
                    )
                    _, new_conf, new_reason, new_wrong, new_action, new_flags = \
                        interpret_vision_decision(confirm_decision, [selected])

                    confidence, reason, is_wrong_step, is_action_visible, quality_flags = \
                        merge_confirmation_result(
                            confidence, reason, is_wrong_step, is_action_visible,
                            quality_flags,
                            new_conf, new_reason, new_wrong, new_action, new_flags,
                            step.step_number,
                        )
                except Exception as exc:
                    logger.debug("Confirmation pass error: %s", exc)

            return (
                selected, confidence,
                f"{pass_name}: {reason}",
                is_wrong_step, is_action_visible, quality_flags,
                all_candidates,
            )

        except Exception as exc:
            logger.warning(
                "Step %02d %s vision failed after retries: %s",
                step.step_number, pass_name, exc,
            )
            if CLIENT_DEMO_MODE:
                return (
                    None, 0.0, f"{pass_name}: vision error: {exc}",
                    True, False, ["vision_api_error"], all_candidates,
                )
            return local_fallback_result(f"vision error: {exc}")

    def is_good_enough(
        selected: Optional[FrameCandidate],
        confidence: float,
        is_wrong_step: bool,
        flags: list[str],
        two_person: bool = False,
    ) -> bool:
        if selected is None:
            return False
        if "local_fallback" in flags or "vision_api_error" in flags:
            return False
        if is_wrong_step and confidence < 0.78:
            return False
        threshold = MULTI_PERSON_RESCUE_CONFIDENCE if two_person else VISION_MIN_CONFIDENCE
        if confidence >= threshold:
            return True
        if ALLOW_PARTIAL_VISUAL_MATCH and confidence >= PARTIAL_MATCH_MIN_CONFIDENCE:
            useful_flags = {
                "partial_match", "completed_state", "close_up", "clear",
                "context", "context_shot", "action_implied", "relevant_part",
                "relevant_tool", "exact_action", "exact_match",
            }
            if any(f in useful_flags for f in flags):
                return True
        return False

    # Pass 1: Normal window
    normal_timestamps = build_candidate_timestamps(
        step, scene_cuts, video_duration, previous_step, next_step,
    )
    (
        normal_selected, normal_confidence, normal_reason,
        normal_wrong_step, normal_action_visible, normal_flags,
        normal_candidates,
    ) = await run_vision_pass(normal_timestamps, "normal", two_person=needs_two)

    if is_good_enough(normal_selected, normal_confidence, normal_wrong_step, normal_flags, needs_two):
        return await commit_frame(
            normal_selected, normal_confidence, normal_reason, "\u2713 normal",
        )

    if should_soft_accept_borderline(
        normal_selected, normal_confidence, normal_wrong_step, normal_action_visible,
    ):
        logger.info(
            "Step %02d: soft-accepting borderline normal frame (conf=%.2f, wrong_step=%s).",
            step.step_number, normal_confidence, normal_wrong_step,
        )
        return await commit_frame(
            normal_selected,
            normal_confidence,
            f"{normal_reason} (soft-accepted to avoid rescue API calls)",
            "\u2713 normal~",
        )

    logger.warning(
        "Step %02d \u2717 normal  conf=%.2f wrong_step=%s",
        step.step_number, normal_confidence, normal_wrong_step,
    )

    # Pass 2.5: Multi-person rescue
    if needs_two and normal_confidence < MULTI_PERSON_RESCUE_CONFIDENCE:
        logger.info(
            "Step %02d: two-person step with conf=%.2f — trying multi-person rescue.",
            step.step_number, normal_confidence,
        )
        mp_timestamps = build_multi_person_rescue_timestamps(
            step=step, video_duration=video_duration,
            previous_step=previous_step, next_step=next_step,
        )
        (
            mp_selected, mp_confidence, mp_reason,
            mp_wrong_step, mp_action_visible, mp_flags,
            mp_candidates,
        ) = await run_vision_pass(mp_timestamps, "multi_person_rescue",
                                   apply_quality_filter=False, two_person=True)

        if mp_selected and mp_confidence >= 0.55:
            screenshots_dir.mkdir(parents=True, exist_ok=True)
            final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
            await extract_frame_at_timestamp(
                video_path, mp_selected.timestamp, final_path, for_final_screenshot=True,
            )
            logger.info(
                "Step %02d \u2713 multi-person-rescue @%.2fs conf=%.2f",
                step.step_number, mp_selected.timestamp, mp_confidence,
            )
            return step.model_copy(update={
                "selected_frame_label":      mp_selected.label,
                "selected_frame_reason":     mp_reason,
                "selected_frame_confidence": mp_confidence,
                "screenshot_timestamp":      mp_selected.timestamp,
                "screenshot_path":           str(final_path).replace("\\", "/"),
            })

    # Pass 2: Wide rescue window
    logger.warning("Step %02d — trying wide rescue pass.", step.step_number)
    rescue_timestamps = build_wide_rescue_timestamps(
        step=step, video_duration=video_duration,
        previous_step=previous_step, next_step=next_step,
    )
    (
        rescue_selected, rescue_confidence, rescue_reason,
        rescue_wrong_step, rescue_action_visible, rescue_flags,
        rescue_candidates,
    ) = await run_vision_pass(rescue_timestamps, "rescue", apply_quality_filter=False,
                               two_person=needs_two)

    if is_good_enough(rescue_selected, rescue_confidence, rescue_wrong_step, rescue_flags):
        screenshots_dir.mkdir(parents=True, exist_ok=True)
        final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
        await extract_frame_at_timestamp(
            video_path, rescue_selected.timestamp, final_path, for_final_screenshot=True,
        )
        logger.info(
            "Step %02d \u2713 rescue  frame %s @%.2fs conf=%.2f",
            step.step_number, rescue_selected.label,
            rescue_selected.timestamp, rescue_confidence,
        )
        return step.model_copy(update={
            "selected_frame_label":      rescue_selected.label,
            "selected_frame_reason":     rescue_reason,
            "selected_frame_confidence": rescue_confidence,
            "screenshot_timestamp":      rescue_selected.timestamp,
            "screenshot_path":           str(final_path).replace("\\", "/"),
        })

    rescue_api_failed = (
        rescue_selected is None
        and rescue_reason
        and ("vision error" in rescue_reason or "local fallback" in rescue_reason)
    )
    if _vision_skip_extended_on_api_error() and rescue_api_failed:
        pool = rescue_candidates or normal_candidates
        local = pick_local_fallback_candidate(pool, step, video_duration)
        if local is not None:
            logger.warning(
                "Step %02d: skipping extended vision after API failure; using local best frame.",
                step.step_number,
            )
            return await commit_frame(
                local,
                max(normal_confidence, rescue_confidence, VISION_SOFT_ACCEPT_CONFIDENCE),
                f"rescue failed ({rescue_reason}); local best-quality frame",
                "\u26a0 local",
            )

    logger.warning(
        "Step %02d \u2717 rescue  conf=%.2f — trying extended fallback.",
        step.step_number, rescue_confidence,
    )

    # Pass 3: Extended fallback
    extended_timestamps = build_extended_fallback_timestamps(
        step=step, video_duration=video_duration,
        previous_step=previous_step, next_step=next_step,
    )
    (
        ext_selected, ext_confidence, ext_reason,
        ext_wrong_step, ext_action_visible, ext_flags,
        ext_candidates,
    ) = await run_vision_pass(extended_timestamps, "extended_fallback",
                               apply_quality_filter=False, two_person=needs_two)

    if ext_selected and ext_confidence >= _extended_fallback_min_confidence():
        screenshots_dir.mkdir(parents=True, exist_ok=True)
        final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
        await extract_frame_at_timestamp(
            video_path, ext_selected.timestamp, final_path, for_final_screenshot=True,
        )
        logger.info(
            "Step %02d \u2713 extended_fallback @%.2fs conf=%.2f",
            step.step_number, ext_selected.timestamp, ext_confidence,
        )
        return step.model_copy(update={
            "selected_frame_label":      ext_selected.label,
            "selected_frame_reason":     ext_reason,
            "selected_frame_confidence": ext_confidence,
            "screenshot_timestamp":      ext_selected.timestamp,
            "screenshot_path":           str(final_path).replace("\\", "/"),
        })

    # Pass 4: Whole-video rescue — search the entire content range using the
    # visual_query, ignoring the transcript anchor. Fixes cases where Claude
    # anchored to outro voice-over / wrong B-roll.
    global_selected = None
    global_confidence = 0.0
    global_reason = ""
    if WHOLE_VIDEO_RESCUE_ENABLED:
        logger.warning(
            "Step %02d \u2717 extended  conf=%.2f \u2014 trying whole-video rescue.",
            step.step_number, ext_confidence,
        )
        global_timestamps = build_global_rescue_timestamps(
            step=step,
            video_duration=video_duration,
            scene_cuts=scene_cuts,
            previous_screenshot_timestamp=previous_screenshot_timestamp,
        )
        (
            global_selected, global_confidence, global_reason,
            _g_wrong, _g_action, _g_flags,
            global_candidates,
        ) = await run_vision_pass(
            global_timestamps, "global_rescue",
            apply_quality_filter=True, two_person=needs_two,
            max_vision_frames=VISION_RESCUE_SHORTLIST_FRAMES,
        )

        if (
            global_selected
            and global_confidence >= WHOLE_VIDEO_RESCUE_MIN_CONFIDENCE
            and global_confidence > max(normal_confidence, rescue_confidence, ext_confidence)
        ):
            screenshots_dir.mkdir(parents=True, exist_ok=True)
            final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
            await extract_frame_at_timestamp(
                video_path, global_selected.timestamp, final_path,
                for_final_screenshot=True,
            )
            logger.info(
                "Step %02d \u2713 global_rescue @%.2fs conf=%.2f (visual-query whole-video scan)",
                step.step_number, global_selected.timestamp, global_confidence,
            )
            return step.model_copy(update={
                "selected_frame_label":      global_selected.label,
                "selected_frame_reason":     f"global_rescue: {global_reason}",
                "selected_frame_confidence": global_confidence,
                "screenshot_timestamp":      global_selected.timestamp,
                "screenshot_path":           str(final_path).replace("\\", "/"),
            })

    # Pass 5: Forced best-quality frame (last resort)
    if ALWAYS_FILL_SCREENSHOTS:
        all_pool = ext_candidates or rescue_candidates or normal_candidates
        if all_pool:
            content_end = content_max_timestamp(video_duration)
            in_content = [c for c in all_pool if c.timestamp <= content_end]
            pool_base = in_content if in_content else all_pool
            non_banner = [c for c in pool_base if not c.has_prototype_banner]
            pool = non_banner if non_banner else pool_base
            anchor = step.timestamp
            fallback = max(
                pool,
                key=lambda c: (
                    c.quality_score
                    - 0.25 * abs(c.timestamp - anchor)
                    - (0.5 if c.timestamp > content_end - 3.0 else 0.0)
                ),
            )
            screenshots_dir.mkdir(parents=True, exist_ok=True)
            final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
            await extract_frame_at_timestamp(
                video_path, fallback.timestamp, final_path, for_final_screenshot=True,
            )
            fallback_conf   = max(
                normal_confidence, rescue_confidence, ext_confidence,
                global_confidence, 0.25,
            )
            fallback_reason = (
                f"FORCED BEST-AVAILABLE SCREENSHOT -- no confident match found. "
                f"Extended fallback: {ext_reason}"
            )
            logger.warning(
                "Step %02d \u26a0 forced  @%.2fs (quality=%.2f)",
                step.step_number, fallback.timestamp, fallback.quality_score,
            )
            return step.model_copy(update={
                "selected_frame_label":      fallback.label,
                "selected_frame_reason":     fallback_reason,
                "selected_frame_confidence": fallback_conf,
                "screenshot_timestamp":      fallback.timestamp,
                "screenshot_path":           str(final_path).replace("\\", "/"),
            })

    return step.model_copy(update={
        "selected_frame_label":      None,
        "selected_frame_reason":     "No frame candidates found after all passes.",
        "selected_frame_confidence": 0.0,
        "screenshot_timestamp":      None,
        "screenshot_path":           None,
    })


async def select_best_screenshots(
    video_path: Path,
    steps: list[InstallationStep],
    scene_cuts: list[SceneCut],
    video_duration: float,
    output_dir: Path,
) -> list[InstallationStep]:
    logger.info("[5/7] Gemini Vision frame selection for %d steps...", len(steps))

    candidates_root = output_dir / "frame_candidates"
    screenshots_dir = output_dir / "screenshots"

    steps = prepare_steps_for_vision(steps, video_duration)

    results: list[InstallationStep] = []
    for idx, step in enumerate(steps):
        previous_step = steps[idx - 1] if idx > 0 else None
        next_step     = steps[idx + 1] if idx + 1 < len(steps) else None
        prev_ss_ts: Optional[float] = results[-1].screenshot_timestamp if results else None

        result = await select_best_frame_for_step(
            video_path=video_path,
            step=step,
            scene_cuts=scene_cuts,
            video_duration=video_duration,
            candidates_root=candidates_root,
            screenshots_dir=screenshots_dir,
            previous_step=previous_step,
            next_step=next_step,
            previous_screenshot_timestamp=prev_ss_ts,
        )
        results.append(result)

    ok = sum(1 for s in results if s.screenshot_path and Path(s.screenshot_path).exists())
    logger.info("Screenshots: %d ok, %d missing.", ok, len(results) - ok)
    return results


# ─────────────────────────────────────────────
# CONSECUTIVE DUPLICATE DETECTION & FIX
# ─────────────────────────────────────────────

async def fix_consecutive_duplicate_screenshots(
    video_path: Path,
    steps: list[InstallationStep],
    scene_cuts: list[SceneCut],
    video_duration: float,
    output_dir: Path,
) -> list[InstallationStep]:
    fixed    = list(steps)
    candidates_root = output_dir / "frame_candidates"
    screenshots_dir = output_dir / "screenshots"
    safe_end = safe_max_timestamp(video_duration)

    for i in range(1, len(fixed)):
        prev_step = fixed[i - 1]
        curr_step = fixed[i]

        prev_ss = prev_step.screenshot_path
        curr_ss = curr_step.screenshot_path

        if not (
            prev_ss and curr_ss
            and Path(prev_ss).exists()
            and Path(curr_ss).exists()
        ):
            continue

        similarity = await asyncio.to_thread(
            image_histogram_similarity_sync, prev_ss, curr_ss,
        )

        if similarity < CONSECUTIVE_FRAME_SIMILARITY_MAX:
            continue

        logger.warning(
            "Steps %02d & %02d have near-identical screenshots (sim=%.3f >= %.2f). "
            "De-dup rescue for step %02d.",
            prev_step.step_number, curr_step.step_number,
            similarity, CONSECUTIVE_FRAME_SIMILARITY_MAX,
            curr_step.step_number,
        )

        curr_ts_start = curr_step.transcript_start or curr_step.timestamp
        curr_ts_end   = curr_step.transcript_end or (curr_step.timestamp + 8.0)
        next_step  = fixed[i + 1] if i + 1 < len(fixed) else None

        dedup_start = min(safe_end, curr_ts_start)
        dedup_end_raw = (next_step.transcript_start or next_step.timestamp) if next_step else safe_end
        dedup_end = min(safe_end, dedup_end_raw)

        if dedup_end - dedup_start < 1.0:
            continue

        n_frames = min(WIDE_RESCUE_FRAMES, max(4, int((dedup_end - dedup_start) / 0.8)))
        if n_frames <= 1:
            dedup_ts = [round((dedup_start + dedup_end) / 2.0, 3)]
        else:
            dedup_ts = [
                dedup_start + j * ((dedup_end - dedup_start) / (n_frames - 1))
                for j in range(n_frames)
            ]
        dedup_ts = unique_sorted_timestamps(dedup_ts, video_duration)

        pass_root  = candidates_root / "dedup"
        all_cands  = await extract_candidate_frames(video_path, curr_step, dedup_ts, pass_root)

        if not all_cands:
            continue

        dedup_pool = filter_by_quality(all_cands)
        if not dedup_pool:
            dedup_pool = all_cands
        dedup_pool = shortlist_candidates_for_vision(dedup_pool, curr_step, video_duration)

        prev_title = prev_step.title
        next_title = next_step.title if next_step else ""
        needs_two = is_multi_person_step(curr_step)
        prev_ss_ts = prev_step.screenshot_timestamp

        try:
            decision = await _ai_with_retry(
                select_frame_with_openai_vision_sync,
                curr_step, dedup_pool,
                prev_title, next_title,
                False, None,
                prev_ss_ts,
                needs_two,
                step_label=f"dedup-step{curr_step.step_number:02d}",
                max_retries=OPENAI_VISION_MAX_RETRIES,
                base_delay=OPENAI_VISION_BASE_DELAY,
                timeout_sec=OPENAI_VISION_TIMEOUT_SEC,
                models=OPENAI_VISION_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
            )
            selected, confidence, reason, is_wrong, _, flags = \
                interpret_vision_decision(decision, dedup_pool)

            if selected and confidence >= VISION_RETRY_MIN_CONFIDENCE:
                final_path = screenshots_dir / f"step_{curr_step.step_number:02d}.jpg"
                await extract_frame_at_timestamp(
                    video_path, selected.timestamp, final_path, for_final_screenshot=True,
                )

                fixed[i] = curr_step.model_copy(update={
                    "selected_frame_label":      selected.label,
                    "selected_frame_reason":     f"dedup-rescue: {reason}",
                    "selected_frame_confidence": confidence,
                    "screenshot_timestamp":      selected.timestamp,
                    "screenshot_path":           str(final_path).replace("\\", "/"),
                })
                logger.info(
                    "Step %02d \u2713 dedup-rescue @%.2fs conf=%.2f",
                    curr_step.step_number, selected.timestamp, confidence,
                )
            else:
                logger.warning(
                    "Step %02d dedup: low vision confidence (%.2f). "
                    "Using best-quality frame from candidate pool.",
                    curr_step.step_number, confidence,
                )
                non_banner = [c for c in all_cands if not c.has_prototype_banner]
                pool = non_banner if non_banner else all_cands
                best = max(pool, key=lambda c: c.quality_score)
                final_path = screenshots_dir / f"step_{curr_step.step_number:02d}.jpg"
                await extract_frame_at_timestamp(
                    video_path, best.timestamp, final_path, for_final_screenshot=True,
                )
                fixed[i] = curr_step.model_copy(update={
                    "selected_frame_label":      best.label,
                    "selected_frame_reason":     f"dedup-quality-fallback @{best.timestamp:.2f}s",
                    "selected_frame_confidence": max(confidence, 0.45),
                    "screenshot_timestamp":      best.timestamp,
                    "screenshot_path":           str(final_path).replace("\\", "/"),
                })
                logger.info(
                    "Step %02d \u2713 dedup-quality-fallback @%.2fs (quality=%.2f)",
                    curr_step.step_number, best.timestamp, best.quality_score,
                )

        except Exception as exc:
            logger.warning(
                "Step %02d dedup: all Gemini retries failed (%s). "
                "Falling back to best-quality frame from candidate pool.",
                curr_step.step_number, exc,
            )
            if all_cands:
                non_banner = [c for c in all_cands if not c.has_prototype_banner]
                pool = non_banner if non_banner else all_cands
                best = max(pool, key=lambda c: c.quality_score)
                final_path = screenshots_dir / f"step_{curr_step.step_number:02d}.jpg"
                try:
                    await extract_frame_at_timestamp(
                        video_path, best.timestamp, final_path, for_final_screenshot=True,
                    )
                    fixed[i] = curr_step.model_copy(update={
                        "selected_frame_label":      best.label,
                        "selected_frame_reason":     f"dedup-error-fallback @{best.timestamp:.2f}s: {exc}",
                        "selected_frame_confidence": 0.45,
                        "screenshot_timestamp":      best.timestamp,
                        "screenshot_path":           str(final_path).replace("\\", "/"),
                    })
                    logger.info(
                        "Step %02d \u2713 dedup-error-fallback @%.2fs (quality=%.2f)",
                        curr_step.step_number, best.timestamp, best.quality_score,
                    )
                except Exception as frame_exc:
                    logger.warning(
                        "Step %02d dedup-error-fallback frame extraction also failed: %s",
                        curr_step.step_number, frame_exc,
                    )

    return fixed


async def refine_low_confidence_screenshots(
    video_path: Path,
    steps: list[InstallationStep],
    video_duration: float,
    output_dir: Path,
) -> list[InstallationStep]:
    """Second-chance dense sampling for steps that finished below the refine threshold."""
    if not ENABLE_LOW_CONFIDENCE_REFINE:
        return steps

    candidates_root = output_dir / "frame_candidates"
    screenshots_dir = output_dir / "screenshots"
    refined = list(steps)
    target_indices = [
        i for i, s in enumerate(refined)
        if (s.selected_frame_confidence or 0.0) < LOW_CONFIDENCE_REFINE_THRESHOLD
        and s.screenshot_path
    ]
    if not target_indices:
        return refined

    logger.info(
        "[5b/7] Refining %d low-confidence screenshot(s) (threshold %.2f)...",
        len(target_indices),
        LOW_CONFIDENCE_REFINE_THRESHOLD,
    )

    for i in target_indices:
        step = refined[i]
        prev_step = refined[i - 1] if i > 0 else None
        next_step = refined[i + 1] if i + 1 < len(refined) else None
        prev_conf = step.selected_frame_confidence or 0.0

        timestamps = build_refine_timestamps(step, video_duration)
        pass_root = candidates_root / "refine"
        all_cands = await extract_candidate_frames(
            video_path, step, timestamps, pass_root / f"step_{step.step_number:02d}",
        )
        if not all_cands:
            continue

        pool = filter_by_quality(all_cands)
        if not pool:
            pool = all_cands
        pool = shortlist_candidates_for_vision(pool, step, video_duration)

        prev_title = prev_step.title if prev_step else ""
        next_title = next_step.title if next_step else ""
        needs_two = is_multi_person_step(step)
        prev_ss_ts = prev_step.screenshot_timestamp if prev_step else None

        try:
            decision = await _ai_with_retry(
                select_frame_with_openai_vision_sync,
                step, pool,
                prev_title, next_title,
                False, None,
                prev_ss_ts,
                needs_two,
                step_label=f"refine-step{step.step_number:02d}",
                max_retries=OPENAI_VISION_MAX_RETRIES,
                base_delay=OPENAI_VISION_BASE_DELAY,
                timeout_sec=OPENAI_VISION_TIMEOUT_SEC,
                models=OPENAI_VISION_MODELS if OPENAI_ENABLE_MODEL_FALLBACK else None,
            )
            selected, confidence, reason, is_wrong, _, _flags = \
                interpret_vision_decision(decision, pool)
        except Exception as exc:
            logger.debug("Step %02d refine pass failed: %s", step.step_number, exc)
            continue

        if not selected or is_wrong:
            continue
        if confidence < prev_conf + LOW_CONFIDENCE_REFINE_MIN_GAIN:
            logger.debug(
                "Step %02d refine: %.2f not better than %.2f (min gain %.2f).",
                step.step_number, confidence, prev_conf, LOW_CONFIDENCE_REFINE_MIN_GAIN,
            )
            continue
        if confidence < VISION_RETRY_MIN_CONFIDENCE:
            continue

        final_path = screenshots_dir / f"step_{step.step_number:02d}.jpg"
        await extract_frame_at_timestamp(
            video_path, selected.timestamp, final_path, for_final_screenshot=True,
        )
        refined[i] = step.model_copy(update={
            "selected_frame_label":      selected.label,
            "selected_frame_reason":     f"refine-pass: {reason}",
            "selected_frame_confidence": confidence,
            "screenshot_timestamp":      selected.timestamp,
            "screenshot_path":           str(final_path).replace("\\", "/"),
        })
        logger.info(
            "Step %02d \u2713 refine @%.2fs conf %.2f \u2192 %.2f",
            step.step_number, selected.timestamp, prev_conf, confidence,
        )

    return refined


# ─────────────────────────────────────────────
# QUALITY REPORT
# ─────────────────────────────────────────────

def build_quality_report(steps: list[InstallationStep]) -> dict[str, Any]:
    total           = len(steps)
    with_screenshots = sum(
        1 for s in steps if s.screenshot_path and Path(s.screenshot_path).exists()
    )
    low_confidence = [
        {
            "step_number": s.step_number, "title": s.title,
            "confidence":  s.selected_frame_confidence,
            "reason":      s.selected_frame_reason,
            "screenshot_timestamp": s.screenshot_timestamp,
        }
        for s in steps if (s.selected_frame_confidence or 0.0) < VISION_MIN_CONFIDENCE
    ]
    missing_screenshots = [
        {"step_number": s.step_number, "title": s.title, "reason": s.selected_frame_reason}
        for s in steps if not s.screenshot_path or not Path(s.screenshot_path).exists()
    ]
    fallback_steps = [
        {
            "step_number": s.step_number, "title": s.title,
            "confidence":  s.selected_frame_confidence, "reason": s.selected_frame_reason,
        }
        for s in steps
        if s.selected_frame_reason and (
            "fallback" in s.selected_frame_reason.lower()
            or "forced best-available" in s.selected_frame_reason.lower()
        )
    ]
    avg_conf = (
        sum((s.selected_frame_confidence or 0.0) for s in steps) / total
        if total else 0.0
    )
    return {
        "total_steps":        total,
        "screenshots_ok":     with_screenshots,
        "screenshots_missing": total - with_screenshots,
        "average_confidence": round(avg_conf, 3),
        "low_confidence_count": len(low_confidence),
        "fallback_count":     len(fallback_steps),
        "low_confidence_steps": low_confidence,
        "missing_screenshots":  missing_screenshots,
        "fallback_steps":       fallback_steps,
        "recommended_manual_review": (
            bool(missing_screenshots)
            if CLIENT_DEMO_MODE
            else bool(low_confidence or missing_screenshots or fallback_steps)
        ),
        "client_demo_mode": CLIENT_DEMO_MODE,
    }


# ─────────────────────────────────────────────
# DESCRIPTION CLEANER
# ─────────────────────────────────────────────

_DESCRIPTION_FIXES: list[tuple[str, str]] = [
    ("hard ware",           "hardware"),
    ("hardwire up top",     "hardware up top"),
    ("hard wire up top",    "hardware up top"),
    ("hardwire",            "hardware"),
    ("hard wire",           "hardware"),
    ("target top",          "tailgate top"),
    ("the target",          "the tailgate"),
    ("hard  top",           "hard top"),
    ("Targa  top",          "Targa top"),
    ("soft  top",           "soft top"),
    ("bed  rail",           "bed rail"),
    ("B  pillar",           "B-pillar"),
    ("b  pillar",           "B-pillar"),
    ("clamp down",          "clamp-down"),
    ("  ",                  " "),
]


def clean_description(text: str) -> str:
    for wrong, right in _DESCRIPTION_FIXES:
        text = text.replace(wrong, right)
    return text.strip()


# ─────────────────────────────────────────────
# LOGO  (v12 — LOGO_PATH only, no directory scan)
# ─────────────────────────────────────────────

def _strip_near_white_pixels(img, bg_color: tuple[int, int, int], threshold: int = 232):
    """Replace baked-in white JPEG margins with page background (dark template)."""
    from PIL import Image

    rgba = img.convert("RGBA")
    px = rgba.load()
    br, bg, bb = bg_color
    for y in range(rgba.height):
        for x in range(rgba.width):
            r, g, b, _a = px[x, y]
            if r >= threshold and g >= threshold and b >= threshold:
                px[x, y] = (br, bg, bb, 255)
    return rgba


def load_real_logo_bytes(
    logo_path: str,
    bg_color: tuple[int, int, int] = (30, 30, 30),
    padding: tuple[int, int] = (12, 8),
    strip_white_background: bool = False,
) -> Optional[bytes]:
    """
    Load a logo image (PNG/JPEG/AVIF/WEBP) and composite it on bg_color so
    white-on-transparent logos are visible on paper.
    Returns PNG bytes for python-docx, or None on failure.

    v12: Only called when LOGO_PATH is explicitly set.
    No directory scanning — avoids silently picking up stray files.
    """
    if not logo_path or not Path(logo_path).exists():
        return None
    try:
        from PIL import Image
        import io as _io
        with open(logo_path, "rb") as f:
            raw = f.read()
        img = Image.open(_io.BytesIO(raw))
        if strip_white_background:
            img = _strip_near_white_pixels(img, bg_color)
        pw, ph = padding
        canvas_w = img.width + pw * 2
        canvas_h = img.height + ph * 2
        canvas = Image.new("RGB", (canvas_w, canvas_h), bg_color)
        if img.mode == "RGBA":
            canvas.paste(img, (pw, ph), mask=img.split()[3])
        elif img.mode in ("LA", "PA"):
            rgba = img.convert("RGBA")
            canvas.paste(rgba, (pw, ph), mask=rgba.split()[3])
        else:
            canvas.paste(img.convert("RGB"), (pw, ph))
        buf = _io.BytesIO()
        canvas.save(buf, "PNG", optimize=True)
        return buf.getvalue()
    except Exception as exc:
        logger.warning("Real logo load failed (%s): %s", logo_path, exc)
        return None


def generate_logo_png_sync(output_path: Path, width: int = 400, height: int = 80) -> bool:
    """Draw a simple navy placeholder logo when no real logo is supplied."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)

        dark_navy  = (26, 26, 46)
        brand_blue = (26, 92, 150)
        white      = (255, 255, 255)

        bar_h = height - 4
        bar_y = 2
        draw.rounded_rectangle([0, bar_y, width, bar_y + bar_h], radius=8, fill=dark_navy)

        font_size_large = max(24, height - 24)
        font_size_small = max(14, height - 40)

        try:
            import platform
            if platform.system() == "Windows":
                font_bold = ImageFont.truetype("arialbd.ttf", font_size_large)
                font_reg  = ImageFont.truetype("arial.ttf",   font_size_small)
            elif platform.system() == "Darwin":
                font_bold = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size_large)
                font_reg  = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size_small)
            else:
                font_bold = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size_large)
                font_reg  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size_small)
        except Exception:
            font_bold = ImageFont.load_default()
            font_reg  = ImageFont.load_default()

        stripe_w = 6
        draw.rectangle([stripe_w + 4, bar_y + 6, stripe_w + 10, bar_y + bar_h - 6],
                        fill=brand_blue)

        text_x = stripe_w + 18
        text_y_main = bar_y + (bar_h - font_size_large) // 2
        text_y_sub  = bar_y + (bar_h - font_size_small) // 2 + 2

        draw.text((text_x, text_y_main), "TURN", font=font_bold, fill=white)

        try:
            bbox = draw.textbbox((text_x, text_y_main), "TURN", font=font_bold)
            turn_width = bbox[2] - bbox[0]
        except Exception:
            turn_width = font_size_large * 4

        draw.text((text_x + turn_width + 8, text_y_sub), "OFFROAD", font=font_reg,
                  fill=(180, 210, 240))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        img.save(str(output_path), "PNG")
        return True

    except Exception as exc:
        logger.debug("Logo generation failed: %s", exc)
        return False


# ─────────────────────────────────────────────
# WORD DOCUMENT helpers (screenshots; v14 Tap layout lives in build_word_doc_sync)
# ─────────────────────────────────────────────


def prepare_screenshot_for_doc_sync(path: str | Path) -> bytes:
    """Re-encode step JPG with optional uniform black border for Word embedding."""
    from io import BytesIO

    p = Path(path)
    raw = p.read_bytes()
    if SCREENSHOT_BORDER_PX <= 0:
        return raw
    try:
        from PIL import Image

        im = Image.open(BytesIO(raw)).convert("RGB")
        b = max(0, SCREENSHOT_BORDER_PX)
        canvas = Image.new("RGB", (im.width + 2 * b, im.height + 2 * b), (0, 0, 0))
        canvas.paste(im, (b, b))
        out = BytesIO()
        canvas.save(out, format="JPEG", quality=94, optimize=True)
        return out.getvalue()
    except Exception as exc:
        logger.warning("Screenshot border failed for %s (%s) — using original file.", p, exc)
        return raw


def _header_title_lines(video_title: str, brand: str) -> tuple[str, str, str]:
    """Match installation_guide_template_style.docx title block."""
    title = (video_title or "Installation Guide").strip().rstrip(".")
    m = re.search(r"^(.+?)\s+installation\s+guide\s*\.?$", title, re.IGNORECASE)
    product = (m.group(1) if m else title).strip()
    brand_name = (brand or "Turn Offroad").strip()
    product = re.sub(
        rf"^{re.escape(brand_name)}'?s?\s+",
        "",
        product,
        flags=re.IGNORECASE,
    )
    product_line = product.upper()
    return product_line, "INSTALLATION GUIDE", f"{brand_name}  \u2022  Official Installation Guide"


def _qr_code_png_bytes(video_url: str) -> Optional[bytes]:
    if not video_url:
        return None
    try:
        import io as _io
        import qrcode

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=4,
            border=2,
        )
        qr.add_data(video_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception as exc:
        logger.warning("QR generation failed: %s", exc)
        return None


def publish_deliverable_doc(
    source_docx: Path,
    *,
    video_url: str = "",
    video_title: str = "",
) -> Optional[Path]:
    """
    Copy the finished Word file to DELIVERABLE_OUTPUT_DIR for the server / web app.
    Writes deliverable_manifest.json alongside the doc.
    """
    if not DELIVERABLE_OUTPUT_DIR:
        return None
    dest_root = Path(DELIVERABLE_OUTPUT_DIR)
    dest_root.mkdir(parents=True, exist_ok=True)

    video_id = "guide"
    if video_url:
        vid_m = re.search(r"(?:v=|youtu\.be/)([A-Za-z0-9_-]{6,})", video_url)
        if vid_m:
            video_id = vid_m.group(1)

    slug = re.sub(r"[^\w\-]+", "_", (video_title or "installation_guide")).strip("_")[:72]
    if not slug:
        slug = "installation_guide"
    if DELIVERABLE_FILENAME_STYLE == "video_id":
        dest_path = dest_root / f"{video_id}.docx"
    else:
        dest_path = dest_root / f"{slug}_{video_id}.docx"
    shutil.copy2(source_docx, dest_path)

    manifest = {
        "docx_path": str(dest_path.resolve()),
        "video_url": video_url,
        "video_title": video_title,
        "generated_at": datetime.datetime.now(datetime.timezone.utc).isoformat(),
    }
    manifest_path = dest_root / "deliverable_manifest.json"
    with open(manifest_path, "w", encoding="utf-8") as mf:
        json.dump(manifest, mf, indent=2)

    logger.info("Deliverable published: %s", dest_path.resolve())
    return dest_path


def _doc_theme_palette(*, dark: bool) -> dict[str, str]:
    """Colors for build_word_doc_sync (dark = client Turn Offroad template)."""
    if dark:
        return {
            "page": "1E1E1E",
            "panel": "2B2B2B",
            "panel_alt": "333333",
            "header_bar": "1A1A1A",
            "title": "FFFFFF",
            "body": "E8E8E8",
            "muted": "B0B0B0",
            "accent": DOC_ACCENT_COLOR or "E8634B",
            "warn_fill": "4A3F1A",
            "white": "FFFFFF",
            "logo_bg": (30, 30, 30),
            "border": "444444",
        }
    return {
        "page": "FFFFFF",
        "panel": "FFFFFF",
        "panel_alt": "F5F5F5",
        "header_bar": "1A1A1A",
        "title": "1A1A1A",
        "body": "333333",
        "muted": "666666",
        "accent": DOC_ACCENT_COLOR or "E8634B",
        "warn_fill": "FFF3CD",
        "white": "FFFFFF",
        "logo_bg": (255, 255, 255),
        "border": "BFBFBF",
    }


def _set_document_page_background(doc, fill_hex: str) -> None:
    """Full-page background (client dark template)."""
    try:
        from docx.oxml import OxmlElement as _Ox
        from docx.oxml.ns import qn as _qn

        background = _Ox("w:background")
        background.set(_qn("w:color"), fill_hex)
        doc.element.insert(0, background)
    except Exception as exc:
        logger.debug("Page background not set: %s", exc)


def _apply_dark_document_defaults(doc, title_hex: str, body_hex: str) -> None:
    """Default paragraph style so gaps between tables are not black-on-white."""
    try:
        from docx.shared import RGBColor

        normal = doc.styles["Normal"]
        normal.font.color.rgb = RGBColor(
            int(body_hex[0:2], 16), int(body_hex[2:4], 16), int(body_hex[4:6], 16),
        )
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            if style_name in doc.styles:
                doc.styles[style_name].font.color.rgb = RGBColor(
                    int(title_hex[0:2], 16), int(title_hex[2:4], 16), int(title_hex[4:6], 16),
                )
    except Exception as exc:
        logger.debug("Document default styles not set: %s", exc)


def _shade_paragraph(paragraph, fill_hex: str) -> None:
    """Paragraph background for text outside tables (dark template)."""
    from docx.oxml import OxmlElement as _Ox
    from docx.oxml.ns import qn as _qn

    p_pr = paragraph._element.get_or_add_pPr()
    shd = _Ox("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), fill_hex)
    p_pr.append(shd)


def _paragraph_bottom_border(paragraph, color_hex: str, *, sz: str = "12") -> None:
    """Single bottom rule on one paragraph (step titles, section headings)."""
    from docx.oxml import OxmlElement as _Ox
    from docx.oxml.ns import qn as _qn

    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = _Ox("w:pBdr")
    bottom = _Ox("w:bottom")
    bottom.set(_qn("w:val"), "single")
    bottom.set(_qn("w:sz"), sz)
    bottom.set(_qn("w:space"), "1")
    bottom.set(_qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_accent_rule(doc, color_hex: str, *, space_before_pt: int = 4, space_after_pt: int = 12) -> None:
    """Thin coral/red line under the header block."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    _paragraph_bottom_border(p, color_hex, sz="14")


def _logo_picture_size_inches(logo_path: Path) -> tuple[float, float]:
    """Fit wide Turn Offroad logo without cropping in the header cell."""
    max_w = DOC_HEADER_LOGO_MAX_WIDTH_INCHES
    max_h = DOC_HEADER_LOGO_MAX_HEIGHT_INCHES
    try:
        from PIL import Image

        with Image.open(logo_path) as im:
            px_w, px_h = im.size
        if px_h <= 0:
            return max_w, max_h
        aspect = px_w / px_h
        width = max_w
        height = width / aspect
        if height > max_h:
            height = max_h
            width = height * aspect
        return round(width, 3), round(height, 3)
    except Exception:
        return max_w, max_h


def _insert_logo_picture(paragraph, logo_path: Path) -> bool:
    try:
        from docx.shared import Inches as _Inches

        w_in, h_in = _logo_picture_size_inches(logo_path)
        paragraph.add_run().add_picture(
            str(logo_path), width=_Inches(w_in), height=_Inches(h_in),
        )
        return True
    except Exception as exc:
        logger.warning("Logo picture insert failed: %s", exc)
        return False


def build_word_doc_sync(
    steps: list[InstallationStep],
    output_path: Path,
    video_title: str,
    quality_report: dict[str, Any],
    video_url: str = "",
    tools_info: dict = None,
    contact_info: dict = None,
    important_note_override: str = "",
) -> Path:
    """
    Build Word doc to match installation_guide_template_style.docx:
    logo + title, video/QR, overview, tools|hardware, step|screenshot rows,
    important note, contact block, footer bar.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    if tools_info is None:
        tools_info = {}
    if contact_info is None:
        contact_info = {}

    # Theme is controlled solely by DOC_DARK_THEME. CLIENT_DEMO_MODE only
    # affects vision/QA strictness, not the document background colour.
    dark_theme = DOC_DARK_THEME
    pal = _doc_theme_palette(dark=dark_theme)
    _PAGE_BG = pal["page"]
    _PANEL_BG = pal["panel"]
    _PANEL_ALT = pal["panel_alt"]
    _DARK_HDR = pal["header_bar"]
    _TITLE = pal["title"]
    _BODY = pal["body"]
    _MUTED = pal["muted"]
    _ACCENT = pal["accent"]
    _WARN_FILL = pal["warn_fill"]
    _WHITE = pal["white"]
    _BORDER = pal.get("border", "BFBFBF")
    brand = contact_info.get("brand", "Turn Offroad")
    logger.info(
        "Word layout theme: %s (DOC_DARK_THEME=%s, CLIENT_DEMO_MODE=%s)",
        "dark" if dark_theme else "light", DOC_DARK_THEME, CLIENT_DEMO_MODE,
    )

    important_default = (
        "\u26a0  IMPORTANT: Recheck all bolts and jam nuts after 48 hours of use. "
        "Do not overtighten the top mounting bolts."
    )
    if important_note_override and important_note_override.strip():
        important_note = important_note_override.strip()
    else:
        important_note = os.getenv("TEMPLATE_IMPORTANT_NOTE", important_default).strip() or important_default

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(DOC_LEFT_MARGIN_INCHES)
        section.right_margin = Inches(DOC_RIGHT_MARGIN_INCHES)
    if dark_theme:
        _set_document_page_background(doc, _PAGE_BG)
        _apply_dark_document_defaults(doc, _TITLE, _BODY)

    def _rgbx(hex6: str):
        from docx.shared import RGBColor
        return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))

    def cell_shad(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def cell_bdr_all(cell, color: str = "CCCCCC", sz: str = "4") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        bdr = OxmlElement("w:tcBorders")
        for side in ("w:top", "w:left", "w:bottom", "w:right"):
            b = OxmlElement(side)
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            bdr.append(b)
        tc_pr.append(bdr)

    def rm_tbl_borders(tbl) -> None:
        tbl_elem = tbl._tbl
        tbl_pr = tbl_elem.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl_elem.insert(0, tbl_pr)
        tbl_bdr = OxmlElement("w:tblBorders")
        for side in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
            b = OxmlElement(side)
            b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "auto")
            tbl_bdr.append(b)
        tbl_pr.append(tbl_bdr)

    def set_cell_valign_top(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        valign = OxmlElement("w:vAlign")
        valign.set(qn("w:val"), "top")
        tc_pr.append(valign)

    def _cell_para(cell, text: str, *, bold: bool = False, size: int = 10, color: str = _BODY) -> None:
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = _rgbx(color)

    def _fill_tool_column(cell, heading: str, items: list[str]) -> None:
        cell.text = ""
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_after = Pt(6)
        hr = hp.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = _rgbx(_TITLE)
        for item in items:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_after = Pt(3)
            br = bp.add_run(f"\u2610  {item}")
            br.font.size = Pt(9)
            br.font.color.rgb = _rgbx(_BODY)

    # ── Logo (LOGO_PATH or placeholder) ───────────────────────────────────
    logo_bg = pal["logo_bg"]
    real_logo_bytes: Optional[bytes] = None
    if LOGO_PATH:
        real_logo_bytes = load_real_logo_bytes(
            LOGO_PATH,
            bg_color=logo_bg,
            strip_white_background=dark_theme,
            padding=(4, 4) if dark_theme else (12, 8),
        )
        if real_logo_bytes:
            logger.info("Using real logo from LOGO_PATH: %s", LOGO_PATH)
        else:
            logger.warning("LOGO_PATH set but logo failed to load: %s", LOGO_PATH)

    logo_path = output_path.parent / "_logo_temp_v16.png"
    logo_ok = False
    if real_logo_bytes:
        logo_ok = True
        with open(logo_path, "wb") as _lf:
            _lf.write(real_logo_bytes)
    else:
        logo_ok = generate_logo_png_sync(logo_path, width=360, height=72)
        if logo_ok:
            logger.info("Using generated placeholder logo.")

    raw_tools_list: list[str] = tools_info.get("tools", [])
    raw_hardware_list: list[str] = tools_info.get("hardware", [])
    if raw_hardware_list:
        display_tools = raw_tools_list
        display_hardware = raw_hardware_list
    else:
        display_tools, display_hardware = _auto_split_tools(raw_tools_list)

    product_line, guide_line, tagline = _header_title_lines(video_title, brand)
    step_img_w = DOC_STEP_IMAGE_WIDTH_INCHES
    step_txt_w = DOC_STEP_TEXT_WIDTH_INCHES
    qr_png = _qr_code_png_bytes(video_url)
    content_w = DOC_CONTENT_WIDTH_INCHES
    logo_col_w = DOC_HEADER_LOGO_MAX_WIDTH_INCHES + 0.35
    title_col_w = max(2.5, content_w - logo_col_w)

    # ── T0: Logo (left) | title block (right, left-aligned like client template) ─
    hdr = doc.add_table(rows=1, cols=2)
    rm_tbl_borders(hdr)
    hdr.autofit = False
    h_logo, h_title = hdr.rows[0].cells[0], hdr.rows[0].cells[1]
    h_logo.width = Inches(logo_col_w)
    h_title.width = Inches(title_col_w)
    set_cell_valign_top(h_logo)
    set_cell_valign_top(h_title)
    cell_shad(h_logo, _PAGE_BG)
    cell_shad(h_title, _PAGE_BG)

    hp_logo = h_logo.paragraphs[0]
    hp_logo.paragraph_format.space_before = Pt(4)
    hp_logo.paragraph_format.space_after = Pt(0)
    if logo_ok and logo_path.exists():
        if not _insert_logo_picture(hp_logo, logo_path):
            _cell_para(h_logo, brand.upper(), bold=True, size=14, color=_TITLE)
    else:
        _cell_para(h_logo, brand.upper(), bold=True, size=14, color=_TITLE)

    ht0 = h_title.paragraphs[0]
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ht0.paragraph_format.space_before = Pt(6)
    ht0.paragraph_format.space_after = Pt(4)
    r0 = ht0.add_run(product_line)
    r0.bold = True
    r0.font.size = Pt(18)
    r0.font.color.rgb = _rgbx(_TITLE)
    ht1 = h_title.add_paragraph()
    ht1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ht1.paragraph_format.space_after = Pt(4)
    r1 = ht1.add_run(guide_line)
    r1.bold = True
    r1.font.size = Pt(15)
    r1.font.color.rgb = _rgbx(_ACCENT)
    ht2 = h_title.add_paragraph()
    ht2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ht2.paragraph_format.space_after = Pt(0)
    r2 = ht2.add_run(tagline)
    r2.font.size = Pt(9)
    r2.font.color.rgb = _rgbx(_BODY if dark_theme else _MUTED)

    _add_accent_rule(doc, _ACCENT, space_before_pt=8, space_after_pt=16)

    # ── T1: QR (left) | YouTube copy (right, dark panel) ─────────────────
    if video_url:
        vid_tbl = doc.add_table(rows=1, cols=2)
        rm_tbl_borders(vid_tbl)
        vid_tbl.autofit = False
        v_qr, v_txt = vid_tbl.rows[0].cells[0], vid_tbl.rows[0].cells[1]
        qr_col_w = 1.65
        v_qr.width = Inches(qr_col_w)
        v_txt.width = Inches(content_w - qr_col_w)
        cell_shad(v_qr, _PAGE_BG)
        cell_shad(v_txt, _PANEL_BG)
        cell_bdr_all(v_txt, _BORDER, sz="4")
        set_cell_valign_top(v_qr)
        set_cell_valign_top(v_txt)
        if qr_png:
            from io import BytesIO
            qp = v_qr.paragraphs[0]
            qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qp.paragraph_format.space_before = Pt(12)
            qp.paragraph_format.space_after = Pt(12)
            qp.add_run().add_picture(BytesIO(qr_png), width=Inches(1.35))
        vp0 = v_txt.paragraphs[0]
        vp0.paragraph_format.space_before = Pt(14)
        vp0.paragraph_format.space_after = Pt(6)
        vh = vp0.add_run("Watch the Official Installation Video")
        vh.bold = True
        vh.font.size = Pt(11)
        vh.font.color.rgb = _rgbx(_TITLE)
        vp1 = v_txt.add_paragraph()
        vp1.paragraph_format.space_after = Pt(6)
        v1 = vp1.add_run("Scan the QR code or visit:")
        v1.font.size = Pt(10)
        v1.font.color.rgb = _rgbx(_BODY)
        vp2 = v_txt.add_paragraph()
        vp2.paragraph_format.space_after = Pt(12)
        v2 = vp2.add_run(video_url)
        v2.font.size = Pt(9)
        v2.font.color.rgb = _rgbx(_ACCENT)

    # ── Overview heading + T2 summary table ───────────────────────────────
    ov_p = doc.add_paragraph()
    if dark_theme:
        _shade_paragraph(ov_p, _PAGE_BG)
    ov_p.paragraph_format.space_before = Pt(16)
    ov_p.paragraph_format.space_after = Pt(6)
    ovr = ov_p.add_run("INSTALLATION OVERVIEW")
    ovr.bold = True
    ovr.font.size = Pt(11)
    ovr.font.color.rgb = _rgbx(_TITLE)

    if tools_info:
        summary_tbl = doc.add_table(rows=2, cols=3)
        rm_tbl_borders(summary_tbl)
        summary_tbl.autofit = False
        hdr_labels = ["Difficulty", "Estimated Time", "People Needed"]
        hdr_values = [
            str(tools_info.get("difficulty", "\u2014")),
            f"{tools_info.get('estimated_time_minutes', '?')} min",
            f"{tools_info.get('people_required', '?')} person(s)",
        ]
        col_w = round(content_w / 3.0, 2)
        for col in range(3):
            hc = summary_tbl.rows[0].cells[col]
            hc.width = Inches(col_w)
            cell_shad(hc, _PANEL_ALT if dark_theme else _DARK_HDR)
            cell_bdr_all(hc, _BORDER, sz="4")
            hp = hc.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.paragraph_format.space_before = Pt(4)
            hp.paragraph_format.space_after = Pt(4)
            hr = hp.add_run(hdr_labels[col])
            hr.bold = True
            hr.font.size = Pt(9)
            hr.font.color.rgb = _rgbx(_TITLE if dark_theme else _WHITE)
            vc = summary_tbl.rows[1].cells[col]
            cell_shad(vc, _PANEL_BG)
            cell_bdr_all(vc, _BORDER, sz="4")
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vp.paragraph_format.space_before = Pt(6)
            vp.paragraph_format.space_after = Pt(6)
            vr = vp.add_run(hdr_values[col])
            vr.bold = True
            vr.font.size = Pt(11)
            vr.font.color.rgb = _rgbx(_TITLE if dark_theme else _DARK_HDR)

    # ── T3: Tools | Hardware (two columns) ────────────────────────────────
    if display_tools or display_hardware:
        tools_tbl = doc.add_table(rows=1, cols=2)
        rm_tbl_borders(tools_tbl)
        tools_tbl.autofit = False
        t_left, t_right = tools_tbl.rows[0].cells[0], tools_tbl.rows[0].cells[1]
        half = round(content_w / 2.0, 2)
        t_left.width = Inches(half)
        t_right.width = Inches(half)
        set_cell_valign_top(t_left)
        set_cell_valign_top(t_right)
        cell_shad(t_left, _PAGE_BG)
        cell_shad(t_right, _PAGE_BG)
        cell_bdr_all(t_left, _BORDER, sz="4")
        cell_bdr_all(t_right, _BORDER, sz="4")
        _fill_tool_column(t_left, "TOOLS REQUIRED", display_tools)
        _fill_tool_column(t_right, "HARDWARE INCLUDED IN KIT", display_hardware)

    # ── Steps: text (left) | screenshot (right) ───────────────────────────
    figure_num = 0
    for step in steps:
        description = clean_description(step.description)
        title = clean_description(step.title)

        row_tbl = doc.add_table(rows=1, cols=2)
        rm_tbl_borders(row_tbl)
        row_tbl.autofit = False
        txt_cell, img_cell = row_tbl.rows[0].cells[0], row_tbl.rows[0].cells[1]
        txt_cell.width = Inches(step_txt_w)
        img_cell.width = Inches(step_img_w)
        set_cell_valign_top(txt_cell)
        set_cell_valign_top(img_cell)

        cell_shad(txt_cell, _PAGE_BG)
        cell_shad(img_cell, _PAGE_BG)

        txt_cell.paragraphs[0].text = ""
        np = txt_cell.paragraphs[0]
        np.paragraph_format.space_before = Pt(10)
        np.paragraph_format.space_after = Pt(0)
        nr = np.add_run(f"{step.step_number}.")
        nr.bold = True
        nr.font.size = Pt(30)
        nr.font.color.rgb = _rgbx(_ACCENT)
        tp1 = txt_cell.add_paragraph()
        tp1.paragraph_format.space_before = Pt(0)
        tp1.paragraph_format.space_after = Pt(6)
        tr = tp1.add_run(title)
        tr.bold = True
        tr.font.size = Pt(11)
        tr.font.color.rgb = _rgbx(_TITLE)
        _paragraph_bottom_border(tp1, _ACCENT, sz="10")
        tp2 = txt_cell.add_paragraph()
        tp2.paragraph_format.space_after = Pt(10)
        bd = tp2.add_run(description)
        bd.font.size = Pt(10)
        bd.font.color.rgb = _rgbx(_BODY)

        if DOC_SHOW_STEP_TIMESTAMPS:
            ts_sec = step_video_timestamp_seconds(step)
            ts_label = format_video_timestamp(ts_sec)
            ts_p = txt_cell.add_paragraph()
            ts_p.paragraph_format.space_before = Pt(6)
            ts_p.paragraph_format.space_after = Pt(0)
            ts_prefix = ts_p.add_run("Video timestamp: ")
            ts_prefix.font.size = Pt(9)
            ts_prefix.font.color.rgb = _rgbx(_MUTED)
            link_url = youtube_watch_url_at(video_url, ts_sec)
            if link_url:
                add_docx_hyperlink(
                    ts_p,
                    ts_label,
                    link_url,
                    color_hex=_ACCENT,
                    font_size_pt=9,
                )
            else:
                ts_val = ts_p.add_run(ts_label)
                ts_val.bold = True
                ts_val.font.size = Pt(9)
                ts_val.font.color.rgb = _rgbx(_ACCENT)

        ss = step.screenshot_path
        if ss and Path(ss).exists():
            try:
                from io import BytesIO

                fig = prepare_screenshot_for_doc_sync(ss)
                ip = img_cell.paragraphs[0]
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(4)
                ip.add_run().add_picture(BytesIO(fig), width=Inches(step_img_w - 0.1))
                figure_num += 1
                if DOC_SHOW_STEP_TIMESTAMPS:
                    ts_sec = step_video_timestamp_seconds(step)
                    ts_label = format_video_timestamp(ts_sec)
                    cap_p = img_cell.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_before = Pt(4)
                    cap_p.paragraph_format.space_after = Pt(8)
                    cap_prefix = cap_p.add_run("Screenshot at ")
                    cap_prefix.font.size = Pt(8)
                    cap_prefix.font.color.rgb = _rgbx(_MUTED)
                    link_url = youtube_watch_url_at(video_url, ts_sec)
                    if link_url:
                        add_docx_hyperlink(
                            cap_p,
                            ts_label,
                            link_url,
                            color_hex=_ACCENT,
                            font_size_pt=8,
                        )
                        cap_hint = cap_p.add_run(" — click to open in YouTube")
                        cap_hint.font.size = Pt(8)
                        cap_hint.font.color.rgb = _rgbx(_MUTED)
                    else:
                        cap_val = cap_p.add_run(ts_label)
                        cap_val.bold = True
                        cap_val.font.size = Pt(8)
                        cap_val.font.color.rgb = _rgbx(_ACCENT)
            except Exception as exc:
                logger.error("Step %02d image insert failed: %s", step.step_number, exc)
                _cell_para(img_cell, f"[Photo unavailable: {exc}]", color="CC0000")
        else:
            _cell_para(img_cell, "[Screenshot not available]", color="CC0000")

    # ── Important note (yellow box, after steps) ──────────────────────────
    imp_tbl = doc.add_table(rows=1, cols=1)
    rm_tbl_borders(imp_tbl)
    ic = imp_tbl.rows[0].cells[0]
    cell_shad(ic, _WARN_FILL)
    ip = ic.paragraphs[0]
    ip.paragraph_format.space_before = Pt(8)
    ip.paragraph_format.space_after = Pt(8)
    ir = ip.add_run(important_note)
    ir.font.size = Pt(9)
    ir.font.color.rgb = _rgbx(_BODY)
    ir.bold = True

    # ── Support / contact (template closing block) ────────────────────────
    help_p = doc.add_paragraph()
    if dark_theme:
        _shade_paragraph(help_p, _PAGE_BG)
    help_p.paragraph_format.space_before = Pt(14)
    hr = help_p.add_run("Need Help With Your Installation?")
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = _rgbx(_TITLE)
    _add_accent_rule(doc, _ACCENT, space_before_pt=2, space_after_pt=8)

    help2 = doc.add_paragraph()
    if dark_theme:
        _shade_paragraph(help2, _PAGE_BG)
    help2.paragraph_format.space_after = Pt(8)
    h2 = help2.add_run(
        "Our team is here to help. Reach out through any of the channels below "
        "and we'll get you sorted."
    )
    h2.font.size = Pt(10)
    h2.font.color.rgb = _rgbx(_BODY)

    website = (contact_info.get("website") or "https://turnoffroad.com").strip()
    email = (contact_info.get("email") or "help@turnoffroad.com").strip()
    phone = (contact_info.get("phone") or "951-505-6437").strip()
    contact_rows = [
        ("Website", website),
        ("Email", email),
        ("Phone", phone),
    ]
    ct_tbl = doc.add_table(rows=len(contact_rows), cols=2)
    rm_tbl_borders(ct_tbl)
    ct_tbl.autofit = False
    _zebra = _PANEL_ALT if not dark_theme else _PANEL_BG
    for r_idx, (label, value) in enumerate(contact_rows):
        lc, vc = ct_tbl.rows[r_idx].cells[0], ct_tbl.rows[r_idx].cells[1]
        lc.width = Inches(1.8)
        vc.width = Inches(4.7)
        fill = _zebra if r_idx % 2 == 0 else _PAGE_BG
        cell_shad(lc, fill)
        cell_shad(vc, fill)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = _rgbx(_TITLE)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = _rgbx(_ACCENT)

    # ── Footer bar (white banner + logo, like client template) ───────────
    brand_upper = brand.upper()
    footer_line = f"{website}  \u2502  {email}  \u2502  {phone}"
    ft_tbl = doc.add_table(rows=1, cols=2)
    rm_tbl_borders(ft_tbl)
    ft_tbl.autofit = False
    flc, frc = ft_tbl.rows[0].cells[0], ft_tbl.rows[0].cells[1]
    flc.width = Inches(2.2)
    frc.width = Inches(4.3)
    _FOOTER_BG = _WHITE
    _FOOTER_TITLE = "1A1A1A"
    _FOOTER_DETAIL = "555555"
    cell_shad(flc, _FOOTER_BG)
    cell_shad(frc, _FOOTER_BG)
    set_cell_valign_top(flc)
    set_cell_valign_top(frc)
    footer_logo = Path(LOGO_PATH) if LOGO_PATH and Path(LOGO_PATH).exists() else logo_path
    if footer_logo.exists():
        flp = flc.paragraphs[0]
        flp.paragraph_format.space_before = Pt(6)
        flp.paragraph_format.space_after = Pt(6)
        try:
            w_in, h_in = _logo_picture_size_inches(footer_logo)
            flp.add_run().add_picture(
                str(footer_logo),
                width=Inches(max(1.0, w_in * 0.5)),
                height=Inches(max(0.22, h_in * 0.5)),
            )
        except Exception:
            pass
    frp = frc.paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frp.paragraph_format.space_before = Pt(8)
    frp.paragraph_format.space_after = Pt(8)
    frr_title = frp.add_run(f"{brand_upper} INSTALLATION SUPPORT\n")
    frr_title.bold = True
    frr_title.font.size = Pt(9)
    frr_title.font.color.rgb = _rgbx(_FOOTER_TITLE)
    frr_detail = frp.add_run(footer_line)
    frr_detail.font.size = Pt(8)
    frr_detail.font.color.rgb = _rgbx(_FOOTER_DETAIL)

    thanks = doc.add_paragraph()
    if dark_theme:
        _shade_paragraph(thanks, _PAGE_BG)
    thanks.paragraph_format.space_before = Pt(10)
    tk = thanks.add_run(f"Thank you for choosing {brand}. Enjoy the trail!")
    tk.bold = True
    tk.font.size = Pt(11)
    tk.font.color.rgb = _rgbx(_TITLE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    try:
        if logo_path.exists():
            logo_path.unlink()
    except Exception:
        pass

    missing = len(steps) - figure_num
    logger.info(
        "Saved (template style): %s  (%d screenshots, %d missing)",
        output_path.name, figure_num, missing,
    )
    published = publish_deliverable_doc(
        output_path, video_url=video_url, video_title=video_title,
    )
    return published or output_path


async def build_word_doc(
    steps: list[InstallationStep],
    output_path: Path,
    video_title: str,
    quality_report: dict[str, Any],
    video_url: str = "",
    tools_info: dict = None,
    contact_info: dict = None,
    important_note_override: str = "",
) -> Path:
    logger.info("[7/7] Building Word document -> %s", output_path)
    return await asyncio.to_thread(
        build_word_doc_sync,
        steps, output_path, video_title, quality_report,
        video_url, tools_info or {}, contact_info or {},
        important_note_override,
    )


# ─────────────────────────────────────────────
# DOC-ONLY MODE
# ─────────────────────────────────────────────

async def doc_only_mode(args: argparse.Namespace) -> None:
    output_dir = Path(args.output_dir)
    setup_logging(output_dir, args.verbose)
    logger.info("=== DOC-ONLY mode: rebuilding document from %s ===", output_dir)

    steps_path = output_dir / "steps.json"
    if not steps_path.exists():
        logger.error("steps.json not found in %s. Run the full pipeline first.", output_dir)
        sys.exit(1)

    with open(steps_path, encoding="utf-8") as f:
        raw_steps = json.load(f)
    steps = [InstallationStep.model_validate(s) for s in raw_steps]

    qr_path = output_dir / "quality_report.json"
    quality_report: dict[str, Any] = {}
    if qr_path.exists():
        with open(qr_path, encoding="utf-8") as f:
            quality_report = json.load(f)
    else:
        quality_report = build_quality_report(steps)

    tools_info: dict = {}
    tools_path = output_dir / "tools_info.json"
    if tools_path.exists():
        with open(tools_path, encoding="utf-8") as f:
            tools_info = json.load(f)

    if getattr(args, "title", None):
        video_title = args.title
    else:
        video_title = read_video_title_from_info_json(output_dir) or "Installation Guide"

    doc_path = output_dir / "installation_guide.docx"
    deliverable = await build_word_doc(
        steps=steps,
        output_path=doc_path,
        video_title=video_title,
        quality_report=quality_report,
        video_url=args.url,
        tools_info=tools_info,
        contact_info=CONTACT_INFO,
    )
    logger.info("DONE — Deliverable: %s", deliverable.resolve())


async def vision_only_mode(args: argparse.Namespace) -> None:
    """Re-run screenshot vision from steps_pre_vision.json (keeps video + transcript)."""
    output_dir = Path(args.output_dir)
    setup_logging(output_dir, args.verbose)
    validate_config()

    video_path = output_dir / "video.mp4"
    if not video_path.exists():
        logger.error("video.mp4 not found in %s. Run the full pipeline first.", output_dir)
        sys.exit(1)

    steps_path = output_dir / "steps_pre_vision.json"
    if not steps_path.exists():
        steps_path = output_dir / "steps.json"
    if not steps_path.exists():
        logger.error("steps_pre_vision.json / steps.json not found in %s.", output_dir)
        sys.exit(1)

    with open(steps_path, encoding="utf-8") as f:
        raw_steps = json.load(f)
    steps = [
        InstallationStep.model_validate(s).model_copy(
            update={
                "screenshot_path": None,
                "selected_frame_label": None,
                "selected_frame_reason": None,
                "selected_frame_confidence": None,
                "screenshot_timestamp": None,
            },
        )
        for s in raw_steps
    ]

    scenes_path = output_dir / "scenes.json"
    scene_cuts: list[SceneCut] = []
    if scenes_path.exists():
        with open(scenes_path, encoding="utf-8") as f:
            scene_cuts = [SceneCut.model_validate(s) for s in json.load(f)]

    video_duration = await get_duration(video_path)
    youtube_url = normalize_youtube_url(args.url)

    for folder in ("screenshots", "frame_candidates"):
        p = output_dir / folder
        if p.exists():
            shutil.rmtree(p)

    logger.info(
        "=== VISION-ONLY mode: re-picking screenshots for %d steps (%s) ===",
        len(steps), output_dir,
    )

    steps = await select_best_screenshots(
        video_path=video_path,
        steps=steps,
        scene_cuts=scene_cuts,
        video_duration=video_duration,
        output_dir=output_dir,
    )

    if not getattr(args, "skip_dedup", False):
        steps = await fix_consecutive_duplicate_screenshots(
            video_path=video_path,
            steps=steps,
            scene_cuts=scene_cuts,
            video_duration=video_duration,
            output_dir=output_dir,
        )

    steps = await refine_low_confidence_screenshots(
        video_path=video_path,
        steps=steps,
        video_duration=video_duration,
        output_dir=output_dir,
    )

    save_json([s.model_dump() for s in steps], output_dir / "steps.json", "final steps")

    quality_report = build_quality_report(steps)
    save_json(quality_report, output_dir / "quality_report.json", "quality report")

    tools_info: dict = {}
    tools_path = output_dir / "tools_info.json"
    if tools_path.exists():
        with open(tools_path, encoding="utf-8") as f:
            tools_info = json.load(f)
    else:
        segments_path = output_dir / "transcript.json"
        if segments_path.exists():
            with open(segments_path, encoding="utf-8") as f:
                segments = [TranscriptSegment.model_validate(s) for s in json.load(f)]
            tools_info = await extract_tools_info_with_ai(segments)
            save_json(tools_info, tools_path, "tools info")

    if getattr(args, "title", None):
        video_title = args.title
    else:
        video_title = read_video_title_from_info_json(output_dir) or "Installation Guide"

    doc_path = output_dir / "installation_guide.docx"
    deliverable = await build_word_doc(
        steps=steps,
        output_path=doc_path,
        video_title=video_title,
        quality_report=quality_report,
        video_url=youtube_url,
        tools_info=tools_info,
        contact_info=CONTACT_INFO,
    )
    logger.info("DONE — Deliverable: %s", deliverable.resolve())
    if quality_report.get("recommended_manual_review") and not CLIENT_DEMO_MODE:
        logger.warning(
            "Manual review recommended: %d low-confidence, %d fallback, %d missing.",
            quality_report.get("low_confidence_count", 0),
            quality_report.get("fallback_count", 0),
            quality_report.get("screenshots_missing", 0),
        )


# ─────────────────────────────────────────────
# PROGRAMMATIC RUNNER (CLI + HTTP API)
# ─────────────────────────────────────────────

async def run_full_pipeline(
    youtube_url: str,
    output_dir: Path,
    *,
    skip_dedup: bool = False,
    title: Optional[str] = None,
    prompt_overrides: Optional[PipelinePromptOverrides] = None,
    verbose: bool = False,
    fresh: bool = False,
    video_file: Optional[str | Path] = None,
) -> PipelineResult:
    """
    Run the full YouTube → Word pipeline. Used by CLI and api_server.py.
    Raises on failure; returns paths to the finished document.
    """
    if fresh and output_dir.exists():
        shutil.rmtree(output_dir)

    setup_logging(output_dir, verbose)
    validate_config()
    youtube_url = normalize_youtube_url(youtube_url)

    overrides = prompt_overrides or PipelinePromptOverrides()
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if video_file:
        video_path = await ingest_local_video(Path(video_file), output_dir)
    else:
        video_path = await download_video(youtube_url, output_dir)
    video_path = await normalize_video_for_processing(video_path)

    video_title = title or await get_video_title(youtube_url, output_dir)
    video_duration = await get_duration(video_path)

    scene_cuts, segments = await asyncio.gather(
        detect_scenes(video_path),
        transcribe_audio(video_path, output_dir, video_duration=video_duration),
    )

    save_json([s.model_dump() for s in scene_cuts], output_dir / "scenes.json", "scenes")
    save_json([s.model_dump() for s in segments], output_dir / "transcript.json", "transcript")

    if not segments:
        raise RuntimeError("No transcript segments returned.")

    steps = await extract_steps_with_ai(
        segments,
        step_prompt_extra=overrides.step_prompt_extra,
        video_duration=video_duration,
    )
    if not steps:
        raise RuntimeError("Claude returned no installation steps.")

    save_json(
        [s.model_dump() for s in steps],
        output_dir / "steps_pre_vision.json",
        "steps before vision",
    )

    steps = await select_best_screenshots(
        video_path=video_path,
        steps=steps,
        scene_cuts=scene_cuts,
        video_duration=video_duration,
        output_dir=output_dir,
    )

    if not skip_dedup:
        steps = await fix_consecutive_duplicate_screenshots(
            video_path=video_path,
            steps=steps,
            scene_cuts=scene_cuts,
            video_duration=video_duration,
            output_dir=output_dir,
        )

    steps = await refine_low_confidence_screenshots(
        video_path=video_path,
        steps=steps,
        video_duration=video_duration,
        output_dir=output_dir,
    )

    save_json([s.model_dump() for s in steps], output_dir / "steps.json", "final steps")

    quality_report = build_quality_report(steps)
    save_json(quality_report, output_dir / "quality_report.json", "quality report")

    tools_info = await extract_tools_info_with_ai(
        segments, tools_prompt_extra=overrides.tools_prompt_extra,
    )
    save_json(tools_info, output_dir / "tools_info.json", "tools info")

    doc_path = output_dir / "installation_guide.docx"
    deliverable = await build_word_doc(
        steps=steps,
        output_path=doc_path,
        video_title=video_title,
        quality_report=quality_report,
        video_url=youtube_url,
        tools_info=tools_info,
        contact_info=CONTACT_INFO,
        important_note_override=overrides.important_note,
    )

    return PipelineResult(
        output_dir=str(output_dir.resolve()),
        deliverable_path=str(deliverable.resolve()),
        video_title=video_title,
        video_url=youtube_url,
        step_count=len(steps),
        quality_report=quality_report,
    )


# ─────────────────────────────────────────────
# CLI / MAIN
# ─────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Convert a YouTube installation video to a Word document (template style).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  Full pipeline:\n"
            "    python v14.py https://youtu.be/XYZ output --verbose\n\n"
            "  Rebuild document only (fast):\n"
            "    python v14.py https://youtu.be/XYZ output --doc-only\n\n"
            "  Override model:\n"
            "    python v14.py https://youtu.be/XYZ output "
            "--model claude-opus-4-20250514\n\n"
            "  Use real logo (PowerShell):\n"
            "    $env:LOGO_PATH='assets/logo.png'; python v14.py ...\n"
            "  Use real logo (bash):\n"
            "    LOGO_PATH=assets/logo.png python v14.py ...\n"
        ),
    )
    p.add_argument("url",        help="YouTube URL")
    p.add_argument("output_dir", nargs="?", default="output", help="Output directory")
    p.add_argument("--verbose",  "-v", action="store_true", help="Show DEBUG logs on console")
    p.add_argument("--fresh",    action="store_true", help="Delete output directory before running")
    p.add_argument("--skip-dedup", action="store_true",
                   help="Skip consecutive duplicate screenshot detection")
    p.add_argument("--doc-only", action="store_true",
                   help="Rebuild document from existing steps.json (skip video processing)")
    p.add_argument("--vision-only", action="store_true",
                   help="Re-run OpenAI screenshot selection from steps_pre_vision.json")
    p.add_argument("--model", default=None,
                   help="Override CLAUDE_MODEL (e.g. claude-opus-4-20250514)")
    p.add_argument("--title", default=None,
                   help="Override video title in the document (useful with --doc-only)")
    p.add_argument(
        "--video-file",
        default=None,
        metavar="PATH",
        help="Skip yt-dlp: use this local .mp4 (copied to output_dir/video.mp4)",
    )
    return p.parse_args()


async def main_async() -> None:
    global CLAUDE_MODEL

    args = parse_args()

    if args.model:
        CLAUDE_MODEL = args.model
        logger.debug("Claude model overridden via --model flag: %s", CLAUDE_MODEL)

    output_dir = Path(args.output_dir)

    if getattr(args, "doc_only", False):
        await doc_only_mode(args)
        return

    if getattr(args, "vision_only", False):
        await vision_only_mode(args)
        return

    logger.info("=" * 70)
    logger.info("YouTube -> Word Doc Pipeline  (installation_guide_template_style layout)")
    if OPENAI_ENABLE_MODEL_FALLBACK:
        logger.info(
            "Transcription  : OpenAI %s → %s",
            OPENAI_TRANSCRIBE_MODEL,
            " → ".join(OPENAI_TRANSCRIBE_MODELS[1:]) or "(no fallback)",
        )
        logger.info("Step AI        : Claude %s", CLAUDE_MODEL)
        logger.info(
            "Vision AI      : OpenAI %s → %s",
            OPENAI_VISION_MODEL,
            " → ".join(OPENAI_VISION_MODELS[1:]) or "(no fallback)",
        )
    else:
        logger.info("Transcription  : OpenAI %s", OPENAI_TRANSCRIBE_MODEL)
        logger.info("Step AI        : Claude %s", CLAUDE_MODEL)
        logger.info("Vision AI      : OpenAI %s", OPENAI_VISION_MODEL)
    logger.info("Candidates     : %d normal, %d rescue, %d extended-fallback",
                MAX_CANDIDATE_FRAMES, WIDE_RESCUE_FRAMES, EXTENDED_FALLBACK_FRAMES)
    logger.info(
        "Quality filter : sharpness>=%.0f  brightness %.0f-%.0f  color_cast<=%.2f",
        SHARPNESS_MIN, BRIGHTNESS_MIN, BRIGHTNESS_MAX, COLOR_CAST_MAX_RATIO,
    )
    logger.info("Proto-banner   : check=%s top=%.0f%% mean<%.0f std>%.0f",
                PROTOTYPE_BANNER_CHECK,
                PROTOTYPE_BANNER_TOP_FRACTION * 100,
                PROTOTYPE_BANNER_MEAN_MAX,
                PROTOTYPE_BANNER_CONTRAST_THRESHOLD)
    logger.info(
        "Confirmation   : enabled=%s threshold=%.2f  floor=%.2f  max_drop=%.2f",
        ENABLE_VISION_CONFIRMATION,
        VISION_CONFIRM_THRESHOLD, VISION_CONFIRM_FLOOR, VISION_CONFIRM_MAX_DROP,
    )
    logger.info("Partial match  : min_conf=%.2f", PARTIAL_MATCH_MIN_CONFIDENCE)
    logger.info(
        "Transcript cap : window<=%.0fs  segment_span<=%.0fs  outro_excl=%.0fs",
        MAX_TRANSCRIPT_WINDOW_SECONDS,
        MAX_SOURCE_SEGMENT_SPAN_SECONDS,
        VIDEO_OUTRO_EXCLUSION_SECONDS,
    )
    logger.info(
        "Vision shortlist: %d frames/step  refine<%s conf=%.2f (%d frames)",
        VISION_SHORTLIST_FRAMES,
        ENABLE_LOW_CONFIDENCE_REFINE,
        LOW_CONFIDENCE_REFINE_THRESHOLD,
        LOW_CONFIDENCE_REFINE_FRAMES,
    )
    logger.info("Outro handling : lookback=%.0fs  title_filter=%d keywords",
                OUTRO_STEP_LOOKBACK_SECONDS, len(OUTRO_STEP_TITLE_KEYWORDS))
    logger.info("De-bias TS     : %s", DEBIAS_TIMESTAMPS)
    logger.info("Min step gap   : %.0fs (enforced post-extraction)", MIN_STEP_GAP_SECONDS)
    logger.info("Client demo    : %s (stricter vision, no manual-review flag)", CLIENT_DEMO_MODE)
    logger.info(
        "Reinstall shift: +%.0fs when gap < %.0fs",
        REINSTALL_EXTRA_FORWARD_SEARCH, REINSTALL_TRIGGER_GAP_SECONDS,
    )
    logger.info("Multi-person   : rescue frames=%d, min conf=%.2f",
                MULTI_PERSON_RESCUE_FRAMES, MULTI_PERSON_RESCUE_CONFIDENCE)
    logger.info("Dedup sim max  : %.2f", CONSECUTIVE_FRAME_SIMILARITY_MAX)
    logger.info("Concurrency    : %d", SCREENSHOT_CONCURRENCY)
    logger.info(
        "Frame pipeline : candidates %dpx wide \u2192 finals %dpx  accurate_seek=%s  border=%dpx",
        FRAME_EXTRACT_WIDTH,
        FINAL_SCREENSHOT_WIDTH,
        FRAME_FINAL_ACCURATE_SEEK,
        SCREENSHOT_BORDER_PX,
    )
    logger.info("Sample interval: %.2fs between candidate frames", VISUAL_SAMPLE_INTERVAL)
    logger.info("Logo           : %s", LOGO_PATH if LOGO_PATH else "(generated placeholder)")
    logger.info("Logo width     : %.1f inches", LOGO_WIDTH_INCHES)
    logger.info(
        "Deliverable dir: %s",
        DELIVERABLE_OUTPUT_DIR if DELIVERABLE_OUTPUT_DIR else "(working dir only)",
    )
    logger.info("Contact info   : %s", {k: v for k, v in CONTACT_INFO.items() if v})
    logger.info("URL            : %s", args.url)
    logger.info("Output         : %s", output_dir.resolve())
    logger.info("=" * 70)

    try:
        result = await run_full_pipeline(
            args.url,
            output_dir,
            skip_dedup=getattr(args, "skip_dedup", False),
            title=getattr(args, "title", None),
            verbose=args.verbose,
            fresh=args.fresh,
            video_file=getattr(args, "video_file", None),
        )

        logger.info("=" * 70)
        logger.info("DONE")
        logger.info("Deliverable     : %s", result.deliverable_path)
        logger.info("Steps           : %d", result.step_count)
        logger.info("Avg confidence  : %.2f", result.quality_report.get("average_confidence", 0.0))
        logger.info("Quality report  : %s", (output_dir / "quality_report.json").resolve())
        logger.info("Log             : %s", (output_dir / "pipeline.log").resolve())
        logger.info("=" * 70)

        if result.quality_report.get("recommended_manual_review") and not CLIENT_DEMO_MODE:
            logger.warning(
                "Manual review recommended: %d low-confidence, %d fallback, %d missing.",
                result.quality_report.get("low_confidence_count", 0),
                result.quality_report.get("fallback_count", 0),
                result.quality_report.get("screenshots_missing", 0),
            )

    except subprocess.CalledProcessError as exc:
        logger.error("Subprocess failed: %s", " ".join(str(c) for c in exc.cmd))
        logger.error("Return code: %s", exc.returncode)
        logger.error("stderr: %s", exc.stderr)
        sys.exit(1)
    except KeyboardInterrupt:
        logger.warning("Interrupted.")
        sys.exit(130)
    except Exception as exc:
        logger.exception("Pipeline failed: %s", exc)
        sys.exit(1)


def main() -> None:
    asyncio.run(main_async())


if __name__ == "__main__":
    main()